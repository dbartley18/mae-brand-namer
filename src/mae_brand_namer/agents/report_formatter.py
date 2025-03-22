#!/usr/bin/env python3
"""
Report Formatter

It pulls raw data from the report_raw_data table and formats it into a polished report.
"""

import os
import re
import json
import logging
import asyncio
import time
import traceback
from datetime import datetime
from typing import Dict, Any, List, Optional, Tuple, Any
from pathlib import Path

import docx
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from langchain.prompts import PromptTemplate, load_prompt
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.schema import SystemMessage, HumanMessage
from pydantic import ValidationError

from ..models.report_sections import (
    BrandContext,
    NameGenerationSection,
    BrandName,
    TOCSection,
    TableOfContentsSection,
    LinguisticAnalysis,
    LinguisticAnalysisDetails,
    SemanticAnalysis,
    SemanticAnalysisDetails,
    CulturalSensitivityAnalysis,
    BrandAnalysis,
    TranslationAnalysis,
    LanguageAnalysis,
    BrandNameEvaluation,
    EvaluationDetails,
    EvaluationLists,
    DomainAnalysis,
    DomainDetails,
    SEOOnlineDiscoverability,
    SEOOnlineDiscoverabilityDetails,
    SEORecommendations,
    CompetitorAnalysis,
    CompetitorDetails,
    BrandCompetitors,
    MarketResearch,
    MarketResearchDetails,
    SurveySimulation,
    SurveyDetails
)
from ..utils.supabase_utils import SupabaseManager
from ..utils.logging import get_logger
from ..config.settings import settings
from .prompts.report_formatter import (
    get_title_page_prompt,
    get_toc_prompt,
    get_executive_summary_prompt,
    get_recommendations_prompt,
    get_system_prompt
)

from ..utils.logging import get_logger

# Define path to prompts directory
PROMPTS_DIR = os.path.join(Path(__file__).parent, "prompts", "report_formatter")

logger = get_logger(__name__)

def _safe_load_prompt(path: str) -> PromptTemplate:
    """
    Safely load a prompt template from a file or use a default one.
    
    Args:
        path: Path to the prompt template file
        
    Returns:
        PromptTemplate: The loaded prompt template or a default one if loading fails
    """
    try:
        # Only load necessary templates for brand_context and executive_summary
        if any(section in path for section in ["brand_context", "executive_summary"]):
            logger.debug(f"Loading prompt template from {path}")
            with open(path, "r") as f:
                content = f.read().strip()
                
            # Extract template variables
            variables = set(re.findall(r"{([^{}]*)}", content))
            
            logger.debug(f"Loaded prompt template with variables: {variables}")
            return PromptTemplate(template=content, input_variables=list(variables))
        else:
            # For other sections, return a minimal template that won't be used
            logger.debug(f"Skipping prompt template loading for {path} - using ETL processing")
            return PromptTemplate(template="ETL processing used instead of LLM", input_variables=["dummy"])
    except Exception as e:
        logger.warning(f"Error loading prompt template from {path}: {str(e)}")
        logger.warning("Using default prompt template")
        return PromptTemplate(
            template="Please format the following data for a section of a brand naming report: {data}",
            input_variables=["data"]
        )

class ReportFormatter:
    """
    Handles formatting and generation of reports using raw data from the report_raw_data table.
    This is the second step in the two-step report generation process.
    """

    # Mapping between DB section names and formatter section names
    SECTION_MAPPING = {
        # DB section name -> Formatter section name (as per notepad.md)
        "brand_context": "Brand Context",
        "brand_name_generation": "Name Generation",
        "linguistic_analysis": "Linguistic Analysis",
        "semantic_analysis": "Semantic Analysis",
        "cultural_sensitivity_analysis": "Cultural Sensitivity",
        "translation_analysis": "Translation Analysis",
        "survey_simulation": "Survey Simulation",
        "brand_name_evaluation": "Name Evaluation",
        "domain_analysis": "Domain Analysis",
        "seo_online_discoverability": "SEO Analysis",
        "competitor_analysis": "Competitor Analysis",
        "market_research": "Market Research",
        "exec_summary": "Executive Summary",
        "final_recommendations": "Strategic Recommendations"
    }

    # Reverse mapping for convenience
    REVERSE_SECTION_MAPPING = {
        # Formatter section name -> DB section name
        "brand_context": "brand_context",
        "brand_name_generation": "brand_name_generation",
        "linguistic_analysis": "linguistic_analysis",
        "semantic_analysis": "semantic_analysis", 
        "cultural_sensitivity_analysis": "cultural_sensitivity_analysis",
        "translation_analysis": "translation_analysis",
        "survey_simulation": "survey_simulation",
        "brand_name_evaluation": "brand_name_evaluation",
        "domain_analysis": "domain_analysis",
        "seo_analysis": "seo_online_discoverability",
        "competitor_analysis": "competitor_analysis",
        "market_research": "market_research",
        "executive_summary": "exec_summary",
        "recommendations": "final_recommendations"
    }

    # Default storage bucket for reports
    STORAGE_BUCKET = "agent_reports"
    
    # Report formats
    FORMAT_DOCX = "docx"
    
    # Section order for report generation (in exact order from notepad.md)
    SECTION_ORDER = [
        "exec_summary",
        "brand_context",
        "brand_name_generation",
        "linguistic_analysis",
        "semantic_analysis",
        "cultural_sensitivity_analysis",
        "translation_analysis",
        "brand_name_evaluation", 
        "domain_analysis",
        "seo_online_discoverability",
        "competitor_analysis",
        "market_research",
        "survey_simulation",
        "final_recommendations"
    ]
    
    def __init__(self, run_id: str, dependencies=None, supabase: SupabaseManager = None):
        """
        Initialize the ReportFormatter.
        
        Args:
            run_id: The ID of the run to fetch data for
            dependencies: Optional dependencies to inject
            supabase: Optional Supabase manager
        """
        # Store run ID
        self.current_run_id = run_id
        
        # Initialize dependencies
        if dependencies is None:
            dependencies = {}
        
        # Initialize Supabase connection
        if supabase:
            self.supabase = supabase
        elif hasattr(dependencies, "supabase"):
            self.supabase = dependencies.supabase
        else:
            self.supabase = SupabaseManager()
        
        # Initialize Gemini model with tracing 
        if hasattr(dependencies, "llm"):
            self.llm = dependencies.llm
        else:
            from ..config.settings import settings
            self.llm = ChatGoogleGenerativeAI(
                model=settings.model_name, 
                temperature=0.2,  # Balanced temperature for analysis 
                google_api_key=settings.google_api_key, 
                convert_system_message_to_human=True, 
                callbacks=settings.get_langsmith_callbacks()
            )
        
        # Set up prompt templates
        self.prompts = {
            # Only load templates for brand_context and executive_summary
            "brand_context": _safe_load_prompt(os.path.join(PROMPTS_DIR, "brand_context.yaml")),
            "exec_summary": _safe_load_prompt(os.path.join(PROMPTS_DIR, "executive_summary.yaml")),
            "executive_summary": _safe_load_prompt(os.path.join(PROMPTS_DIR, "executive_summary.yaml"))
        }
        
        # Log available prompt templates
        logger.debug(f"Loaded {len(self.prompts)} prompt templates: {list(self.prompts.keys())}")
        
        # Initialize error tracking
        self.formatting_errors = {}
        self.missing_sections = set()
        
        # Initialize current run ID
        self.current_run_id = run_id
        
        # Create transformers mapping
        self.transformers = {
            "brand_context": self._transform_brand_context,
            "brand_name_generation": self._transform_name_generation,
            "semantic_analysis": self._transform_semantic_analysis,
            "linguistic_analysis": self._transform_linguistic_analysis,
            "cultural_sensitivity_analysis": self._transform_cultural_sensitivity,
            "brand_name_evaluation": self._transform_name_evaluation,
            "translation_analysis": self._transform_translation_analysis,
            "market_research": self._transform_market_research,
            "competitor_analysis": self._transform_competitor_analysis,
            "domain_analysis": self._transform_domain_analysis,
            "survey_simulation": self._transform_survey_simulation,
            "seo_online_discoverability": self._transform_seo_analysis,
            # Add transformers as needed
        }
        
        logger.info("Initializing ReportFormatter for run_id: %s", run_id)
    def _get_format_instructions(self, section_name: str) -> str:
        """
        Get formatting instructions for a section.
        Only used for brand_context and executive_summary sections.
        
        Args:
            section_name: The name of the section
            
        Returns:
            The formatting instructions
        """
        # Only brand_context and executive_summary use format instructions
        if section_name not in ["brand_context", "executive_summary"]:
            return "ETL processing used instead"
            
        # Base instructions for all sections
        base_instructions = """
        Format the provided data into a well-structured, professionally written report section.
        The output should be a JSON object containing fields appropriate for this section.
        Keep the tone professional, informative, and engaging.
        """
        
        # Section-specific instructions
        if section_name == "brand_context":
            return base_instructions + """
            Include the following fields in your JSON output:
            - "overview": A comprehensive overview of the brand context
            - "industry": Analysis of the industry context
            - "target_audience": Description of the target audience
            - "brand_values": Core brand values and personality
            - "positioning": Brand positioning statement and strategy
            - "brand_voice": Tone and voice guidelines
            - "naming_objectives": Specific goals for the naming project
            """
        elif section_name == "executive_summary":
            return base_instructions + """
            Include the following fields in your JSON output:
            - "project_overview": Brief overview of the brand naming project
            - "process_summary": Summary of the naming process and methodology
            - "key_findings": The most important findings from the analysis
            - "recommended_names": The top recommended brand name options
            - "rationale": Explanation of why these names are recommended
            - "next_steps": Suggested next steps in the brand naming process
            """
        else:
            return base_instructions

    def _format_template(self, template_name: str, format_data: Dict[str, Any], section_name: str = None) -> str:
        """
        Format a template with the given data.
        
        Args:
            template_name: The name of the template to format
            format_data: The data to format the template with
            section_name: The section name to use for logging
            
        Returns:
            The formatted template
        """
        # Only brand_context and executive_summary use templates
        if section_name not in ["brand_context", "executive_summary"]:
            logger.debug(f"Skipping template formatting for {section_name} - using ETL processing")
            return "ETL processing used instead"
            
        try:
            if template_name in self.prompts:
                # Format the template
                template = self.prompts[template_name]
                
                # Include format instructions if available and not in the data
                if "format_instructions" not in format_data and section_name:
                    format_data["format_instructions"] = self._get_format_instructions(section_name)
                
                # Check for missing variables
                missing_vars = []
                for var in template.input_variables:
                    if var not in format_data:
                        missing_vars.append(var)
                        format_data[var] = f"[No data available for {var}]"
                
                if missing_vars:
                    logger.warning(f"Missing variables for template {template_name}: {missing_vars}")
                
                # Format the template
                return template.format(**format_data)
            else:
                # Return a default prompt if template not found
                logger.warning(f"Template {template_name} not found in prompts dictionary")
                return f"Please format the following {template_name} data: {json.dumps(format_data)}"
        except Exception as e:
            logger.error(f"Error formatting template {template_name}: {str(e)}")
            return f"Error formatting template: {str(e)}"

    def _get_system_content(self, fallback: str = None) -> str:
        """Get system prompt content with fallback.
        
        Args:
            fallback: Fallback content if system prompt not found
            
        Returns:
            System prompt content
        """
        if "system" in self.prompts:
            if hasattr(self.prompts["system"], "template"):
                return self.prompts["system"].template
            else:
                return self.prompts["system"].get("template", "")
        return fallback or "You are an expert report formatter."
            
    async def _safe_llm_invoke(self, messages, section_name=None, fallback_response=None):
        """Safe wrapper for LLM invocation with error handling.
        
        Args:
            messages: The messages to send to the LLM
            section_name: Optional name of the section being processed (for logging)
            fallback_response: Optional fallback response to return if invocation fails
            
        Returns:
            The LLM response or fallback response on failure
        """
        context = f" for section {section_name}" if section_name else ""
        max_retries = 2
        
        for attempt in range(max_retries):
            try:
                return await self.llm.ainvoke(messages)
            except Exception as e:
                logger.error(f"Error during LLM call{context} (attempt {attempt+1}/{max_retries}): {str(e)}")
                if attempt == max_retries - 1:
                    logger.error(f"All LLM retry attempts failed{context}. Using fallback.")
                    
                    if fallback_response:
                        return fallback_response
                    
                    # Create a simple response object that mimics the LLM response structure
                    class FallbackResponse:
                        def __init__(self, content):
                            self.content = content
                    
                    error_msg = f"LLM processing failed after {max_retries} attempts. Error: {str(e)}"
                    return FallbackResponse(json.dumps({
                        "title": f"Error Processing {section_name if section_name else 'Content'}",
                        "content": error_msg,
                        "sections": [{"heading": "Error Details", "content": error_msg}]
                    }))

    def _transform_brand_context(self, data: Dict[str, Any]) -> Dict[str, Any]:
        if not data:
            return {}
        try:
            # Check if data has nested 'brand_context' key and extract it if so
            brand_context_data = data.get('brand_context', data)
            
            # Log what we're trying to validate
            logger.debug(f"Validating brand context: {str(brand_context_data)[:200]}...")
            
            # Validate against the model
            brand_context_model = BrandContext.model_validate(brand_context_data)
            return brand_context_model.model_dump()
        except ValidationError as e:
            logger.error(f"Validation error for brand context data: {str(e)}")
            # Return the original data as fallback
            return data.get('brand_context', data)

    def _transform_name_generation(self, data: Dict[str, Any]) -> Any:
        """
        Transform brand name generation data into a structured format for the report.
        
        This handles various formats of brand name generation data:
        1. Raw data from brand_name_generation.md which has:
           - Categories as top-level keys
           - Lists of brand names as values for each category
           - Each brand name entry has attributes like rationale, brand_promise_alignment, etc.
        
        2. Data already structured as a NameGenerationSection model with categories
        
        3. Data with a nested 'brand_name_generation' key containing any of the above formats
        
        Returns:
            A structured dictionary matching the NameGenerationSection model structure
        """
        try:
            # Log original data structure
            if isinstance(data, dict):
                logger.debug(f"Name generation data keys: {list(data.keys())}")
                
                # Check if data is nested under 'brand_name_generation'
                if "brand_name_generation" in data:
                    data = data["brand_name_generation"]
                    logger.debug(f"Extracted nested brand_name_generation data, keys: {list(data.keys()) if isinstance(data, dict) else 'not a dict'}")
            
            # First, handle the case where we have a raw format like in brand_name_generation.md
            # where categories are top-level keys and values are lists of brand names
            if isinstance(data, dict) and all(isinstance(value, list) for value in data.values() if isinstance(value, list)):
                logger.debug("Detected raw brand name generation data format with categories as keys")
                
                # Convert to NameGenerationSection structure
                categories = []
                for category_name, names_list in data.items():
                    # Skip non-list values and special keys that aren't categories
                    if not isinstance(names_list, list) or category_name in ["methodology_and_approach", "introduction", "summary"]:
                        continue
                    
                    # Process names in this category
                    processed_names = []
                    for name_data in names_list:
                        if isinstance(name_data, dict) and "brand_name" in name_data:
                            processed_names.append(name_data)
                    
                    # Create category entry
                    if processed_names:
                        categories.append({
                            "category_name": category_name,
                            "category_description": "",  # Default empty description
                            "names": processed_names
                        })
                
                # Create structured data
                structured_data = {
                    "introduction": data.get("introduction", ""),
                    "methodology_and_approach": data.get("methodology_and_approach", ""),
                    "summary": data.get("summary", ""),
                    "categories": categories,
                    "generated_names_overview": {
                        "total_count": sum(len(cat.get("names", [])) for cat in categories)
                    },
                    "evaluation_metrics": {}  # Default empty metrics
                }
                
                logger.debug(f"Converted raw data to structured format with {len(categories)} categories")
                
                # Try to validate with NameGenerationSection model
                try:
                    from mae_brand_namer.models.report_sections import NameGenerationSection
                    validated_data = NameGenerationSection.model_validate(structured_data).model_dump()
                    logger.debug("Successfully validated structured data with NameGenerationSection model")
                    return validated_data
                except Exception as validation_error:
                    logger.warning(f"Validation of structured name generation data failed: {str(validation_error)}")
                    # Return the structured data even if validation fails
                    return structured_data
            
            # Attempt to directly validate the data if it's already in the right structure
            if isinstance(data, dict) and "categories" in data:
                try:
                    from mae_brand_namer.models.report_sections import NameGenerationSection
                    validated_data = NameGenerationSection.model_validate(data).model_dump()
                    logger.debug("Successfully validated existing structured data with NameGenerationSection model")
                    return validated_data
                except Exception as validation_error:
                    logger.warning(f"Validation of existing structured name generation data failed: {str(validation_error)}")
            
            # Try to parse from JSON directly
            if isinstance(data, str):
                try:
                    parsed_data = json.loads(data)
                    return self._transform_name_generation(parsed_data)  # Recursively process the parsed data
                except json.JSONDecodeError:
                    logger.warning("Failed to parse name generation data as JSON string")
            
            # If still here, use NameGenerationSection.from_raw_json as fallback
            try:
                from mae_brand_namer.models.report_sections import NameGenerationSection
                section = NameGenerationSection.from_raw_json(data)
                logger.debug("Successfully created NameGenerationSection from raw JSON data")
                return section.model_dump()
            except Exception as e:
                logger.warning(f"Failed to create NameGenerationSection from raw JSON: {str(e)}")
            
            # Last resort: return the original data
            logger.debug("Returning original name generation data as no transformation succeeded")
            return data
            
        except Exception as e:
            logger.error(f"Error transforming name generation data: {str(e)}")
            logger.debug(f"Error details: {traceback.format_exc()}")
            return data

    def _transform_semantic_analysis(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """
        Transform semantic analysis data to match expected model format.
        
        The semantic analysis data can be structured in several ways:
        1. Under a "semantic_analysis" key in the top-level data
        2. As a dictionary of brand names with semantic analysis details
        3. As a nested structure under semantic_analysis.semantic_analysis
        """
        if not data:
            return {}
        try:
            # First check for semantic_analysis key
            semantic_data = data.get("semantic_analysis", data)
            
            # Log what we're trying to validate
            logger.debug(f"Validating semantic analysis data: {str(semantic_data)[:200]}...")
            
            # If semantic_data exists but is nested one level deeper
            if isinstance(semantic_data, dict) and "semantic_analysis" in semantic_data:
                semantic_data = semantic_data["semantic_analysis"]
                
            # Create the expected structure
            semantic_analysis = {"semantic_analysis": semantic_data}
            
            # Validate against the model
            validated_data = SemanticAnalysis.model_validate(semantic_analysis)
            
            # Transform the data for the template format
            template_data = {
                "brand_analyses": []
            }
            
            # Convert dictionary of brand details to a list of brand analyses
            for brand_name, details in validated_data.semantic_analysis.items():
                # Make sure brand_name is included in the details
                brand_analysis = details.model_dump()
                # Ensure brand_name is present
                if "brand_name" not in brand_analysis:
                    brand_analysis["brand_name"] = brand_name
                    
                template_data["brand_analyses"].append(brand_analysis)
                
            logger.info(f"Successfully transformed semantic analysis data with {len(template_data['brand_analyses'])} brand analyses")
            return template_data
            
        except ValidationError as e:
            logger.error(f"Validation error for semantic analysis data: {str(e)}")
            # Return the original data as fallback
            return data.get('semantic_analysis', data)

    def _transform_linguistic_analysis(self, data: Dict[str, Any]) -> Dict[str, Any]:
        if not data:
            return {}
        try:
            linguistic_analysis_data = LinguisticAnalysis.model_validate(data)
            return linguistic_analysis_data.model_dump()
        except ValidationError as e:
            logger.error(f"Validation error for linguistic analysis data: {str(e)}")
            return {}

    def _transform_cultural_sensitivity(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """
        Transform cultural sensitivity analysis data to match expected format.
        
        The cultural sensitivity analysis data should contain brand name analyses with
        attributes such as symbolic_meanings, cultural_connotations, etc.
        
        Args:
            data: Raw cultural sensitivity analysis data
            
        Returns:
            Transformed data with brand_analyses and summary fields
        """
        if not data:
            return {}
        try:
            # Validate data against the model
            cultural_sensitivity_data = CulturalSensitivityAnalysis.model_validate(data)
            
            # Transform to a format more suitable for the template
            transformed_data = {
                "brand_analyses": [],
                "summary": "The cultural sensitivity analysis evaluated brand names across multiple dimensions including symbolic meanings, cultural connotations, religious sensitivities, and regional variations. This analysis helps identify names that minimize cultural risks while maximizing positive associations across diverse markets."
            }
            
            # Process each brand analysis
            if hasattr(cultural_sensitivity_data, "cultural_sensitivity_analysis"):
                for brand_name, analysis in cultural_sensitivity_data.cultural_sensitivity_analysis.items():
                    # Create a brand analysis entry with all fields from the model
                    brand_analysis = analysis.model_dump()
                    # Ensure brand_name is included
                    if "brand_name" not in brand_analysis:
                        brand_analysis["brand_name"] = brand_name
                    
                    transformed_data["brand_analyses"].append(brand_analysis)
            
            logger.info(f"Successfully transformed cultural sensitivity data with {len(transformed_data['brand_analyses'])} brand analyses")
            return transformed_data
            
        except ValidationError as e:
            logger.error(f"Validation error for cultural sensitivity analysis data: {str(e)}")
            return {}

    def _transform_name_evaluation(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """
        Transform brand name evaluation data to match expected model format.
        
        The brand name evaluation data should have:
        1. A "brand_name_evaluation" key with shortlisted_names and other_names lists
        2. Each name evaluation contains brand_name, overall_score, shortlist_status, and evaluation_comments
        
        Args:
            data: Raw brand name evaluation data
            
        Returns:
            Transformed data in the format expected by the formatter
        """
        if not data:
            return {}
        try:
            # First, validate data against the model
            name_evaluation_data = BrandNameEvaluation.model_validate(data)
            
            # Transform to a format more suitable for the template
            transformed_data = {
                "shortlisted_names": [],
                "other_names": [],
                "evaluation_methodology": "Each brand name was evaluated based on multiple criteria including distinctiveness, memorability, strategic alignment, and cultural considerations.",
                "comparative_analysis": "The shortlisted names demonstrated stronger overall performance across evaluation criteria, particularly in areas of strategic alignment with the brand's values and positioning in the market."
            }
            
            # Process shortlisted names
            if hasattr(name_evaluation_data.brand_name_evaluation, "shortlisted_names"):
                transformed_data["shortlisted_names"] = [
                    {
                        "brand_name": item.brand_name,
                        "overall_score": item.overall_score,
                        "evaluation_comments": item.evaluation_comments
                    }
                    for item in name_evaluation_data.brand_name_evaluation.shortlisted_names
                ]
                logger.info(f"Processed {len(transformed_data['shortlisted_names'])} shortlisted names")
            
            # Process other names
            if hasattr(name_evaluation_data.brand_name_evaluation, "other_names"):
                transformed_data["other_names"] = [
                    {
                        "brand_name": item.brand_name,
                        "overall_score": item.overall_score,
                        "evaluation_comments": item.evaluation_comments
                    }
                    for item in name_evaluation_data.brand_name_evaluation.other_names
                ]
                logger.info(f"Processed {len(transformed_data['other_names'])} other names")
                
            # Add final rankings based on scores
            all_names = transformed_data["shortlisted_names"] + transformed_data["other_names"]
            rankings = {
                item["brand_name"]: item["overall_score"] for item in all_names
            }
            transformed_data["final_rankings"] = rankings
            
            logger.info(f"Successfully transformed name evaluation data with {len(transformed_data['shortlisted_names'])} shortlisted names and {len(transformed_data['other_names'])} other names")
            logger.debug(f"Transformed data: {transformed_data}")
            return transformed_data
            
        except ValidationError as e:
            logger.error(f"Validation error for name evaluation data: {str(e)}")
            
            # Try a fallback approach if the validation fails
            try:
                # Check if the data is already in the expected format
                if "brand_name_evaluation" in data and isinstance(data["brand_name_evaluation"], dict):
                    eval_lists = data["brand_name_evaluation"]
                    
                    transformed_data = {
                        "shortlisted_names": [],
                        "other_names": [],
                        "evaluation_methodology": "Each brand name was evaluated based on multiple criteria including distinctiveness, memorability, strategic alignment, and cultural considerations.",
                        "comparative_analysis": "The shortlisted names demonstrated stronger overall performance across evaluation criteria, particularly in areas of strategic alignment with the brand's values and positioning in the market."
                    }
                    
                    # Process shortlisted names
                    if "shortlisted_names" in eval_lists and isinstance(eval_lists["shortlisted_names"], list):
                        transformed_data["shortlisted_names"] = [
                            {
                                "brand_name": item.get("brand_name", ""),
                                "overall_score": item.get("overall_score", 0),
                                "evaluation_comments": item.get("evaluation_comments", "")
                            }
                            for item in eval_lists["shortlisted_names"]
                        ]
                    
                    # Process other names
                    if "other_names" in eval_lists and isinstance(eval_lists["other_names"], list):
                        transformed_data["other_names"] = [
                            {
                                "brand_name": item.get("brand_name", ""),
                                "overall_score": item.get("overall_score", 0),
                                "evaluation_comments": item.get("evaluation_comments", "")
                            }
                            for item in eval_lists["other_names"]
                        ]
                    
                    # Add final rankings
                    all_names = transformed_data["shortlisted_names"] + transformed_data["other_names"]
                    rankings = {
                        item["brand_name"]: item["overall_score"] for item in all_names
                    }
                    transformed_data["final_rankings"] = rankings
                    
                    logger.info(f"Fallback transformation successful with {len(transformed_data['shortlisted_names'])} shortlisted names and {len(transformed_data['other_names'])} other names")
                    return transformed_data
            except Exception as fallback_error:
                logger.error(f"Fallback transformation also failed: {str(fallback_error)}")
            
            return {}

    def _transform_translation_analysis(self, data: Dict[str, Any]) -> Dict[str, Any]:
        if not data:
            return {}
        try:
            translation_analysis_data = TranslationAnalysis.model_validate(data)
            return translation_analysis_data.model_dump()
        except ValidationError as e:
            logger.error(f"Validation error for translation analysis data: {str(e)}")
            return {}

    def _transform_market_research(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """
        Transform market research data to match expected format.
        
        The market research data contains:
        1. A "market_research" key containing a list of brand analyses
        2. Each analysis has market size, industry name, trends, etc.
        
        Args:
            data: Raw market research data
            
        Returns:
            Transformed data suitable for the formatter
        """
        if not data:
            return {}

        try:
            # Extract the list of market research entries
            market_research_list = data.get("market_research", [])
            if not market_research_list:
                return {}

            # Transform into dictionary keyed by brand name
            transformed_data = {
                "market_research": {
                    entry["brand_name"]: {
                        "brand_name": entry["brand_name"],
                        "market_size": entry["market_size"],
                        "industry_name": entry["industry_name"],
                        "emerging_trends": entry["emerging_trends"],
                        "key_competitors": entry["key_competitors"],
                        "potential_risks": entry["potential_risks"],
                        "recommendations": entry["recommendations"],
                        "market_viability": entry["market_viability"],
                        "market_growth_rate": entry["market_growth_rate"],
                        "market_opportunity": entry["market_opportunity"],
                        "target_audience_fit": entry["target_audience_fit"],
                        "competitive_analysis": entry["competitive_analysis"],
                        "customer_pain_points": entry["customer_pain_points"],
                        "market_entry_barriers": entry["market_entry_barriers"]
                    }
                    for entry in market_research_list
                }
            }
            
            return transformed_data

        except (KeyError, AttributeError) as e:
            logger.error(f"Error transforming market research data: {str(e)}")
            return {}

    def _transform_competitor_analysis(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """
        Transform competitor analysis data to match expected format.
        
        The competitor analysis data contains:
        1. A "competitor_analysis" key with a list of brand competitor analyses
        2. Each brand analysis has a "brand_name" and "competitors" array
        3. Each competitor has details like name, strengths, weaknesses, etc.
        
        Args:
            data: Raw competitor analysis data
            
        Returns:
            Transformed data suitable for the formatter
        """
        if not data:
            return {}

        # Initialize transformed data structure
        transformed_data = {}

        try:
            # Try Pydantic validation first
            competitor_analysis_data = CompetitorAnalysis.model_validate(data)
            
            # Process each brand's competitor analysis
            if hasattr(competitor_analysis_data, "competitor_analysis"):
                for brand_analysis in competitor_analysis_data.competitor_analysis:
                    brand_name = brand_analysis.brand_name
                    transformed_data[brand_name] = {
                        "brand_name": brand_name,
                        "competitors": [
                            {
                                "competitor_name": competitor.competitor_name,
                                "competitor_naming_style": competitor.competitor_naming_style,
                                "competitor_keywords": competitor.competitor_keywords,
                                "competitor_positioning": competitor.competitor_positioning,
                                "competitor_strengths": competitor.competitor_strengths,
                                "competitor_weaknesses": competitor.competitor_weaknesses,
                                "competitor_differentiation_opportunity": competitor.competitor_differentiation_opportunity,
                                "differentiation_score": competitor.differentiation_score,
                                "risk_of_confusion": competitor.risk_of_confusion,
                                "target_audience_perception": competitor.target_audience_perception,
                                "competitive_advantage_notes": competitor.competitive_advantage_notes,
                                "trademark_conflict_risk": competitor.trademark_conflict_risk
                            }
                            for competitor in brand_analysis.competitors
                        ]
                    }
            return transformed_data

        except (ValidationError, AttributeError) as e:
            logger.error(f"Validation error in competitor analysis: {str(e)}")
            return {}

    def _transform_domain_analysis(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """
        Transform domain analysis data to match expected format.
        
        Args:
            data: Raw domain analysis data
            
        Returns:
            Transformed data suitable for the formatter
        """
        if not data:
            return {}

        # Initialize transformed data structure
        transformed_data = {}

        # Case 1: The expected structure with "domain_analysis" key containing a dictionary of brand analyses
        if isinstance(data, dict) and "domain_analysis" in data:
            if isinstance(data["domain_analysis"], dict):
                transformed_data = data["domain_analysis"]
        # Case 2: Already has expected structure
        elif isinstance(data, dict) and any(isinstance(v, dict) and "brand_name" in v for v in data.values()):
            transformed_data = data
        
        return transformed_data

    def _transform_survey_simulation(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """Transform survey simulation data into a structured format.
        
        The raw data has a different structure than other sections, containing a list of
        survey responses directly in the root, rather than a nested structure.
        """
        if not data:
            return {}
            
        try:
            # Log the data structure to help with debugging
            logger.debug(f"Survey simulation data type: {type(data)}")
            logger.debug(f"Survey simulation data keys: {list(data.keys()) if isinstance(data, dict) else 'Not a dictionary'}")
            
            # Check if the data is already a list of survey details
            if isinstance(data, list):
                # Data is already a list of survey details
                survey_items = data
            elif isinstance(data, dict) and "survey_simulation" in data:
                # Data follows the model structure
                survey_items = data["survey_simulation"]
            else:
                # Data might be in a different format, try to extract it
                survey_items = []
                if isinstance(data, dict):
                    # Loop through all keys to find any that might contain survey items
                    for key, value in data.items():
                        if isinstance(value, list) and value:
                            survey_items = value
                            break
                
            # Now transform each survey item
            formatted_survey_details = []
            for item in survey_items:
                try:
                    survey_detail = SurveyDetails.model_validate(item)
                    formatted_survey_details.append(survey_detail.model_dump())
                except ValidationError as e:
                    logger.error(f"Validation error for survey item: {str(e)}")
                    # Add the item as-is, even if it doesn't fully validate
                    if isinstance(item, dict):
                        formatted_survey_details.append(item)
            
            return {"survey_simulation": formatted_survey_details}
            
        except Exception as e:
            logger.error(f"Error transforming survey simulation data: {str(e)}")
            # Return the original data as a fallback
            if isinstance(data, dict):
                return data
            return {"survey_simulation": data if isinstance(data, list) else []}

    def _transform_seo_analysis(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """
        Transform SEO analysis data to match expected format.
        
        The SEO analysis data contains:
        1. A "seo_online_discoverability" key with a list of brand SEO analyses
        2. Each analysis has search volume, keyword alignment, etc.
        
        Args:
            data: Raw SEO analysis data
            
        Returns:
            Transformed data suitable for the formatter
        """
        if not data:
            return {}

        try:
            # Extract the list of SEO analysis entries
            seo_analysis_list = data.get("seo_online_discoverability", [])
            if not seo_analysis_list:
                return {}

            # Transform into dictionary keyed by brand name
            transformed_data = {
                "seo_analysis": {
                    entry["brand_name"]: {
                        "brand_name": entry["brand_name"],
                        "search_volume": entry["search_volume"],
                        "keyword_alignment": entry["keyword_alignment"],
                        "keyword_competition": entry["keyword_competition"],
                        "seo_recommendations": entry["seo_recommendations"],
                        "seo_viability_score": entry["seo_viability_score"],
                        "negative_search_results": entry["negative_search_results"],
                        "unusual_spelling_impact": entry["unusual_spelling_impact"],
                        "branded_keyword_potential": entry["branded_keyword_potential"],
                        "name_length_searchability": entry["name_length_searchability"],
                        "social_media_availability": entry["social_media_availability"],
                        "competitor_domain_strength": entry["competitor_domain_strength"],
                        "exact_match_search_results": entry["exact_match_search_results"],
                        "social_media_discoverability": entry["social_media_discoverability"],
                        "negative_keyword_associations": entry["negative_keyword_associations"],
                        "non_branded_keyword_potential": entry["non_branded_keyword_potential"],
                        "content_marketing_opportunities": entry["content_marketing_opportunities"]
                    }
                    for entry in seo_analysis_list
                }
            }
            
            return transformed_data

        except (KeyError, AttributeError) as e:
            logger.error(f"Error transforming SEO analysis data: {str(e)}")
            return {}

    def _validate_section_data(self, section_name: str, data: Any) -> Tuple[bool, List[str]]:
        """
        Validate section data for common issues and provide diagnostics.
        
        Args:
            section_name: Name of the section being validated
            data: Data to validate
            
        Returns:
            Tuple of (is_valid, list_of_issues)
        """
        issues = []
        
        # Check if data exists
        if data is None:
            issues.append(f"Section '{section_name}' data is None")
            return False, issues
            
        # Check if data is empty
        if isinstance(data, dict) and not data:
            issues.append(f"Section '{section_name}' contains an empty dictionary")
        elif isinstance(data, list) and not data:
            issues.append(f"Section '{section_name}' contains an empty list")
        elif isinstance(data, str) and not data.strip():
            issues.append(f"Section '{section_name}' contains an empty string")
            
        # Check if data is in expected format
        if section_name in ["brand_context"] and not isinstance(data, dict):
            issues.append(f"Section '{section_name}' should be a dictionary, got {type(data).__name__}")
            return False, issues
            
        if section_name in ["name_generation", "linguistic_analysis", "semantic_analysis", 
                          "cultural_sensitivity", "translation_analysis", "competitor_analysis", 
                          "domain_analysis", "market_research"]:
            # These sections should typically have arrays of items
            list_fields = [
                "name_generations", "linguistic_analyses", "semantic_analyses", 
                "cultural_analyses", "translation_analyses", "competitor_analyses", 
                "domain_analyses", "market_researches"
            ]
            
            # Check if any expected list fields are present and valid
            found_valid_list = False
            for field in list_fields:
                if isinstance(data, dict) and field in data:
                    if not isinstance(data[field], list):
                        issues.append(f"Field '{field}' in section '{section_name}' should be a list, got {type(data[field]).__name__}")
                    elif not data[field]:
                        issues.append(f"Field '{field}' in section '{section_name}' is an empty list")
                    else:
                        found_valid_list = True
                        
                        # Check if items in the list have expected format
                        invalid_items = [i for i, item in enumerate(data[field]) 
                                        if not isinstance(item, dict) or "brand_name" not in item]
                        if invalid_items:
                            issues.append(f"{len(invalid_items)} items in '{field}' missing 'brand_name' or not dictionaries")
            
            if not found_valid_list and isinstance(data, dict):
                issues.append(f"No valid list fields found in section '{section_name}' data")
        
        # For LLM-generated sections, just check they're dictionaries with some content
        if section_name in ["executive_summary", "recommendations"]:
            if not isinstance(data, dict):
                issues.append(f"LLM-generated section '{section_name}' should be a dictionary, got {type(data).__name__}")
                return False, issues
        
        # Return overall validity and list of issues
        return len(issues) == 0, issues

    async def fetch_raw_data(self, run_id: str) -> Dict[str, Any]:
        """Fetch all raw data for a specific run_id from the report_raw_data table."""
        logger.info(f"Fetching all raw data for run_id: {run_id}")
        
        try:
            # Update to use the new execute_with_retry API
            result = await self.supabase.execute_with_retry(
                "select",
                "report_raw_data",
                {
                    "run_id": f"eq.{run_id}",
                    "select": "section_name,raw_data",
                    "order": "section_name"
                }
            )
            
            logger.info(f"Query result for run_id {run_id}: {len(result) if result else 0} rows found")
            
            if not result or len(result) == 0:
                logger.warning(f"No data found for run_id: {run_id}")
                return {}
            
            # Debug log to see what's in the result
            section_names = [row['section_name'] for row in result]
            logger.info(f"Found sections: {section_names}")
            
            # Critical sections we need to make sure are included
            critical_sections = ['domain_analysis', 'competitor_analysis', 'market_research']
            for section in critical_sections:
                if section not in section_names:
                    logger.warning(f"Critical section missing from query results: {section}")
            
            # Track data quality issues
            data_quality_issues = {}
            
            # Transform results into a dictionary with section_name as keys
            sections_data = {}
            for row in result:
                db_section_name = row['section_name']
                raw_data = row['raw_data']
                
                # Debug each row
                logger.debug(f"Processing row for section: {db_section_name} (data length: {len(str(raw_data))})")
                
                # Use the DB section name as the key directly since that's what we use in SECTION_ORDER
                # This way we maintain the exact mapping between database names and our processing
                formatter_section_name = db_section_name
                
                # Make sure raw_data is a properly formatted dict
                if isinstance(raw_data, str):
                    try:
                        raw_data = json.loads(raw_data)
                        logger.debug(f"Converted string raw_data to dict for {formatter_section_name}")
                    except json.JSONDecodeError:
                        logger.warning(f"Raw data for {formatter_section_name} is a string but not valid JSON")
                
                if not isinstance(raw_data, dict) and not isinstance(raw_data, list):
                    logger.warning(f"Raw data for {formatter_section_name} is not a dict or list: {type(raw_data)}")
                    # Convert to dict to avoid crashes
                    raw_data = {"raw_content": str(raw_data)}
                
                # Apply transformation if available
                if formatter_section_name in self.transformers:
                    transformer = self.transformers[formatter_section_name]
                    try:
                        logger.info(f"Applying transformer for {formatter_section_name}")
                        transformed_data = transformer(raw_data)
                        
                        # Validate transformed data
                        is_valid, issues = self._validate_section_data(formatter_section_name, transformed_data)
                        if not is_valid or issues:
                            for issue in issues:
                                logger.warning(f"Data quality issue in {formatter_section_name}: {issue}")
                            if issues:
                                data_quality_issues[formatter_section_name] = issues
                        
                        sections_data[formatter_section_name] = transformed_data
                        logger.info(f"Successfully transformed data for {formatter_section_name}")
                    except Exception as e:
                        logger.error(f"Error transforming data for section {formatter_section_name}: {str(e)}")
                        logger.error(f"Transformation error details: {traceback.format_exc()}")
                        logger.debug(f"Raw data for failed transformation: {str(raw_data)[:500]}...")
                        # Store error but continue with other sections
                        data_quality_issues[formatter_section_name] = [f"Transformation error: {str(e)}"]
                        # Still store raw data for potential fallback handling
                        sections_data[formatter_section_name] = raw_data
                else:
                    # Use raw data if no transformer is available
                    logger.info(f"No transformer available for {formatter_section_name}, using raw data")
                    # Also validate raw data
                    is_valid, issues = self._validate_section_data(formatter_section_name, raw_data)
                    if not is_valid or issues:
                        for issue in issues:
                            logger.warning(f"Data quality issue in {formatter_section_name}: {issue}")
                        if issues:
                            data_quality_issues[formatter_section_name] = issues
                            
                    sections_data[formatter_section_name] = raw_data
            
            # Store data quality issues for later reporting
            if data_quality_issues:
                logger.warning(f"Found {len(data_quality_issues)} sections with data quality issues")
                sections_data["_data_quality_issues"] = data_quality_issues
                
            logger.info(f"Successfully processed {len(sections_data)} sections for run_id: {run_id}")
            
            # Log the keys to verify what we're returning
            logger.info(f"Returning data for sections: {sorted(list(sections_data.keys()))}")
            
            # Check for expected sections
            missing_sections = [section for section in self.SECTION_ORDER if section not in sections_data]
            if missing_sections:
                logger.warning(f"Missing expected sections: {missing_sections}")
            
            return sections_data
            
        except Exception as e:
            logger.error(f"Error fetching raw data for run_id {run_id}: {str(e)}")
            logger.error(f"Error details: {traceback.format_exc()}")
            return {}

    async def _fetch_domain_analysis_directly(self, run_id: str) -> Dict[str, Any]:
        """Fallback method to fetch domain analysis data directly from the database."""
        try:
            # Query the domain_analysis table directly
            result = await self.supabase.execute_with_retry(
                "select",
                "domain_analysis",
                {
                    "run_id": f"eq.{run_id}",
                    "select": "*"
                }
            )
            
            if not result or len(result) == 0:
                logger.warning(f"No domain analysis data found for run_id: {run_id}")
                return {}
                
            # Transform to expected structure
            transformed_data = {"domain_analysis": []}
            for item in result:
                transformed_data["domain_analysis"].append(item)
                
            return transformed_data
        except Exception as e:
            logger.error(f"Error fetching domain analysis directly: {str(e)}")
            return {}
            
    async def _fetch_competitor_analysis_directly(self, run_id: str) -> Dict[str, Any]:
        """Fallback method to fetch competitor analysis data directly from the database."""
        try:
            # Query the competitor_analysis table directly
            result = await self.supabase.execute_with_retry(
                "select",
                "competitor_analysis",
                {
                    "run_id": f"eq.{run_id}",
                    "select": "*"
                }
            )
            
            if not result or len(result) == 0:
                logger.warning(f"No competitor analysis data found for run_id: {run_id}")
                return {}
                
            # Group by brand name while preserving brand_name at top level
            by_brand = {}
            for item in result:
                brand_name = item["brand_name"]
                if brand_name not in by_brand:
                    by_brand[brand_name] = {
                        "brand_name": brand_name,  # Keep brand_name at top level
                        "competitors": []
                    }
                # Remove brand_name from competitor data since it's at top level
                competitor_data = {k: v for k, v in item.items() if k != "brand_name"}
                by_brand[brand_name]["competitors"].append(competitor_data)
            
            # Convert to list preserving the structure
            formatted_data = list(by_brand.values())
            return {"competitor_analysis": formatted_data}
            
        except Exception as e:
            logger.error(f"Error fetching competitor analysis directly: {str(e)}")
            return {}
            
    async def _fetch_market_research_directly(self, run_id: str) -> Dict[str, Any]:
        """Fallback method to fetch market research data directly from the database."""
        try:
            # Query the market_research table directly
            result = await self.supabase.execute_with_retry(
                "select",
                "market_research",
                {
                    "run_id": f"eq.{run_id}",
                    "select": "*"
                }
            )
            
            if not result or len(result) == 0:
                logger.warning(f"No market research data found for run_id: {run_id}")
                return {}
                
            # Structure data as array while preserving brand_name in each entry
            formatted_data = []
            for item in result:
                formatted_data.append(item)
            
            return {"market_research": formatted_data}
            
        except Exception as e:
            logger.error(f"Error fetching market research directly: {str(e)}")
            return {}

    async def fetch_user_prompt(self, run_id: str) -> str:
        """Fetch the user prompt from the workflow_state table for a specific run_id."""
        logger.info(f"Fetching user prompt for run_id: {run_id}")
        
        try:
            # Query the workflow_state table to get the state data
            result = await self.supabase.execute_with_retry(
                "select",
                "workflow_state",
                {
                    "run_id": f"eq.{run_id}",
                    "select": "state",
                    "order": "created_at.desc",
                    "limit": 1  # Get the most recent state
                }
            )
            
            if not result or len(result) == 0:
                logger.warning(f"No state data found for run_id: {run_id}")
                return "No user prompt available"
            
            # Extract the state data from the result
            state_data = result[0]['state']
            
            # Parse the state data if it's a string
            if isinstance(state_data, str):
                try:
                    state_data = json.loads(state_data)
                except json.JSONDecodeError:
                    logger.warning(f"State data for {run_id} is not valid JSON")
                    return "No user prompt available (invalid state format)"
            
            # Extract the user prompt from the state data
            if isinstance(state_data, dict) and 'user_prompt' in state_data:
                user_prompt = state_data['user_prompt']
                logger.info(f"Successfully fetched user prompt for run_id: {run_id}")
                return user_prompt
            else:
                logger.warning(f"User prompt not found in state data for run_id: {run_id}")
                return "No user prompt available (not in state)"
            
        except Exception as e:
            logger.error(f"Error fetching user prompt for run_id {run_id}: {str(e)}")
            logger.error(f"Error details: {traceback.format_exc()}")
            return "No user prompt available (error)"

    def _setup_document_styles(self, doc: Document) -> None:
        """Set up document styles."""
        try:
            # Add styles
            styles = doc.styles
            
            # Add a style for bullet points
            if 'List Bullet' not in styles:
                bullet_style = styles.add_style('List Bullet', 1)
                bullet_style.base_style = styles['Normal']
                bullet_style.font.size = Pt(11)
                bullet_style.paragraph_format.left_indent = Inches(0.5)
            
            # Add a style for quotes
            if 'Intense Quote' not in styles:
                quote_style = styles.add_style('Intense Quote', 1)
                quote_style.base_style = styles['Normal']
                quote_style.font.size = Pt(11)
                quote_style.font.italic = True
                quote_style.paragraph_format.left_indent = Inches(0.5)
                quote_style.paragraph_format.right_indent = Inches(0.5)
            
            # Add a style for code blocks
            if 'Code Block' not in styles:
                code_style = styles.add_style('Code Block', 1)
                code_style.base_style = styles['Normal']
                code_style.font.name = 'Courier New'
                code_style.font.size = Pt(10)
                code_style.paragraph_format.left_indent = Inches(0.5)
                code_style.paragraph_format.right_indent = Inches(0.5)
            
            # Add a style for tables
            if 'Table Grid' not in styles:
                table_style = styles.add_style('Table Grid', 1)
                table_style.base_style = styles['Table Grid']
                table_style.font.size = Pt(10)
                table_style.paragraph_format.space_before = Pt(6)
                table_style.paragraph_format.space_after = Pt(6)
                
            # Add a style for TOC entries
            if 'TOC 1' not in styles:
                toc_style = styles.add_style('TOC 1', 1)
                toc_style.base_style = styles['Normal']
                toc_style.font.size = Pt(12)
                toc_style.paragraph_format.left_indent = Inches(0.25)
                toc_style.paragraph_format.space_after = Pt(6)
                
        except Exception as e:
            logger.error(f"Error setting up document styles: {str(e)}")
            # Add a basic error message to the document
            doc.add_paragraph("Error setting up document styles. Some formatting may be incorrect.", style='Intense Quote')
    
    def _format_generic_section_fallback(self, doc: Document, section_name: str, data: Dict[str, Any]) -> None:
        """Fallback method for formatting a section when LLM or other approaches fail."""
        try:
            doc.add_paragraph(f"Section data:", style='Quote')
            
            # Add section data in a readable format
            for key, value in data.items():
                if isinstance(value, str) and value:
                    doc.add_heading(key.replace("_", " ").title(), level=2)
                    doc.add_paragraph(value)
                elif isinstance(value, list) and value:
                    doc.add_heading(key.replace("_", " ").title(), level=2)
                    for item in value:
                        if isinstance(item, str):
                            bullet = doc.add_paragraph(style='List Bullet')
                            bullet.add_run(str(item))
                        elif isinstance(item, dict):
                            # Create a subheading for each item if it has a name or title
                            item_title = item.get("name", item.get("title", item.get("brand_name", f"Item {value.index(item) + 1}")))
                            doc.add_heading(item_title, level=3)
                            
                            # Add item details
                            for sub_key, sub_value in item.items():
                                if sub_key not in ["name", "title", "brand_name"] and sub_value:
                                    p = doc.add_paragraph()
                                    p.add_run(f"{sub_key.replace('_', ' ').title()}: ").bold = True
                                    p.add_run(str(sub_value))
        except Exception as e:
            logger.error(f"Error in fallback formatting for {section_name}: {str(e)}")
            doc.add_paragraph(f"Error in fallback formatting: {str(e)}", style='Intense Quote')

    async def _format_survey_simulation(self, doc: Document, data: Dict[str, Any]) -> None:
        """
        Format the survey simulation section using direct ETL process.
        
        Args:
            doc: The document to add content to
            data: The raw survey simulation data
        """
        try:
            # Add section title
            doc.add_heading("Survey Simulation", level=1)
            
            # Add introduction
            doc.add_paragraph(
                "This section presents a simulated consumer survey that evaluates the brand name options "
                "based on target audience preferences, emotional responses, and overall appeal. "
                "The simulation helps predict how real consumers might respond to each brand name."
            )
            
            # Transform data using model
            transformed_data = self._transform_survey_simulation(data)
            logger.debug(f"Transformed survey simulation data: {len(str(transformed_data))} chars")
            
            # Check if we have analysis data - specifically looking for the survey_simulation key
            if transformed_data and isinstance(transformed_data, dict) and "survey_simulation" in transformed_data:
                survey_items = transformed_data["survey_simulation"]
                
                if not survey_items:
                    doc.add_paragraph("No survey simulation data available for this brand naming project.")
                    return
                
                # Add methodology section
                doc.add_heading("Methodology", level=2)
                doc.add_paragraph(
                    "The survey simulation represents feedback from key target personas in the "
                    "intended market. Each persona evaluation includes emotional associations, "
                    "personality fit scores, and detailed qualitative feedback on each brand name."
                )
                
                # Process each survey participant
                doc.add_heading("Survey Results by Persona", level=2)
                
                for i, item in enumerate(survey_items, 1):
                    if not isinstance(item, dict):
                        continue
                        
                    # Extract the brand name and company
                    brand_name = item.get("brand_name", f"Brand {i}")
                    company = item.get("company_name", "")
                    
                    # Add persona heading
                    heading_text = f"Persona {i}: {item.get('job_title', '')}"
                    if company:
                        heading_text += f" at {company}"
                    doc.add_heading(heading_text, level=3)
                    
                    # Add persona details if available
                    if "industry" in item:
                        p = doc.add_paragraph()
                        p.add_run("Industry: ").bold = True
                        p.add_run(str(item.get("industry", "")))
                    
                    if "seniority" in item:
                        p = doc.add_paragraph()
                        p.add_run("Seniority: ").bold = True
                        p.add_run(str(item.get("seniority", "")))
                    
                    # Add brand name evaluation
                    doc.add_heading(f"Evaluation of {brand_name}", level=4)
                    
                    # Add emotional association
                    if "emotional_association" in item:
                        p = doc.add_paragraph()
                        p.add_run("Emotional Association: ").bold = True
                        p.add_run(str(item.get("emotional_association", "")))
                    
                    # Add personality fit score
                    if "personality_fit_score" in item:
                        p = doc.add_paragraph()
                        p.add_run("Personality Fit Score: ").bold = True
                        p.add_run(f"{item.get('personality_fit_score', 0)}/10")
                    
                    # Add qualitative feedback summary
                    if "qualitative_feedback_summary" in item:
                        doc.add_heading("Qualitative Feedback Summary", level=5)
                        doc.add_paragraph(str(item.get("qualitative_feedback_summary", "")))
                    
                    # Add raw qualitative feedback
                    if "raw_qualitative_feedback" in item:
                        doc.add_heading("Detailed Feedback", level=5)
                        raw_feedback = item.get("raw_qualitative_feedback", {})
                        
                        if isinstance(raw_feedback, dict):
                            # Handle different structures
                            if "value" in raw_feedback and isinstance(raw_feedback["value"], str):
                                # Try to parse JSON string
                                try:
                                    parsed_feedback = json.loads(raw_feedback["value"])
                                    if isinstance(parsed_feedback, dict):
                                        for key, value in parsed_feedback.items():
                                            p = doc.add_paragraph()
                                            p.add_run(f"{key}: ").bold = True
                                            p.add_run(str(value))
                                    else:
                                        doc.add_paragraph(str(raw_feedback["value"]))
                                except (json.JSONDecodeError, TypeError):
                                    doc.add_paragraph(str(raw_feedback["value"]))
                            else:
                                # Regular dict
                                for key, value in raw_feedback.items():
                                    p = doc.add_paragraph()
                                    p.add_run(f"{key}: ").bold = True
                                    p.add_run(str(value))
                        elif isinstance(raw_feedback, str):
                            doc.add_paragraph(raw_feedback)
                    
                    # Add competitor benchmarking score
                    if "competitor_benchmarking_score" in item:
                        p = doc.add_paragraph()
                        p.add_run("Competitor Benchmarking Score: ").bold = True
                        p.add_run(f"{item.get('competitor_benchmarking_score', 0)}/10")
                    
                    # Add brand promise perception score
                    if "brand_promise_perception_score" in item:
                        p = doc.add_paragraph()
                        p.add_run("Brand Promise Perception Score: ").bold = True
                        p.add_run(f"{item.get('brand_promise_perception_score', 0)}/10")
                    
                    # Add simulated market adoption score
                    if "simulated_market_adoption_score" in item:
                        p = doc.add_paragraph()
                        p.add_run("Simulated Market Adoption Score: ").bold = True
                        p.add_run(f"{item.get('simulated_market_adoption_score', 0)}/10")
                    
                    # Add competitive differentiation score
                    if "competitive_differentiation_score" in item:
                        p = doc.add_paragraph()
                        p.add_run("Competitive Differentiation Score: ").bold = True
                        p.add_run(f"{item.get('competitive_differentiation_score', 0)}/10")
                    
                    # Add final recommendation
                    if "final_survey_recommendation" in item:
                        doc.add_heading("Recommendation", level=5)
                        doc.add_paragraph(str(item.get("final_survey_recommendation", "")))
                    
                    # Add separator between personas
                    if i < len(survey_items):
                        doc.add_paragraph("")
                
                # Add a summary section
                doc.add_heading("Survey Summary", level=2)
                
                # Generate aggregate scores
                brand_names = set()
                total_personality_fit = 0
                total_competitor_benchmark = 0
                total_brand_promise = 0
                total_market_adoption = 0
                total_differentiation = 0
                count = 0
                
                for item in survey_items:
                    if not isinstance(item, dict):
                        continue
                        
                    count += 1
                    if "brand_name" in item:
                        brand_names.add(item.get("brand_name", ""))
                    if "personality_fit_score" in item:
                        total_personality_fit += item.get("personality_fit_score", 0)
                    if "competitor_benchmarking_score" in item:
                        total_competitor_benchmark += item.get("competitor_benchmarking_score", 0)
                    if "brand_promise_perception_score" in item:
                        total_brand_promise += item.get("brand_promise_perception_score", 0)
                    if "simulated_market_adoption_score" in item:
                        total_market_adoption += item.get("simulated_market_adoption_score", 0)
                    if "competitive_differentiation_score" in item:
                        total_differentiation += item.get("competitive_differentiation_score", 0)
                
                # Add summary paragraph
                if count > 0:
                    summary = f"The survey simulation evaluated {len(brand_names)} brand names across {count} personas. "
                    
                    # Add average scores
                    summary += "Average scores across all personas:\n"
                    doc.add_paragraph(summary)
                    
                    scores_table = doc.add_table(rows=1, cols=2)
                    scores_table.style = 'Table Grid'
                    
                    # Add header row
                    header_cells = scores_table.rows[0].cells
                    header_cells[0].text = "Metric"
                    header_cells[1].text = "Average Score (0-10)"
                    
                    # Add scores rows
                    metrics = [
                        ("Personality Fit", total_personality_fit / count if count > 0 else 0),
                        ("Competitor Benchmarking", total_competitor_benchmark / count if count > 0 else 0),
                        ("Brand Promise Perception", total_brand_promise / count if count > 0 else 0),
                        ("Market Adoption Potential", total_market_adoption / count if count > 0 else 0),
                        ("Competitive Differentiation", total_differentiation / count if count > 0 else 0)
                    ]
                    
                    for metric, score in metrics:
                        row = scores_table.add_row()
                        row.cells[0].text = metric
                        row.cells[1].text = f"{score:.1f}"
            else:
                # No survey simulation data available
                doc.add_paragraph("No survey simulation data available for this brand naming project.")
                
        except Exception as e:
            logger.error(f"Error formatting survey simulation section: {str(e)}")
            logger.debug(f"Error details: {traceback.format_exc()}")
            doc.add_paragraph(f"Error formatting survey simulation section: {str(e)}")
            # Add a generic error message to the document
            doc.add_paragraph("Unable to format the survey simulation section due to an error in processing the data.")

    async def _format_linguistic_analysis(self, doc: Document, data: Dict[str, Any]) -> None:
        """
        Format the linguistic analysis section using direct ETL process.
        
        Args:
            doc: The document to add content to
            data: The raw linguistic analysis data
        """
        try:
            # Add section title
            doc.add_heading("Linguistic Analysis", level=1)
            
            # Add introduction
            doc.add_paragraph(
                "This section analyzes the linguistic characteristics of the brand name options, "
                "including pronunciation, spelling, phonetic patterns, and other linguistic elements "
                "that influence brand perception and usability."
            )
            
            # Transform data using model
            transformed_data = self._transform_linguistic_analysis(data)
            logger.info(f"Transformed linguistic analysis data: {len(str(transformed_data))} chars")
            
            # Check if we have linguistic analysis data
            if "linguistic_analysis" in transformed_data and transformed_data["linguistic_analysis"]:
                linguistic_analyses = transformed_data["linguistic_analysis"]
                
                # Add a summary of the number of brand names analyzed
                doc.add_paragraph(f"Linguistic analysis was conducted for {len(linguistic_analyses)} brand name candidates.")
                
                # Process each brand name analysis
                for brand_name, analysis in linguistic_analyses.items():
                    # Add brand name as heading
                    doc.add_heading(brand_name, level=2)
                    
                    # Process attributes in a structured way
                    attributes = [
                        ("word_class", "Word Class"),
                        ("sound_symbolism", "Sound Symbolism"),
                        ("rhythm_and_meter", "Rhythm and Meter"),
                        ("pronunciation_ease", "Pronunciation Ease"),
                        ("euphony_vs_cacophony", "Euphony vs Cacophony"),
                        ("inflectional_properties", "Inflectional Properties"),
                        ("neologism_appropriateness", "Neologism Appropriateness"),
                        ("overall_readability_score", "Overall Readability Score"),
                        ("morphological_transparency", "Morphological Transparency"),
                        ("naturalness_in_collocations", "Naturalness in Collocations"),
                        ("ease_of_marketing_integration", "Ease of Marketing Integration"),
                        ("phoneme_frequency_distribution", "Phoneme Frequency Distribution"),
                        ("semantic_distance_from_competitors", "Semantic Distance from Competitors"),
                        ("notes", "Notes")
                    ]
                    
                    for attr_key, attr_display in attributes:
                        if attr_key in analysis and analysis[attr_key]:
                            # Format based on value type
                            value = analysis[attr_key]
                            
                            # Add attribute
                            p = doc.add_paragraph()
                            p.add_run(f"{attr_display}: ").bold = True
                            
                            # Handle different value types
                            if isinstance(value, dict):
                                # Add each dictionary item as a bullet point
                                for sub_key, sub_value in value.items():
                                    bullet_p = doc.add_paragraph(style="List Bullet")
                                    bullet_p.add_run(f"{sub_key.replace('_', ' ').title()}: ").bold = True
                                    bullet_p.add_run(str(sub_value))
                            elif isinstance(value, list):
                                # Add the attribute value as a comma-separated list
                                p.add_run(", ".join(str(item) for item in value))
                            else:
                                p.add_run(str(value))
                    
                    # Add separator between brand analyses (except after the last one)
                    if brand_name != list(linguistic_analyses.keys())[-1]:
                        doc.add_paragraph("", style="Normal")
                
                # Add summary section
                doc.add_heading("Linguistic Analysis Summary", level=2)
                doc.add_paragraph(
                    "The linguistic analysis examines how each brand name functions from a linguistic perspective, "
                    "considering pronunciation, rhythm, sound symbolism, and other factors that influence "
                    "memorability and effectiveness. This analysis helps identify names that are not only "
                    "semantically appropriate but also linguistically strong."
                )
                    
            else:
                # No linguistic analysis data available
                doc.add_paragraph("No linguistic analysis data available for this brand naming project.")
                
        except Exception as e:
            logger.error(f"Error formatting linguistic analysis section: {str(e)}")
            logger.debug(f"Error details: {traceback.format_exc()}")
            # Add a generic error message to the document
            doc.add_paragraph("Unable to format the linguistic analysis section due to an error in processing the data.")

    async def _format_cultural_sensitivity(self, doc: Document, data: Dict[str, Any]) -> None:
        """
        Format the cultural sensitivity analysis section using direct ETL process.
        
        Args:
            doc: The document to add content to
            data: The raw cultural sensitivity analysis data
        """
        try:
            # Add section title
            doc.add_heading("Cultural Sensitivity Analysis", level=1)
            
            # Add introduction
            doc.add_paragraph(
                "This section analyzes the cultural implications of the brand name options across "
                "different regions and cultural contexts, identifying potential sensitivities and risks "
                "that should be considered in the brand naming decision process."
            )
            
            # Transform data using model
            transformed_data = self._transform_cultural_sensitivity(data)
            logger.info(f"Transformed cultural sensitivity data: {len(str(transformed_data))} chars")
            
            # Check if we have analysis data
            if transformed_data and isinstance(transformed_data, dict):
                # Check for brand_analyses key (main data structure)
                if "brand_analyses" in transformed_data and transformed_data["brand_analyses"]:
                    brand_analyses = transformed_data["brand_analyses"]
                    
                    # Add a summary of the number of brand names analyzed
                    doc.add_paragraph(f"Cultural sensitivity analysis was conducted for {len(brand_analyses)} brand name candidates.")
                    
                    # Process each brand name analysis
                    for analysis in brand_analyses:
                        brand_name = analysis.get("brand_name", "Unknown Brand")
                        
                        # Add brand name as heading
                        doc.add_heading(brand_name, level=2)
                        
                        # Process attributes in a structured way - updated to match BrandAnalysis model fields
                        attributes = [
                            ("symbolic_meanings", "Symbolic Meanings"),
                            ("historical_meaning", "Historical Meaning"),
                            ("overall_risk_rating", "Overall Risk Rating"),
                            ("regional_variations", "Regional Variations"),
                            ("cultural_connotations", "Cultural Connotations"),
                            ("current_event_relevance", "Current Event Relevance"),
                            ("religious_sensitivities", "Religious Sensitivities"), 
                            ("social_political_taboos", "Social Political Taboos"),
                            ("age_related_connotations", "Age-Related Connotations"),
                            ("alignment_with_cultural_values", "Alignment with Cultural Values"),
                            ("notes", "Notes")
                        ]
                        
                        for attr_key, attr_display in attributes:
                            if attr_key in analysis and analysis[attr_key]:
                                # Format based on value type
                                value = analysis[attr_key]
                                
                                # Add attribute
                                p = doc.add_paragraph()
                                p.add_run(f"{attr_display}: ").bold = True
                                
                                # Handle dictionary attributes
                                if isinstance(value, dict):
                                    # Add each dictionary item as a bullet point
                                    for sub_key, sub_value in value.items():
                                        bullet_p = doc.add_paragraph(style="List Bullet")
                                        bullet_p.add_run(f"{sub_key.replace('_', ' ').title()}: ").bold = True
                                        bullet_p.add_run(str(sub_value))
                                # Handle list attributes
                                elif isinstance(value, list):
                                    # Add the attribute value as a bullet list
                                    for item in value:
                                        bullet_p = doc.add_paragraph(style="List Bullet")
                                        bullet_p.add_run(str(item))
                                # Handle simple string/number attributes
                                else:
                                    p.add_run(str(value))
                        
                        # Add separator between brand analyses (except after the last one)
                        if analysis != brand_analyses[-1]:
                            doc.add_paragraph("")
                    
                    # Add summary if available
                    if "summary" in transformed_data and transformed_data["summary"]:
                        doc.add_heading("Cultural Sensitivity Summary", level=2)
                        doc.add_paragraph(transformed_data["summary"])
                else:
                    # No brand analyses available
                    doc.add_paragraph("No cultural sensitivity analysis data is available for the brand names.")
            else:
                # No cultural sensitivity data available
                doc.add_paragraph("No cultural sensitivity analysis data available for this brand naming project.")
                
        except Exception as e:
            logger.error(f"Error formatting cultural sensitivity section: {str(e)}")
            logger.debug(f"Error details: {traceback.format_exc()}")
            # Add a generic error message to the document
            doc.add_paragraph("Unable to format the cultural sensitivity section due to an error in processing the data.")

    async def _format_name_evaluation(self, doc: Document, data: Dict[str, Any]) -> None:
        """
        Format the brand name evaluation section using direct ETL process.
        
        Args:
            doc: The document to add content to
            data: The raw brand name evaluation data
        """
        try:
            # Add section title
            doc.add_heading("Brand Name Evaluation", level=1)
            
            # Add introduction
            doc.add_paragraph(
                "This section provides a comprehensive evaluation of brand name candidates "
                "based on multiple criteria, highlighting their strengths and weaknesses "
                "to support the final recommendation process."
            )
            
            # Transform data using model
            transformed_data = self._transform_name_evaluation(data)
            logger.info(f"Transformed name evaluation data: {len(str(transformed_data))} chars")
            
            # Check if we have evaluation data
            if transformed_data and isinstance(transformed_data, dict):
                # Process evaluation methodology
                if "evaluation_methodology" in transformed_data and transformed_data["evaluation_methodology"]:
                    doc.add_heading("Evaluation Methodology", level=2)
                    doc.add_paragraph(transformed_data["evaluation_methodology"])
                
                # Process shortlisted names
                if "shortlisted_names" in transformed_data and transformed_data["shortlisted_names"]:
                    doc.add_heading("Shortlisted Brand Names", level=2)
                    doc.add_paragraph(
                        f"The following {len(transformed_data['shortlisted_names'])} brand names were shortlisted "
                        f"based on their strong performance across evaluation criteria."
                    )
                    
                    # Process each shortlisted name
                    for name_data in transformed_data["shortlisted_names"]:
                        doc.add_heading(name_data["brand_name"], level=3)
                        
                        # Add overall score
                        p = doc.add_paragraph()
                        p.add_run("Overall Score: ").bold = True
                        p.add_run(f"{name_data['overall_score']}/10")
                        
                        # Add evaluation comments
                        if "evaluation_comments" in name_data and name_data["evaluation_comments"]:
                            doc.add_paragraph(name_data["evaluation_comments"])
                        
                        # Add separator between brand evaluations (except for the last one)
                        if name_data != transformed_data["shortlisted_names"][-1]:
                            doc.add_paragraph("")
                
                # Process other names
                if "other_names" in transformed_data and transformed_data["other_names"]:
                    doc.add_heading("Other Evaluated Brand Names", level=2)
                    doc.add_paragraph(
                        f"The following {len(transformed_data['other_names'])} brand names were evaluated "
                        f"but not shortlisted."
                    )
                    
                    # Process each non-shortlisted name
                    for name_data in transformed_data["other_names"]:
                        doc.add_heading(name_data["brand_name"], level=3)
                        
                        # Add overall score
                        p = doc.add_paragraph()
                        p.add_run("Overall Score: ").bold = True
                        p.add_run(f"{name_data['overall_score']}/10")
                        
                        # Add evaluation comments
                        if "evaluation_comments" in name_data and name_data["evaluation_comments"]:
                            doc.add_paragraph(name_data["evaluation_comments"])
                        
                        # Add separator between brand evaluations (except for the last one)
                        if name_data != transformed_data["other_names"][-1]:
                            doc.add_paragraph("")
                
                # Process comparative analysis
                if "comparative_analysis" in transformed_data and transformed_data["comparative_analysis"]:
                    doc.add_heading("Comparative Analysis", level=2)
                    doc.add_paragraph(transformed_data["comparative_analysis"])
                
                # Process final rankings
                if "final_rankings" in transformed_data and transformed_data["final_rankings"]:
                    doc.add_heading("Final Rankings", level=2)
                    
                    rankings = transformed_data["final_rankings"]
                    if isinstance(rankings, dict):
                        # Create a table for rankings with scores
                        table = doc.add_table(rows=len(rankings)+1, cols=3)
                        table.style = 'Table Grid'
                        
                        # Add header row
                        header_cells = table.rows[0].cells
                        header_cells[0].text = "Rank"
                        header_cells[1].text = "Brand Name"
                        header_cells[2].text = "Score"
                        
                        # Sort rankings by score (descending)
                        sorted_rankings = sorted(rankings.items(), key=lambda x: x[1], reverse=True)
                        
                        # Add ranking rows
                        for i, (brand_name, score) in enumerate(sorted_rankings, 1):
                            cells = table.rows[i].cells
                            cells[0].text = str(i)
                            cells[1].text = str(brand_name)
                            cells[2].text = str(score)
                    else:
                        # Simple string or list format
                        doc.add_paragraph(str(rankings))
                
                # Add summary section
                doc.add_heading("Evaluation Summary", level=2)
                doc.add_paragraph(
                    "The brand name evaluation process identified names that best align with "
                    "the brand's strategic objectives, target audience, and positioning. "
                    "The shortlisted names demonstrated stronger performance across key criteria "
                    "including distinctiveness, memorability, and strategic alignment."
                )
            else:
                # No evaluation data available
                doc.add_paragraph("No brand name evaluation data available for this brand naming project.")
                
        except Exception as e:
            logger.error(f"Error formatting brand name evaluation section: {str(e)}")
            logger.debug(f"Error details: {traceback.format_exc()}")
            # Add a generic error message to the document
            doc.add_paragraph("Unable to format the brand name evaluation section due to an error in processing the data.")

    async def _format_seo_analysis(self, doc: Document, data: Dict[str, Any]) -> None:
        """
        Format the SEO analysis section using direct ETL process.
        
        Args:
            doc: The document to add content to
            data: The raw SEO analysis data
        """
        try:
            # Add section title
            doc.add_heading("SEO Analysis", level=1)
            
            # Add introduction
            doc.add_paragraph(
                "This section analyzes the SEO potential and online discoverability of each brand name, "
                "including search metrics, keyword opportunities, and recommendations for improving "
                "online visibility."
            )
            
            # Transform data using model
            transformed_data = self._transform_seo_analysis(data)
            logger.debug(f"Transformed SEO analysis data: {len(str(transformed_data))} chars")
            
            # Check if we have analysis data
            if transformed_data and "seo_analysis" in transformed_data:
                # Process each brand's SEO analysis
                for brand_name, details in transformed_data["seo_analysis"].items():
                    # Add brand name as heading
                    doc.add_heading(f"SEO Analysis for {brand_name}", level=2)
                    
                    # Search Metrics
                    doc.add_heading("Search Metrics", level=3)
                    metrics_table = doc.add_table(rows=4, cols=2)
                    metrics_table.style = 'Table Grid'
                    
                    # Add search metrics details
                    rows = [
                        ("Search Volume", str(details["search_volume"])),
                        ("Exact Match Results", details["exact_match_search_results"]),
                        ("SEO Viability Score", f"{details['seo_viability_score']}/10"),
                        ("Keyword Competition", details["keyword_competition"])
                    ]
                    
                    for i, (label, value) in enumerate(rows):
                        cells = metrics_table.rows[i].cells
                        cells[0].text = label
                        cells[1].text = str(value)
                    
                    doc.add_paragraph()  # Add spacing
                    
                    # Keyword Analysis
                    doc.add_heading("Keyword Analysis", level=3)
                    keyword_table = doc.add_table(rows=5, cols=2)
                    keyword_table.style = 'Table Grid'
                    
                    # Add keyword analysis details
                    rows = [
                        ("Keyword Alignment", details["keyword_alignment"]),
                        ("Branded Keyword Potential", details["branded_keyword_potential"]),
                        ("Non-Branded Keyword Potential", details["non_branded_keyword_potential"]),
                        ("Negative Keyword Associations", details["negative_keyword_associations"]),
                        ("Name Length Impact", details["name_length_searchability"])
                    ]
                    
                    for i, (label, value) in enumerate(rows):
                        cells = keyword_table.rows[i].cells
                        cells[0].text = label
                        cells[1].text = str(value)
                    
                    doc.add_paragraph()  # Add spacing
                    
                    # Online Presence
                    doc.add_heading("Online Presence", level=3)
                    presence_table = doc.add_table(rows=4, cols=2)
                    presence_table.style = 'Table Grid'
                    
                    # Add online presence details
                    rows = [
                        ("Social Media Availability", "Available" if details["social_media_availability"] else "Not Available"),
                        ("Social Media Discoverability", details["social_media_discoverability"]),
                        ("Competitor Domain Strength", details["competitor_domain_strength"]),
                        ("Negative Search Results", "Yes" if details["negative_search_results"] else "No")
                    ]
                    
                    for i, (label, value) in enumerate(rows):
                        cells = presence_table.rows[i].cells
                        cells[0].text = label
                        cells[1].text = str(value)
                    
                    doc.add_paragraph()  # Add spacing
                    
                    # Content Marketing Opportunities
                    doc.add_heading("Content Marketing Opportunities", level=3)
                    doc.add_paragraph(details["content_marketing_opportunities"])
                    
                    # SEO Recommendations
                    doc.add_heading("SEO Recommendations", level=3)
                    recommendations = details["seo_recommendations"]
                    if isinstance(recommendations, str):
                        try:
                            # Try to evaluate the string as a list if it starts with [ and ends with ]
                            if recommendations.startswith("[") and recommendations.endswith("]"):
                                recommendations = eval(recommendations)
                        except:
                            # If evaluation fails, split by comma
                            recommendations = [r.strip() for r in recommendations.split(",")]
                    
                    for recommendation in recommendations:
                        # Remove any remaining quotes
                        recommendation = str(recommendation).strip("'\"")
                        doc.add_paragraph(recommendation, style='List Bullet')
                    
                    doc.add_paragraph()  # Add spacing between brand analyses
            else:
                doc.add_paragraph("No SEO analysis data available.")
                
        except Exception as e:
            logger.error(f"Error formatting SEO analysis: {str(e)}")
            doc.add_paragraph("Error formatting SEO analysis section.")

    async def _format_brand_context(self, doc: Document, data: Dict[str, Any]) -> None:
        """
        Format the brand context section using direct ETL process.
        
        Args:
            doc: The document to add content to
            data: The raw brand context data
        """
        try:
            # Add section title
            doc.add_heading("Brand Context", level=1)
            
            # Add introduction
            doc.add_paragraph(
                "This section provides an overview of the brand's foundation, including its values, mission, "
                "target audience, and market positioning. It serves as the strategic framework for the "
                "brand naming process and ensures alignment with the organization's goals."
            )
            
            # Transform data using model
            transformed_data = self._transform_brand_context(data)
            logger.debug(f"Transformed brand context data: {len(str(transformed_data))} chars")
            
            # Check if we have brand context data
            if transformed_data and isinstance(transformed_data, dict):
                # Brand Values
                if "brand_values" in transformed_data and transformed_data["brand_values"]:
                    doc.add_heading("Brand Values", level=2)
                    values = transformed_data["brand_values"]
                    if isinstance(values, list):
                        for value in values:
                            doc.add_paragraph(f"• {value}", style="List Bullet")
                    else:
                        doc.add_paragraph(str(values))
                
                # Brand Mission
                if "brand_mission" in transformed_data and transformed_data["brand_mission"]:
                    doc.add_heading("Brand Mission", level=2)
                    doc.add_paragraph(transformed_data["brand_mission"])
                
                # Brand Promise
                if "brand_promise" in transformed_data and transformed_data["brand_promise"]:
                    doc.add_heading("Brand Promise", level=2)
                    doc.add_paragraph(transformed_data["brand_promise"])
                
                # Brand Purpose
                if "brand_purpose" in transformed_data and transformed_data["brand_purpose"]:
                    doc.add_heading("Brand Purpose", level=2)
                    doc.add_paragraph(transformed_data["brand_purpose"])
                
                # Customer Needs
                if "customer_needs" in transformed_data and transformed_data["customer_needs"]:
                    doc.add_heading("Customer Needs", level=2)
                    needs = transformed_data["customer_needs"]
                    if isinstance(needs, list):
                        for need in needs:
                            doc.add_paragraph(f"• {need}", style="List Bullet")
                    else:
                        doc.add_paragraph(str(needs))
                
                # Industry Focus
                if "industry_focus" in transformed_data and transformed_data["industry_focus"]:
                    doc.add_heading("Industry Focus", level=2)
                    doc.add_paragraph(transformed_data["industry_focus"])
                
                # Industry Trends
                if "industry_trends" in transformed_data and transformed_data["industry_trends"]:
                    doc.add_heading("Industry Trends", level=2)
                    trends = transformed_data["industry_trends"]
                    if isinstance(trends, list):
                        for trend in trends:
                            doc.add_paragraph(f"• {trend}", style="List Bullet")
                    else:
                        doc.add_paragraph(str(trends))
                
                # Target Audience
                if "target_audience" in transformed_data and transformed_data["target_audience"]:
                    doc.add_heading("Target Audience", level=2)
                    doc.add_paragraph(transformed_data["target_audience"])
                
                # Brand Personality
                if "brand_personality" in transformed_data and transformed_data["brand_personality"]:
                    doc.add_heading("Brand Personality", level=2)
                    personality = transformed_data["brand_personality"]
                    if isinstance(personality, list):
                        for trait in personality:
                            doc.add_paragraph(f"• {trait}", style="List Bullet")
                    else:
                        doc.add_paragraph(str(personality))
                
                # Market Positioning
                if "market_positioning" in transformed_data and transformed_data["market_positioning"]:
                    doc.add_heading("Market Positioning", level=2)
                    doc.add_paragraph(transformed_data["market_positioning"])
                
                # Brand Tone of Voice
                if "brand_tone_of_voice" in transformed_data and transformed_data["brand_tone_of_voice"]:
                    doc.add_heading("Brand Tone of Voice", level=2)
                    doc.add_paragraph(transformed_data["brand_tone_of_voice"])
                
                # Brand Identity Brief
                if "brand_identity_brief" in transformed_data and transformed_data["brand_identity_brief"]:
                    doc.add_heading("Brand Identity Brief", level=2)
                    doc.add_paragraph(transformed_data["brand_identity_brief"])
                
                # Competitive Landscape
                if "competitive_landscape" in transformed_data and transformed_data["competitive_landscape"]:
                    doc.add_heading("Competitive Landscape", level=2)
                    doc.add_paragraph(transformed_data["competitive_landscape"])
                
                # Conclusion
                doc.add_heading("Summary", level=2)
                doc.add_paragraph(
                    f"This brand context provides the foundation for the naming process. "
                    f"The ideal brand name should align with the brand's values of "
                    f"{', '.join(transformed_data.get('brand_values', [''])[:3])} and "
                    f"appeal to the target audience of {transformed_data.get('target_audience', '').split(',')[0]}. "
                    f"It should reflect the brand's {', '.join(transformed_data.get('brand_personality', [''])[:2])} "
                    f"personality and support its positioning as {transformed_data.get('market_positioning', '').split('.')[0]}."
                )
            else:
                # No brand context data available
                doc.add_paragraph("No brand context data available for this brand naming project.")
        
        except Exception as e:
            logger.error(f"Error formatting brand context section: {str(e)}")
            logger.debug(f"Error details: {traceback.format_exc()}")
            # Add a generic error message to the document
            doc.add_paragraph("Unable to format the brand context section due to an error in processing the data.")

    async def _format_name_generation(self, doc: Document, data: Dict[str, Any]) -> None:
        """
        Format the brand name generation section using direct ETL process.
        
        Args:
            doc: The document to add content to
            data: The raw name generation data
        """
        try:
            # Add section title
            doc.add_heading("Brand Name Generation", level=1)
            
            # Transform data using model
            transformed_data = self._transform_name_generation(data)
            logger.info(f"Transformed name generation data: {len(str(transformed_data))} chars")
            
            # Add introduction if available
            if "introduction" in transformed_data and transformed_data["introduction"]:
                doc.add_paragraph(transformed_data["introduction"])
            else:
                # Add default introduction
                doc.add_paragraph("The following section presents brand name candidates generated based on the brand context and naming strategy requirements.")
            
            # Add methodology and approach
            if "methodology_and_approach" in transformed_data and transformed_data["methodology_and_approach"]:
                doc.add_heading("Methodology and Approach", level=2)
                doc.add_paragraph(transformed_data["methodology_and_approach"])
            
            # Process categories and names
            if "categories" in transformed_data and isinstance(transformed_data["categories"], list):
                doc.add_heading("Brand Name Categories", level=2)
                
                # Add each category
                for category in transformed_data["categories"]:
                    # Add category heading
                    category_name = category.get("category_name")
                    if category_name:
                        doc.add_heading(category_name, level=3)
                    
                    # Add category description
                    category_description = category.get("category_description")
                    if category_description:
                        doc.add_paragraph(category_description)
                    
                    # Process names in this category
                    names = category.get("names", [])
                    if names:
                        # Add a count of names in this category
                        doc.add_paragraph(f"This category includes {len(names)} brand name options:")
                        
                        # Process each name in this category
                        for name in names:
                            # Add name heading
                            brand_name = name.get("brand_name")
                            if brand_name:
                                doc.add_heading(brand_name, level=4)
                            
                            # Add name attributes in a structured format
                            attributes = [
                                ("brand_personality_alignment", "Brand Personality Alignment"),
                                ("brand_promise_alignment", "Brand Promise Alignment"),
                                ("name_generation_methodology", "Methodology"),
                                ("memorability_score_details", "Memorability"),
                                ("pronounceability_score_details", "Pronounceability"),
                                ("visual_branding_potential_details", "Visual Branding Potential"),
                                ("target_audience_relevance_details", "Target Audience Relevance"),
                                ("market_differentiation_details", "Market Differentiation"),
                                ("rationale", "Rationale"),
                                ("trademark_status", "Trademark Status"),
                                ("cultural_considerations", "Cultural Considerations")
                            ]
                            
                            for attr_key, attr_display in attributes:
                                if attr_key in name and name[attr_key]:
                                    p = doc.add_paragraph()
                                    p.add_run(f"{attr_display}: ").bold = True
                                    p.add_run(str(name[attr_key]))
                            
                            # Add a separator between names (except for the last one)
                            if name != names[-1]:
                                doc.add_paragraph("", style="Normal")
            
            # Add generated names overview if available
            if "generated_names_overview" in transformed_data and transformed_data["generated_names_overview"]:
                doc.add_heading("Generated Names Overview", level=2)
                
                if isinstance(transformed_data["generated_names_overview"], dict):
                    # Get total count
                    total_count = transformed_data["generated_names_overview"].get("total_count", 0)
                    
                    # If we have a total count, display it
                    if total_count:
                        doc.add_paragraph(f"A total of {total_count} names were generated across various naming categories.")
                    
                    # Add any other overview information
                    for key, value in transformed_data["generated_names_overview"].items():
                        if key != "total_count" and value:
                            p = doc.add_paragraph()
                            p.add_run(f"{key.replace('_', ' ').title()}: ").bold = True
                            p.add_run(str(value))
            
            # Add evaluation metrics if available
            if "evaluation_metrics" in transformed_data and transformed_data["evaluation_metrics"]:
                doc.add_heading("Evaluation Metrics", level=2)
                
                if isinstance(transformed_data["evaluation_metrics"], dict):
                    for key, value in transformed_data["evaluation_metrics"].items():
                        p = doc.add_paragraph()
                        p.add_run(f"{key.replace('_', ' ').title()}: ").bold = True
                        p.add_run(str(value))
            
            # Add summary if available
            if "summary" in transformed_data and transformed_data["summary"]:
                doc.add_heading("Summary", level=2)
                doc.add_paragraph(transformed_data["summary"])
                
        except Exception as e:
            logger.error(f"Error formatting name generation section: {str(e)}")
            logger.debug(f"Error details: {traceback.format_exc()}")
            # Add a generic error message to the document
            doc.add_paragraph("Unable to format the brand name generation section due to an error in processing the data.")

    async def _format_competitor_analysis(self, doc: Document, data: Dict[str, Any]) -> None:
        """
        Format the competitor analysis section using direct ETL process.
        
        Args:
            doc: The document to add content to
            data: The raw competitor analysis data
        """
        try:
            # Add section title
            doc.add_heading("Competitor Analysis", level=1)
            
            # Add introduction
            doc.add_paragraph(
                "This section analyzes competitors' brand naming strategies and industry positioning, "
                "providing valuable context for the evaluation of proposed brand names. "
                "Understanding competitor naming patterns helps inform differentiation strategies "
                "and identify potential risks of market confusion."
            )
            
            # Transform data using model
            transformed_data = self._transform_competitor_analysis(data)
            logger.info(f"Transformed competitor analysis data: {len(str(transformed_data))} chars")
            
            # Check if we have analysis data
            if transformed_data and isinstance(transformed_data, dict):
                # Process each brand name analysis
                for brand_name, analysis in transformed_data.items():
                    # Add brand name as heading
                    doc.add_heading(f"Competitor Analysis for {brand_name}", level=2)
                    
                    # Create comparison table for all competitors
                    if analysis.get("competitors"):
                        doc.add_heading("Competitor Overview", level=3)
                        comparison_table = doc.add_table(rows=1 + len(analysis["competitors"]), cols=4)
                        comparison_table.style = 'Table Grid'
                        
                        # Add header row
                        header_cells = comparison_table.rows[0].cells
                        header_cells[0].text = "Competitor"
                        header_cells[1].text = "Naming Style"
                        header_cells[2].text = "Risk of Confusion"
                        header_cells[3].text = "Differentiation Score"
                        
                        # Add competitor rows
                        for i, competitor in enumerate(analysis["competitors"], 1):
                            cells = comparison_table.rows[i].cells
                            cells[0].text = competitor["competitor_name"]
                            cells[1].text = competitor["competitor_naming_style"]
                            cells[2].text = competitor["risk_of_confusion"]
                            cells[3].text = f"{competitor['differentiation_score']}/10"
                        
                        doc.add_paragraph()  # Add spacing
                        
                        # Add detailed analysis for each competitor
                        doc.add_heading("Detailed Competitor Analysis", level=3)
                        for competitor in analysis["competitors"]:
                            # Add competitor name as subheading
                            doc.add_heading(competitor["competitor_name"], level=4)
                            
                            # Create detailed table
                            detail_table = doc.add_table(rows=10, cols=2)
                            detail_table.style = 'Table Grid'
                            
                            # Add all competitor details
                            rows = [
                                ("Keywords", competitor["competitor_keywords"]),
                                ("Positioning", competitor["competitor_positioning"]),
                                ("Strengths", competitor["competitor_strengths"]),
                                ("Weaknesses", competitor["competitor_weaknesses"]),
                                ("Target Audience Perception", competitor["target_audience_perception"]),
                                ("Differentiation Opportunity", competitor["competitor_differentiation_opportunity"]),
                                ("Competitive Advantage", competitor["competitive_advantage_notes"]),
                                ("Trademark Risk", competitor["trademark_conflict_risk"]),
                                ("Risk of Confusion", competitor["risk_of_confusion"]),
                                ("Differentiation Score", f"{competitor['differentiation_score']}/10")
                            ]
                            
                            for i, (label, value) in enumerate(rows):
                                cells = detail_table.rows[i].cells
                                cells[0].text = label
                                cells[1].text = str(value)
                            
                            doc.add_paragraph()  # Add spacing
                    
                    # Add separator between brand analyses (except for the last one)
                    if brand_name != list(transformed_data.keys())[-1]:
                        doc.add_paragraph("")
            else:
                doc.add_paragraph("No competitor analysis data available.")
                
        except Exception as e:
            logger.error(f"Error formatting competitor analysis: {str(e)}")
            doc.add_paragraph("Error formatting competitor analysis section.")

    async def _add_table_of_contents(self, doc: Document) -> None:
        """Add a table of contents to the document."""
        logger.info("Adding table of contents")
        
        try:
            # Add the TOC heading
            doc.add_heading("Table of Contents", level=1)
            
            # Define the hardcoded TOC sections with descriptions
            toc_sections = [
                {
                    "title": "1. Executive Summary",
                    "description": "Project overview, key findings, and top recommendations."
                },
                {
                    "title": "2. Brand Context",
                    "description": "Brand identity, values, target audience, market positioning, and industry context."
                },
                {
                    "title": "3. Name Generation",
                    "description": "Methodology, generated names overview, and initial evaluation metrics."
                },
                {
                    "title": "4. Linguistic Analysis",
                    "description": "Pronunciation, euphony, rhythm, sound symbolism, and overall readability."
                },
                {
                    "title": "5. Semantic Analysis",
                    "description": "Meaning, etymology, emotional valence, and brand personality fit."
                },
                {
                    "title": "6. Cultural Sensitivity",
                    "description": "Cultural connotations, symbolic meanings, sensitivities, and risk assessment."
                },
                {
                    "title": "7. Translation Analysis",
                    "description": "Translations, semantic shifts, pronunciation, and global consistency."
                },
                {
                    "title": "8. Name Evaluation",
                    "description": "Strategic alignment, distinctiveness, memorability, and shortlisted names."
                },
                {
                    "title": "9. Domain Analysis",
                    "description": "Domain availability, alternative TLDs, social media availability, and SEO potential."
                },
                {
                    "title": "10. SEO/Online Discoverability",
                    "description": "Keyword alignment, search volume potential, and content opportunities."
                },
                {
                    "title": "11. Competitor Analysis",
                    "description": "Competitor naming styles, market positioning, and differentiation opportunities."
                },
                {
                    "title": "12. Market Research",
                    "description": "Market opportunity, target audience fit, viability, and industry insights."
                },
                {
                    "title": "13. Survey Simulation",
                    "description": "Persona demographics, brand perception, emotional associations, and feedback."
                },
                {
                    "title": "14. Strategic Recommendations",
                    "description": "Final name recommendations, implementation strategy, and next steps."
                }
            ]
            
            # Add each section to the document
            for section in toc_sections:
                # Add section title and description
                p = doc.add_paragraph()
                p.add_run(section["title"]).bold = True
                p.add_run(": " + section["description"])
                p.add_run().add_break()  # Add line break after each section
            
            # Add a page break after section descriptions
            doc.add_page_break()
            
            # Add the actual table of contents field (Word will populate this)
            doc.add_heading("Document Outline", level=1)
            paragraph = doc.add_paragraph()
            run = paragraph.add_run()
            fld_char = OxmlElement('w:fldChar')
            fld_char.set(qn('w:fldCharType'), 'begin')
            
            instr_text = OxmlElement('w:instrText')
            instr_text.set(qn('xml:space'), 'preserve')
            instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'
            
            fld_char_end = OxmlElement('w:fldChar')
            fld_char_end.set(qn('w:fldCharType'), 'end')
            
            r_element = run._r
            r_element.append(fld_char)
            r_element.append(instr_text)
            r_element.append(fld_char_end)
            
            # Add a page break after TOC
            doc.add_page_break()
                
        except Exception as e:
            logger.error(f"Error adding table of contents: {str(e)}")
            doc.add_heading("Table of Contents", level=1)
            doc.add_paragraph("An error occurred while generating the table of contents.")
            
            # Still add the TOC field so document navigation works
            paragraph = doc.add_paragraph()
            run = paragraph.add_run()
            fld_char = OxmlElement('w:fldChar')
            fld_char.set(qn('w:fldCharType'), 'begin')
            
            instr_text = OxmlElement('w:instrText')
            instr_text.set(qn('xml:space'), 'preserve')
            instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'
            
            fld_char_end = OxmlElement('w:fldChar')
            fld_char_end.set(qn('w:fldCharType'), 'end')
            
            r_element = run._r
            r_element.append(fld_char)
            r_element.append(instr_text)
            r_element.append(fld_char_end)
            
            doc.add_page_break()
    
    async def _add_executive_summary(self, doc: Document, data: Dict[str, Any]) -> None:
        """Add an executive summary to the document."""
        logger.info("Adding executive summary")
        
        try:
            # Extract necessary data
            all_data = await self.fetch_raw_data(self.current_run_id)
            brand_context = all_data.get("brand_context", {})
            
            # Fetch the user prompt from the state
            user_prompt = await self.fetch_user_prompt(self.current_run_id)
            
            # Count total names generated
            total_names = 0
            shortlisted_names = []
            
            # Get name generation data if available
            if "brand_name_generation" in all_data:
                name_gen_data = all_data["brand_name_generation"]
                
                # Count names across categories
                if isinstance(name_gen_data, dict) and "categories" in name_gen_data and isinstance(name_gen_data["categories"], list):
                    for category in name_gen_data["categories"]:
                        if "names" in category and isinstance(category["names"], list):
                            total_names += len(category["names"])
                # Alternative format: name_categories
                elif isinstance(name_gen_data, dict) and "name_categories" in name_gen_data and isinstance(name_gen_data["name_categories"], list):
                    for category in name_gen_data["name_categories"]:
                        if "names" in category and isinstance(category["names"], list):
                            total_names += len(category["names"])
            
            # Get shortlisted names if available
            if "brand_name_evaluation" in all_data:
                eval_data = all_data["brand_name_evaluation"]
                
                # Extract shortlisted names - handle different structures
                if isinstance(eval_data, dict) and "brand_name_evaluation" in eval_data and isinstance(eval_data["brand_name_evaluation"], dict):
                    # Format where brand_name_evaluation is nested
                    if "shortlisted_names" in eval_data["brand_name_evaluation"] and isinstance(eval_data["brand_name_evaluation"]["shortlisted_names"], list):
                        for item in eval_data["brand_name_evaluation"]["shortlisted_names"]:
                            if isinstance(item, dict) and "brand_name" in item:
                                shortlisted_names.append(item["brand_name"])
                elif isinstance(eval_data, dict) and "shortlisted_names" in eval_data and isinstance(eval_data["shortlisted_names"], list):
                    # Direct format with shortlisted_names at root
                    for item in eval_data["shortlisted_names"]:
                        if isinstance(item, dict) and "brand_name" in item:
                            shortlisted_names.append(item["brand_name"])
                else:
                    # Iterate through items to find shortlisted ones
                    for name, details in eval_data.items():
                        if isinstance(details, dict) and details.get("shortlist_status", False):
                            shortlisted_names.append(name)
            
            # Prepare data for all sections to include in executive summary
            sections_data = {}
            for section_name in self.SECTION_ORDER:
                if section_name in all_data and section_name not in ["exec_summary", "final_recommendations"]:
                    sections_data[section_name] = all_data[section_name]
            
            # Create format data for the executive summary template
            format_data = {
                "run_id": self.current_run_id,
                "sections_data": json.dumps(sections_data, indent=2),
                "brand_context": json.dumps(brand_context, indent=2),
                "total_names": total_names,
                "shortlisted_names": shortlisted_names,
                "user_prompt": user_prompt
            }
            
            # Make sure we have the executive_summary template loaded
            if "executive_summary" not in self.prompts:
                template_path = os.path.join(PROMPTS_DIR, "executive_summary.yaml")
                self.prompts["executive_summary"] = _safe_load_prompt(template_path)
                logger.info(f"Loaded executive summary template from {template_path}")
            
            # Format the template
            formatted_prompt = self._format_template("executive_summary", format_data, "executive_summary")
            
            # Get system content
            system_content = self._get_system_content("You are an expert report formatter creating a professional executive summary.")
            
            # Call LLM to generate executive summary
            logger.info(f"Calling LLM to generate executive summary with prompt length: {len(formatted_prompt)}")
            response = await self._safe_llm_invoke([
                SystemMessage(content=system_content),
                HumanMessage(content=formatted_prompt)
            ], section_name="Executive Summary")
            
            if response:
                logger.info(f"Received LLM response for executive summary with length: {len(response.content)}")
            else:
                logger.error("No response received from LLM for executive summary")
            
            # Try to parse the response
            summary_data = {}
            try:
                content_str = response.content
                # Extract JSON if in code blocks
                json_match = re.search(r'```(?:json)?\s*([\s\S]*?)\s*```', content_str)
                if json_match:
                    content_str = json_match.group(1)
                summary_data = json.loads(content_str)
                logger.info(f"Successfully parsed JSON from executive summary response: {list(summary_data.keys())}")
            except Exception as e:
                logger.error(f"Error parsing executive summary response: {str(e)}")
                # If JSON parsing fails, use the raw content
                summary_data = {"summary": response.content}
            
            # Add the executive summary heading
            doc.add_heading("Executive Summary", level=1)
            
            # Add the summary content
            if "summary" in summary_data and summary_data["summary"]:
                doc.add_paragraph(summary_data["summary"])
            
            # Add key points if available
            if "key_points" in summary_data and isinstance(summary_data["key_points"], list):
                doc.add_heading("Key Points", level=2)
                for point in summary_data["key_points"]:
                    bullet = doc.add_paragraph(style='List Bullet')
                    bullet.add_run(point)
            
            # Add page break after executive summary
            doc.add_page_break()
            
        except Exception as e:
            logger.error(f"Error adding executive summary: {str(e)}")
            logger.error(traceback.format_exc())
            
            # Add a simple executive summary as fallback
            doc.add_heading("Executive Summary", level=1)
            doc.add_paragraph(
                "This report presents a comprehensive brand naming analysis based on linguistic, semantic, "
                "cultural, and market research. It provides detailed evaluations of potential brand names "
                "and offers strategic recommendations for the final name selection."
            )
            doc.add_page_break()

    async def _add_recommendations(self, doc: Document, data: Dict[str, Any]) -> None:
        """Add strategic recommendations to the document."""
        logger.info("Adding strategic recommendations")
        
        try:
            # Extract brand context
            brand_context = data.get("brand_context", {})
            brand_name = brand_context.get("brand_name", "")
            industry = brand_context.get("industry", "")
            
            # Create format data
            format_data = {
                "run_id": self.current_run_id,
                "brand_name": brand_name,
                "industry": industry,
                "brand_context": json.dumps(brand_context, indent=2),
                "data": json.dumps(data, indent=2)
            }
            
            # Format the template using the helper method
            formatted_prompt = self._format_template("recommendations", format_data, "Strategic Recommendations")
            
            # Get system content
            system_content = self._get_system_content("You are an expert report formatter creating professional strategic recommendations.")
            
            # Call LLM to generate recommendations
            response = await self._safe_llm_invoke([
                SystemMessage(content=system_content),
                HumanMessage(content=formatted_prompt)
            ], section_name="Strategic Recommendations")
            
            # Add the recommendations to the document
            doc.add_heading("Strategic Recommendations", level=1)
            
            # Parse the response
            try:
                content = json.loads(response.content)
                
                # Add introduction if available
                if "introduction" in content and content["introduction"]:
                    doc.add_paragraph(content["introduction"])
                
                # Add recommendations if available
                if "recommendations" in content and isinstance(content["recommendations"], list):
                    for i, rec in enumerate(content["recommendations"]):
                        if isinstance(rec, dict) and "title" in rec and "details" in rec:
                            # Rich recommendation format
                            doc.add_heading(f"{i+1}. {rec['title']}", level=2)
                            doc.add_paragraph(rec["details"])
                            
                            # Add sub-recommendations if available
                            if "steps" in rec and isinstance(rec["steps"], list):
                                for step in rec["steps"]:
                                    bullet = doc.add_paragraph(style='List Bullet')
                                    bullet.add_run(step)
                        else:
                            # Simple recommendation format
                            doc.add_heading(f"Recommendation {i+1}", level=2)
                            doc.add_paragraph(str(rec))
                
                # Add conclusion if available
                if "conclusion" in content and content["conclusion"]:
                    doc.add_heading("Conclusion", level=2)
                    doc.add_paragraph(content["conclusion"])
                    
            except json.JSONDecodeError:
                # Fallback to using raw text
                doc.add_paragraph(response.content)
                
        except Exception as e:
            logger.error(f"Error adding recommendations: {str(e)}")
            doc.add_paragraph(f"Error occurred while adding recommendations: {str(e)}", style='Intense Quote')

    async def _format_generic_section(self, doc: Document, section_name: str, data: Dict[str, Any]) -> None:
        """
        Format a generic section of the report.
        
        Args:
            doc: The document to add content to
            section_name: The name of the section to format
            data: The data to format
        """
        try:
            # Check if we have a specific format method for this section
            if section_name in self.REVERSE_SECTION_MAPPING:
                db_section_name = self.REVERSE_SECTION_MAPPING[section_name]
            else:
                db_section_name = section_name
                
            # Use the specific formatter method if it exists
            method_name = f"_format_{db_section_name}"
            if hasattr(self, method_name) and callable(getattr(self, method_name)):
                logger.debug(f"Using specific formatter for {section_name} ({method_name})")
                await getattr(self, method_name)(doc, data)
                return
                
            # Get the display section name from the mapping or capitalize the section name
            display_section_name = self.SECTION_MAPPING.get(db_section_name, section_name.replace("_", " ").title())
            
            # Add a section heading
            doc.add_heading(display_section_name, level=1)
            
            # Only use LLM for formatting if it's brand_context
            if db_section_name == "brand_context" and self.llm:
                try:
                    # Format data for the prompt
                    format_data = {
                        "run_id": self.current_run_id,
                        section_name: json.dumps(data, indent=2) if isinstance(data, dict) else str(data),
                        "format_instructions": self._get_format_instructions(section_name)
                    }
                    
                    # Create prompt
                    prompt_content = self._format_template(section_name, format_data, section_name)
                    
                    # Create messages
                    system_content = self._get_system_content("You are an expert report formatter helping to create a professional brand naming report.")
                    messages = [
                        SystemMessage(content=system_content),
                        HumanMessage(content=prompt_content)
                    ]
                    
                    # Invoke LLM
                    response = await self._safe_llm_invoke(messages, section_name)
                    
                    # Extract JSON content
                    json_content = self._extract_json_from_response(response.content, section_name)
                    
                    if json_content:
                        # Add formatted content
                        for sub_section_name, content in json_content.items():
                            if content and isinstance(content, str):
                                doc.add_heading(sub_section_name.replace("_", " ").title(), level=2)
                                doc.add_paragraph(content)
                            elif content and isinstance(content, list):
                                doc.add_heading(sub_section_name.replace("_", " ").title(), level=2)
                                for item in content:
                                    bullet = doc.add_paragraph(style='List Bullet')
                                    bullet.add_run(item)
                        
                        return  # Successfully formatted with LLM
                        
                except Exception as e:
                    logger.error(f"Error formatting section {section_name} with LLM: {str(e)}")
                    # Fall back to standard formatting
            
            # For any other section, use direct formatting
            # Add a paragraph with the section content
            if isinstance(data, dict):
                # Process the dictionary data
                for key, value in data.items():
                    if isinstance(value, str) and value.strip():
                        doc.add_heading(key.replace("_", " ").title(), level=2)
                        doc.add_paragraph(value.strip())
                    elif isinstance(value, list):
                        doc.add_heading(key.replace("_", " ").title(), level=2)
                        for item in value:
                            if isinstance(item, str) and item.strip():
                                bullet = doc.add_paragraph(style='List Bullet')
                                bullet.add_run(item.strip())
                            elif isinstance(item, dict):
                                for sub_key, sub_value in item.items():
                                    if isinstance(sub_value, str) and sub_value.strip():
                                        doc.add_heading(sub_key.replace("_", " ").title(), level=3)
                                        doc.add_paragraph(sub_value.strip())
                    elif isinstance(value, dict):
                        doc.add_heading(key.replace("_", " ").title(), level=2)
                        for sub_key, sub_value in value.items():
                            if isinstance(sub_value, str) and sub_value.strip():
                                doc.add_heading(sub_key.replace("_", " ").title(), level=3)
                                doc.add_paragraph(sub_value.strip())
            elif isinstance(data, str) and data.strip():
                doc.add_paragraph(data.strip())
            else:
                # Add a generic message for empty or non-string, non-dict data
                doc.add_paragraph(f"No {display_section_name} data available.")
                
        except Exception as e:
            logger.error(f"Error in _format_generic_section for {section_name}: {str(e)}")
            doc.add_paragraph(f"Error formatting {section_name} section: {str(e)}")
            # Add a more generic message
            doc.add_paragraph(f"Unable to format the {display_section_name} section due to an error in processing the data.")

    async def upload_report_to_storage(self, file_path: str, run_id: str) -> str:
        """
        Upload the generated report to Supabase Storage.
        
        Args:
            file_path: Path to the local report file
            run_id: The run ID associated with the report
            
        Returns:
            The URL of the uploaded report
        """
        logger.info(f"Uploading report to Supabase Storage: {file_path}")
        
        # Extract filename from path
        filename = os.path.basename(file_path)
        
        # Define the storage path - organize by run_id
        storage_path = f"{run_id}/{filename}"
        
        # Get file size in KB
        file_size_kb = os.path.getsize(file_path) // 1024
        
        try:
            # Check if Supabase client is initialized
            if not self.supabase:
                logger.error("Supabase client is not initialized. Cannot upload report.")
                return None
                
            # Read the file content
            with open(file_path, 'rb') as f:
                file_content = f.read()
            
            # Upload to Supabase Storage with proper content type
            result = await self.supabase.storage_upload_with_retry(
                bucket=self.STORAGE_BUCKET,
                path=storage_path,
                file=file_content,
                file_options={"content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"}
            )
            
            if not result:
                logger.error("Upload completed but no result returned from Supabase")
                return None
                
            # Get the public URL
            public_url = await self.supabase.storage_get_public_url(
                bucket=self.STORAGE_BUCKET,
                path=storage_path
            )
            
            if not public_url:
                logger.error("Failed to get public URL for uploaded file")
                return None
                
            logger.info(f"Report uploaded successfully to: {public_url}")
            
            # Store metadata in report_metadata table
            await self.store_report_metadata(
                run_id=run_id,
                report_url=public_url,
                file_size_kb=file_size_kb,
                format=self.FORMAT_DOCX
            )
            
            return public_url
            
        except FileNotFoundError:
            logger.error(f"File not found: {file_path}")
            return None
        except Exception as e:
            logger.error(f"Error uploading report to storage: {str(e)}")
            logger.error(f"Error details: {traceback.format_exc()}")
            # Return None instead of raising to allow the process to continue
            return None

    async def store_report_metadata(self, run_id: str, report_url: str, file_size_kb: float, format: str) -> None:
        """
        Store metadata about the generated report in the report_metadata table.
        
        Args:
            run_id: The run ID associated with the report
            report_url: The URL where the report is accessible
            file_size_kb: The size of the report file in kilobytes
            format: The format of the report (e.g., 'docx')
        """
        logger.info(f"Storing report metadata for run_id: {run_id}")
        
        try:
            # Get the current timestamp
            timestamp = datetime.now().isoformat()
            
            # Get the current version number
            version = 1
            try:
                # Query for existing versions
                query = f"""
                SELECT MAX(version) as max_version FROM report_metadata 
                WHERE run_id = '{run_id}'
                """
                result = await self.supabase.execute_with_retry(query, {})
                if result and len(result) > 0 and result[0]['max_version'] is not None:
                    version = int(result[0]['max_version']) + 1
                    logger.info(f"Found existing reports, using version: {version}")
            except Exception as e:
                logger.warning(f"Error getting version number: {str(e)}")
                # Continue with default version 1
            
            # Insert metadata into the report_metadata table
            metadata = {
                "run_id": run_id,
                "report_url": report_url,
                "created_at": timestamp,
                "last_updated": timestamp,
                "format": format,
                "file_size_kb": file_size_kb,
                "version": version,
                "notes": f"Generated by ReportFormatter v{settings.version}"
            }
            
            result = await self.supabase.execute_with_retry(
                "insert",
                "report_metadata",
                metadata
            )
            
            logger.info(f"Report metadata stored successfully for run_id: {run_id}, version: {version}")
            
        except Exception as e:
            logger.error(f"Error storing report metadata: {str(e)}")
            logger.error(f"Error details: {traceback.format_exc()}")
            # Don't raise the exception - this is not critical for report generation

    def _extract_json_from_response(self, response_content: str, section_name: str = None) -> Dict[str, Any]:
        """
        Extract JSON content from LLM response.
        
        Args:
            response_content: The raw response content from the LLM
            section_name: Optional section name for logging
            
        Returns:
            Dict containing the parsed JSON, or None if extraction fails
        """
        try:
            # Try to extract JSON between triple backticks
            json_match = re.search(r"```(?:json)?\s*([\s\S]*?)\s*```", response_content)
            if json_match:
                json_str = json_match.group(1).strip()
                return json.loads(json_str)
            
            # If no JSON in backticks, try to parse the entire response
            return json.loads(response_content)
        except (json.JSONDecodeError, AttributeError, ValueError) as e:
            context = f" for {section_name}" if section_name else ""
            logger.error(f"Error extracting JSON from response{context}: {str(e)}")
            return None

    async def _format_market_research(self, doc: Document, data: Dict[str, Any]) -> None:
        """
        Format the market research section using direct ETL process.
        
        Args:
            doc: The document to add content to
            data: The raw market research data
        """
        try:
            # Add section title
            doc.add_heading("Market Research", level=1)
            
            # Add introduction
            doc.add_paragraph(
                "This section presents market research findings relevant to the brand naming process, "
                "including industry trends, target audience insights, and competitive landscape analysis. "
                "This information helps contextualize the brand name options within the current market environment."
            )
            
            # Transform data using model
            transformed_data = self._transform_market_research(data)
            logger.debug(f"Transformed market research data: {len(str(transformed_data))} chars")
            
            # Check if we have analysis data
            if transformed_data and "market_research" in transformed_data:
                # Process each brand's market research
                for brand_name, details in transformed_data["market_research"].items():
                    # Add brand name as heading
                    doc.add_heading(f"Market Research for {brand_name}", level=2)
                    
                    # Industry Overview
                    doc.add_heading("Industry Overview", level=3)
                    overview_table = doc.add_table(rows=4, cols=2)
                    overview_table.style = 'Table Grid'
                    
                    # Add industry overview details
                    rows = [
                        ("Industry", details["industry_name"]),
                        ("Market Size", details["market_size"]),
                        ("Growth Rate", details["market_growth_rate"]),
                        ("Market Viability", details["market_viability"])
                    ]
                    
                    for i, (label, value) in enumerate(rows):
                        cells = overview_table.rows[i].cells
                        cells[0].text = label
                        cells[1].text = str(value)
                    
                    doc.add_paragraph()  # Add spacing
                    
                    # Market Dynamics
                    doc.add_heading("Market Dynamics", level=3)
                    dynamics_table = doc.add_table(rows=5, cols=2)
                    dynamics_table.style = 'Table Grid'
                    
                    # Parse emerging trends if it's a string representation of a list
                    emerging_trends = details["emerging_trends"]
                    if isinstance(emerging_trends, str):
                        try:
                            # Try to evaluate the string as a list if it starts with [ and ends with ]
                            if emerging_trends.startswith("[") and emerging_trends.endswith("]"):
                                emerging_trends = eval(emerging_trends)
                        except:
                            # If evaluation fails, keep as is
                            pass
                    
                    # Format emerging trends
                    if isinstance(emerging_trends, list):
                        emerging_trends = "\n".join(f"• {trend.strip()}" for trend in emerging_trends)
                    
                    # Add market dynamics details
                    rows = [
                        ("Emerging Trends", emerging_trends),
                        ("Market Opportunity", details["market_opportunity"]),
                        ("Entry Barriers", details["market_entry_barriers"]),
                        ("Potential Risks", details["potential_risks"]),
                        ("Target Audience Fit", details["target_audience_fit"])
                    ]
                    
                    for i, (label, value) in enumerate(rows):
                        cells = dynamics_table.rows[i].cells
                        cells[0].text = label
                        cells[1].text = str(value)
                    
                    doc.add_paragraph()  # Add spacing
                    
                    # Competitive Landscape
                    doc.add_heading("Competitive Landscape", level=3)
                    
                    # Key Competitors
                    doc.add_paragraph("Key Competitors:")
                    competitors = details["key_competitors"]
                    if isinstance(competitors, str):
                        try:
                            # Try to evaluate the string as a list if it starts with [ and ends with ]
                            if competitors.startswith("[") and competitors.endswith("]"):
                                competitors = eval(competitors)
                        except:
                            # If evaluation fails, split by comma
                            competitors = [c.strip() for c in competitors.split(",")]
                    
                    for competitor in competitors:
                        # Remove any remaining quotes
                        competitor = str(competitor).strip("'\"")
                        doc.add_paragraph(competitor, style='List Bullet')
                    
                    # Competitive Analysis
                    doc.add_paragraph(details["competitive_analysis"])
                    
                    doc.add_paragraph()  # Add spacing
                    
                    # Customer Insights
                    doc.add_heading("Customer Pain Points", level=3)
                    pain_points = details["customer_pain_points"]
                    if isinstance(pain_points, str):
                        try:
                            # Try to evaluate the string as a list if it starts with [ and ends with ]
                            if pain_points.startswith("[") and pain_points.endswith("]"):
                                pain_points = eval(pain_points)
                        except:
                            # If evaluation fails, split by comma
                            pain_points = [p.strip() for p in pain_points.split(",")]
                    
                    for pain_point in pain_points:
                        # Remove any remaining quotes
                        pain_point = str(pain_point).strip("'\"")
                        doc.add_paragraph(pain_point, style='List Bullet')
                    
                    # Recommendations
                    doc.add_heading("Recommendations", level=3)
                    doc.add_paragraph(details["recommendations"])
                    
                    doc.add_paragraph()  # Add spacing between brand analyses
            else:
                doc.add_paragraph("No market research data available.")
                
        except Exception as e:
            logger.error(f"Error formatting market research: {str(e)}")
            doc.add_paragraph("Error formatting market research section.")

    async def _format_semantic_analysis(self, doc: Document, data: Dict[str, Any]) -> None:
        """
        Format the semantic analysis section using direct ETL process.
        
        Args:
            doc: The document to add content to
            data: The raw semantic analysis data
        """
        try:
            # Add section title
            doc.add_heading("Semantic Analysis", level=1)
            
            # Add introduction
            doc.add_paragraph(
                "This section analyzes the semantic aspects of the brand name options, "
                "including etymology, meaning, sound symbolism, and other semantic dimensions "
                "that influence brand perception and memorability."
            )
            
            # Transform data using model
            transformed_data = self._transform_semantic_analysis(data)
            logger.info(f"Transformed semantic analysis data: {len(str(transformed_data))} chars")
            
            # Check if we have brand analyses
            if "brand_analyses" in transformed_data and transformed_data["brand_analyses"]:
                brand_analyses = transformed_data["brand_analyses"]
                
                # Add a summary of the number of brand names analyzed
                doc.add_paragraph(f"Semantic analysis was conducted for {len(brand_analyses)} brand name candidates.")
                
                # Process each brand name analysis
                for analysis in brand_analyses:
                    brand_name = analysis.get("brand_name", "Unknown Brand")
                    
                    # Add brand name as heading
                    doc.add_heading(brand_name, level=2)
                    
                    # Process attributes in a structured way - updated to match the exact fields in SemanticAnalysisDetails model
                    attributes = [
                        ("etymology", "Etymology"),
                        ("sound_symbolism", "Sound Symbolism"),
                        ("brand_personality", "Brand Personality"),
                        ("emotional_valence", "Emotional Valence"),
                        ("denotative_meaning", "Denotative Meaning"),
                        ("figurative_language", "Figurative Language"),
                        ("phoneme_combinations", "Phoneme Combinations"),
                        ("sensory_associations", "Sensory Associations"),
                        ("word_length_syllables", "Word Length (Syllables)"),
                        ("alliteration_assonance", "Alliteration/Assonance"),
                        ("compounding_derivation", "Compounding/Derivation"),
                        ("semantic_trademark_risk", "Semantic Trademark Risk")
                    ]
                    
                    for attr_key, attr_display in attributes:
                        if attr_key in analysis and analysis[attr_key]:
                            # Format based on value type
                            value = analysis[attr_key]
                            
                            # Add heading
                            p = doc.add_paragraph()
                            p.add_run(f"{attr_display}: ").bold = True
                            
                            # Handle dictionary attributes
                            if isinstance(value, dict):
                                # Add each dictionary item as a bullet point
                                for sub_key, sub_value in value.items():
                                    bullet_p = doc.add_paragraph(style="List Bullet")
                                    bullet_p.add_run(f"{sub_key.replace('_', ' ').title()}: ").bold = True
                                    bullet_p.add_run(str(sub_value))
                            # Handle list attributes
                            elif isinstance(value, list):
                                # Add the attribute value as a comma-separated list
                                p.add_run(", ".join(str(item) for item in value))
                            # Handle special formatting for boolean values
                            elif isinstance(value, bool):
                                p.add_run("Yes" if value else "No")
                            # Handle simple string/number attributes
                            else:
                                p.add_run(str(value))
                    
                    # Add separator between brand analyses (except after the last one)
                    if analysis != brand_analyses[-1]:
                        doc.add_paragraph("")
                
                # Add summary section
                doc.add_heading("Semantic Analysis Summary", level=2)
                doc.add_paragraph(
                    "The semantic analysis reveals how each brand name option communicates meaning through various "
                    "linguistic dimensions. This analysis helps identify names with strong semantic foundations that "
                    "align with the desired brand positioning and minimize potential semantic confusion or negative associations."
                )
            else:
                # No semantic analysis data available
                doc.add_paragraph("No semantic analysis data available for this brand naming project.")
                
        except Exception as e:
            logger.error(f"Error formatting semantic analysis section: {str(e)}")
            logger.debug(f"Error details: {traceback.format_exc()}")
            # Add a generic error message to the document
            doc.add_paragraph("Unable to format the semantic analysis section due to an error in processing the data.")

    async def _format_translation_analysis(self, doc: Document, data: Dict[str, Any]) -> None:
        """
        Format the translation analysis section using direct ETL process.
        
        Args:
            doc: The document to add content to
            data: The raw translation analysis data
        """
        try:
            # Add section title
            doc.add_heading("Translation Analysis", level=1)
            
            # Add introduction
            doc.add_paragraph(
                "This section analyzes how the brand name options translate across different languages and cultures, "
                "identifying potential issues, unintended meanings, or pronunciation challenges that could impact "
                "global brand perception."
            )
            
            # Transform data using model
            transformed_data = self._transform_translation_analysis(data)
            logger.info(f"Transformed translation analysis data: {len(str(transformed_data))} chars")
            
            # Check if we have analysis data with the expected structure
            if transformed_data and "translation_analysis" in transformed_data:
                translation_analysis = transformed_data["translation_analysis"]
                
                # Add a summary of the number of brand names analyzed
                brand_count = len(translation_analysis)
                doc.add_paragraph(f"Translation analysis was conducted for {brand_count} brand name candidates.")
                
                # Process each brand name
                for brand_name, languages in translation_analysis.items():
                    # Add brand name as heading
                    doc.add_heading(brand_name, level=2)
                    
                    # Process each language analysis for this brand name
                    for language_name, language_analysis in languages.items():
                        # Add language heading
                        doc.add_heading(f"{language_name}", level=3)
                        
                        # Process attributes in a structured way
                        attributes = [
                            ("target_language", "Target Language"),
                            ("direct_translation", "Direct Translation"),
                            ("semantic_shift", "Semantic Shift"),
                            ("phonetic_retention", "Phonetic Retention"),
                            ("pronunciation_difficulty", "Pronunciation Difficulty"),
                            ("adaptation_needed", "Adaptation Needed"),
                            ("proposed_adaptation", "Proposed Adaptation"),
                            ("cultural_acceptability", "Cultural Acceptability"),
                            ("brand_essence_preserved", "Brand Essence Preserved"),
                            ("global_consistency_vs_localization", "Global Consistency vs Localization"),
                            ("notes", "Additional Notes")
                        ]
                        
                        # Display each attribute if present
                        for attr, label in attributes:
                            if attr in language_analysis and language_analysis[attr] is not None:
                                value = language_analysis[attr]
                                
                                # Special handling for boolean values
                                if isinstance(value, bool):
                                    value = "Yes" if value else "No"
                                
                                p = doc.add_paragraph()
                                p.add_run(f"{label}: ").bold = True
                                p.add_run(str(value))
                
                        # Add separator between brand names (except for the last one)
                        if brand_name != list(translation_analysis.keys())[-1]:
                            doc.add_paragraph("", style="Normal")
                    
                    # Add a summary section
                    doc.add_heading("Translation Analysis Summary", level=2)
                    doc.add_paragraph(
                        "The analysis above evaluates how each brand name option would perform across different languages "
                        "and cultural contexts. This assessment is crucial for brands with global aspirations, as it "
                        "identifies potential challenges and opportunities for international brand positioning."
                    )
                
            else:
                # No translation analysis data available or unexpected format
                doc.add_paragraph("No translation analysis data available for this brand naming project.")
                
        except Exception as e:
            logger.error(f"Error formatting translation analysis section: {str(e)}")
            logger.debug(f"Error details: {traceback.format_exc()}")
            # Add a generic error message to the document
            doc.add_paragraph("Unable to format the translation analysis section due to an error in processing the data.")

    async def _format_domain_analysis(self, doc: Document, data: Dict[str, Any]) -> None:
        """
        Format the domain analysis section using direct ETL process.
        
        Args:
            doc: The document to add content to
            data: The raw domain analysis data
        """
        try:
            # Add section title
            doc.add_heading("Domain Analysis", level=1)
            
            # Add introduction
            doc.add_paragraph(
                "This section analyzes the domain availability and suitability of the brand name options, "
                "evaluating factors such as domain extension options, pricing, and potential alternatives. "
                "Domain availability is a critical factor in today's digital business environment."
            )
            
            # Transform data using model
            transformed_data = self._transform_domain_analysis(data)
            logger.debug(f"Transformed domain analysis data: {len(str(transformed_data))} chars")
            
            # Check if we have analysis data
            if transformed_data and isinstance(transformed_data, dict):
                # First, create a comparison table between all brands
                doc.add_heading("Domain Analysis Overview", level=2)
                
                # Create comparison table
                comparison_table = doc.add_table(rows=1 + len(transformed_data), cols=5)
                comparison_table.style = 'Table Grid'
                
                # Add header row
                header_cells = comparison_table.rows[0].cells
                header_cells[0].text = "Brand Name"
                header_cells[1].text = "Domain Available"
                header_cells[2].text = "Acquisition Cost"
                header_cells[3].text = "Brand Name Clarity"
                header_cells[4].text = "Available TLDs"
                
                # Add data rows
                for i, (brand_name, analysis) in enumerate(transformed_data.items(), 1):
                    cells = comparison_table.rows[i].cells
                    cells[0].text = brand_name
                    cells[1].text = "Yes" if analysis.get("domain_exact_match", False) else "No"
                    cells[2].text = analysis.get("acquisition_cost", "Unknown")
                    cells[3].text = analysis.get("brand_name_clarity_in_url", "Unknown")
                    cells[4].text = ", ".join(f".{tld}" for tld in analysis.get("alternative_tlds", []))
                
                doc.add_paragraph()  # Add spacing
                
                # Then show detailed analysis for each brand
                doc.add_heading("Detailed Domain Analysis by Brand", level=2)
                
                # Process each brand name analysis
                for brand_name, analysis in transformed_data.items():
                    # Add brand name as heading
                    doc.add_heading(brand_name, level=3)
                    
                    # Domain Quality Metrics Table
                    quality_table = doc.add_table(rows=7, cols=2)
                    quality_table.style = 'Table Grid'
                    
                    # Fill quality table with all model fields
                    cells = quality_table.rows[0].cells
                    cells[0].text = "Exact Match Domain"
                    cells[1].text = "Available" if analysis.get("domain_exact_match", False) else "Not Available"
                    
                    cells = quality_table.rows[1].cells
                    cells[0].text = "Acquisition Cost"
                    cells[1].text = analysis.get("acquisition_cost", "Unknown")
                    
                    cells = quality_table.rows[2].cells
                    cells[0].text = "Domain Format"
                    cells[1].text = "Clean format" if not analysis.get("hyphens_numbers_present", False) else "Contains hyphens or numbers"
                    
                    cells = quality_table.rows[3].cells
                    cells[0].text = "Brand Name Clarity"
                    cells[1].text = analysis.get("brand_name_clarity_in_url", "Unknown")
                    
                    cells = quality_table.rows[4].cells
                    cells[0].text = "Domain Length and Readability"
                    cells[1].text = analysis.get("domain_length_readability", "Unknown")
                    
                    cells = quality_table.rows[5].cells
                    cells[0].text = "Misspellings Available"
                    cells[1].text = "Yes" if analysis.get("misspellings_variations_available", False) else "No"
                    
                    cells = quality_table.rows[6].cells
                    cells[0].text = "Available TLDs"
                    cells[1].text = ", ".join(f".{tld}" for tld in analysis.get("alternative_tlds", []))
                    
                    doc.add_paragraph()  # Add spacing
                    
                    # Social Media Availability
                    if analysis.get("social_media_availability"):
                        doc.add_heading("Social Media Handles", level=4)
                        social_table = doc.add_table(rows=len(analysis["social_media_availability"]), cols=1)
                        social_table.style = 'Table Grid'
                        
                        for i, handle in enumerate(analysis["social_media_availability"]):
                            social_table.rows[i].cells[0].text = handle
                        
                        doc.add_paragraph()  # Add spacing
                    
                    # Scalability and Future-Proofing
                    if analysis.get("scalability_future_proofing"):
                        doc.add_heading("Scalability Assessment", level=4)
                        doc.add_paragraph(analysis.get("scalability_future_proofing", ""))
                    
                    # Analysis Notes
                    if analysis.get("notes"):
                        doc.add_heading("Analysis and Recommendations", level=4)
                        doc.add_paragraph(analysis.get("notes", ""))
                    
                    # Add separator between brand analyses (except for the last one)
                    if brand_name != list(transformed_data.keys())[-1]:
                        doc.add_paragraph("")
            else:
                # No domain analysis data available
                doc.add_paragraph("No domain analysis data available for this brand naming project.")
                
        except Exception as e:
            logger.error(f"Error formatting domain analysis section: {str(e)}")
            logger.debug(f"Error details: {traceback.format_exc()}")
            doc.add_paragraph("Unable to format the domain analysis section due to an error in processing the data.")

    async def _add_title_page(self, doc: Document, data: Dict[str, Any]) -> None:
        """Add a title page to the document."""
        logger.info("Adding title page")
        
        try:
            # Extract brand context for the title page
            brand_context = data.get("brand_context", {})
            
            # Fetch the user prompt from the state
            user_prompt = await self.fetch_user_prompt(self.current_run_id)
            
            # Format data for the prompt
            format_data = {
                "run_id": self.current_run_id,
                "brand_context": json.dumps(brand_context, indent=2) if isinstance(brand_context, dict) else str(brand_context),
                "user_prompt": user_prompt
            }
            
            # No need to call the LLM for standard title and subtitle
            # Use the exact formats from notesforformatter.md
            title = "Brand Naming Report"
            subtitle = "Generated by Mae Brand Naming Expert"
            
            # Create the title page
            # Set title with large font and center alignment
            title_para = doc.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_para.add_run(title)
            title_run.font.size = Pt(24)
            title_run.font.bold = True
            title_para.space_after = Pt(12)
            
            # Add subtitle
            subtitle_para = doc.add_paragraph()
            subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle_run = subtitle_para.add_run(subtitle)
            subtitle_run.font.size = Pt(16)
            subtitle_run.italic = True
            subtitle_para.space_after = Pt(36)
            
            # Add date
            date_para = doc.add_paragraph()
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            date_para.add_run(f"Generated on: {datetime.now().strftime('%B %d, %Y')}")
            date_para.space_after = Pt(12)
            
            # Add run ID
            run_id_para = doc.add_paragraph()
            run_id_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_id_para.add_run(f"Run ID: {self.current_run_id}")
            run_id_para.space_after = Pt(24)
            
            # Add user prompt note
            user_prompt_para = doc.add_paragraph()
            user_prompt_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            user_prompt_run = user_prompt_para.add_run(
                f"This report was generated using the Maestro AI Agent workflow and all content "
                f"has been derived from this prompt: \"{user_prompt}\""
            )
            user_prompt_run.italic = True
            user_prompt_run.font.size = Pt(10)
            
            # Add page break after title page
            doc.add_page_break()
            
        except Exception as e:
            logger.error(f"Error adding title page: {str(e)}")
            # Add fallback title page
            doc.add_paragraph("Brand Naming Report").style = 'Title'
            doc.add_paragraph("Generated by Mae Brand Naming Expert").style = 'Subtitle'
            doc.add_paragraph(f"Generated on: {datetime.now().strftime('%B %d, %Y')}")
            doc.add_paragraph(f"Run ID: {self.current_run_id}")
            doc.add_page_break()

    async def store_report_metadata(self, run_id: str, report_url: str, file_size_kb: float, format: str) -> None:
        """Store report metadata in the database."""
        logger.info(f"Storing report metadata for run_id: {run_id}")
        
        try:
            # Get the current timestamp
            timestamp = datetime.now().isoformat()
            
            # Get the current version number
            version = 1
            try:
                # Query for existing versions
                query = f"""
                SELECT MAX(version) as max_version FROM report_metadata 
                WHERE run_id = '{run_id}'
                """
                result = await self.supabase.execute_with_retry(query, {})
                if result and len(result) > 0 and result[0]['max_version'] is not None:
                    version = int(result[0]['max_version']) + 1
                    logger.info(f"Found existing reports, using version: {version}")
            except Exception as e:
                logger.warning(f"Error getting version number: {str(e)}")
                # Continue with default version 1
            
            # Insert metadata into the report_metadata table
            metadata = {
                "run_id": run_id,
                "report_url": report_url,
                "created_at": timestamp,
                "last_updated": timestamp,
                "format": format,
                "file_size_kb": file_size_kb,
                "version": version,
                "notes": f"Generated by ReportFormatter v{settings.version}"
            }
            
            result = await self.supabase.execute_with_retry(
                "insert",
                "report_metadata",
                metadata
            )
            
            logger.info(f"Report metadata stored successfully for run_id: {run_id}, version: {version}")
            
        except Exception as e:
            logger.error(f"Error storing report metadata: {str(e)}")
            logger.error(f"Error details: {traceback.format_exc()}")
            # Don't raise the exception - this is not critical for report generation

    async def generate_report(self, run_id: str, upload_to_storage: bool = True) -> tuple[str, str]:
        """
        Generate a formatted report for the given run ID.
        
        Args:
            run_id: The run ID to generate a report for
            upload_to_storage: Whether to upload the report to Supabase storage (default: True)
            
        Returns:
            tuple[str, str]: A tuple containing (file_path, report_url)
                             report_url will be empty if upload_to_storage is False
        """
        logger.info(f"Generating report for run ID: {run_id}")
        
        # Initialize run ID
        self.current_run_id = run_id
        
        # Initialize LLM if not already done
        if not self.llm:
            from ..config.settings import settings
            self.llm = ChatGoogleGenerativeAI(
                model=settings.model_name, 
                temperature=0.2,  # Balanced temperature for analysis 
                google_api_key=settings.google_api_key, 
                convert_system_message_to_human=True, 
                callbacks=settings.get_langsmith_callbacks()
            )
        
        # Fetch raw data
        data = await self.fetch_raw_data(run_id)
        
        if not data:
            logger.error(f"No data found for run ID: {run_id}")
            return None, None
        
        # Create document
        doc = Document()
        
        # Set up document styles
        self._setup_document_styles(doc)
        
        # Add title page - Required to be first!
        await self._add_title_page(doc, data)
        
        # Add table of contents - Required to be second!
        await self._add_table_of_contents(doc)
        
        # Add executive summary - Required to be third!
        # Always call the executive summary generation regardless of whether the data exists
        await self._add_executive_summary(doc, {})
        
        # Process each section in order
        for section_name in self.SECTION_ORDER:
            # Skip already processed sections
            if section_name in ["exec_summary", "final_recommendations"]:
                continue
                
            # Check if section exists in data
            if section_name in data:
                section_data = data[section_name]
                logger.info(f"Processing section: {section_name}")
                
                # Map section name to formatter method name
                formatter_method_name = f"_format_{self.REVERSE_SECTION_MAPPING.get(section_name, section_name)}"
                
                # Format the section based on its name
                if hasattr(self, formatter_method_name):
                    # Call the specific formatter method
                    formatter_method = getattr(self, formatter_method_name)
                    logger.info(f"Using specific formatter method: {formatter_method_name}")
                    await formatter_method(doc, section_data)
                else:
                    # Use generic formatter as fallback
                    logger.info(f"Using generic formatter for section: {section_name}")
                    await self._format_generic_section(doc, section_name, section_data)
            else:
                logger.warning(f"Section {section_name} not found in data")
        
        # Add recommendations section at the end
        if "final_recommendations" in data:
            await self._add_recommendations(doc, data["final_recommendations"])
        
        # Save document to a temporary file
        temp_dir = Path("tmp")
        temp_dir.mkdir(exist_ok=True)
        file_path = temp_dir / f"report_{run_id}.docx"
        
        doc.save(str(file_path))
        logger.info(f"Report saved to: {file_path}")
        
        # Upload to storage if available and requested
        report_url = None
        if upload_to_storage and self.supabase:
            try:
                logger.info("Attempting to upload report to Supabase storage")
                report_url = await self.upload_report_to_storage(str(file_path), run_id)
                
                if report_url:
                    logger.info(f"Report successfully uploaded to: {report_url}")
                    # File size in KB
                    file_size_kb = file_path.stat().st_size / 1024
                    # Store metadata
                    await self.store_report_metadata(run_id, report_url, file_size_kb, self.FORMAT_DOCX)
                else:
                    logger.error("Failed to upload report to Supabase storage - report_url is None")
            except Exception as e:
                logger.error(f"Unexpected error during upload process: {str(e)}")
                logger.error(traceback.format_exc())
        elif not upload_to_storage:
            logger.info("Skipping upload to storage as requested (upload_to_storage=False)")
        else:
            logger.warning("Supabase client not available - skipping upload to storage")
        
        return str(file_path), report_url

async def main(run_id: str = None, upload_to_storage: bool = True):
    """Main function to run the formatter."""
    if not run_id:
        # Use a default run ID for testing
        run_id = "mae_20250312_141302_d45cccde"  # Replace with an actual run ID
        
    formatter = ReportFormatter()
    output_path, report_url = await formatter.generate_report(run_id, upload_to_storage=upload_to_storage)
    print(f"Report generated at: {output_path}")
    print(f"Report URL: {report_url}")


if __name__ == "__main__":
    # Command-line interface
    import argparse
    
    parser = argparse.ArgumentParser(description="Generate a formatted report from raw data")
    parser.add_argument("run_id", help="The run ID to generate a report for")
    parser.add_argument("--no-upload", dest="upload_to_storage", action="store_false", 
                        help="Skip uploading the report to Supabase storage")
    parser.set_defaults(upload_to_storage=True)
    args = parser.parse_args()
    
    asyncio.run(main(args.run_id, args.upload_to_storage)) 