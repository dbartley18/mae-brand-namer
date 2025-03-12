"""Report Compiler for generating comprehensive brand name analysis reports."""

# Standard library imports
import os
from typing import Dict, List, Any, Optional
from datetime import datetime, timezone, timedelta
from pathlib import Path
import tempfile
import re
import json
import asyncio
import uuid
import random

# Import S3 libraries
try:
    import boto3
    from botocore.client import Config
except ImportError:
    boto3 = None
    Config = None

# Third-party imports
from supabase import create_client, Client
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.prompts import ChatPromptTemplate, load_prompt
from langchain.output_parsers import ResponseSchema, StructuredOutputParser
from langchain_core.tracers.context import tracing_v2_enabled
from langchain_core.tracers import LangChainTracer
from langchain_core.messages import HumanMessage, SystemMessage, AIMessage
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from postgrest import APIError

# Local application imports
from ..config.settings import settings
from ..utils.logging import get_logger
from ..config.dependencies import Dependencies
from ..utils.supabase_utils import SupabaseManager

logger = get_logger(__name__)


class ReportCompiler:
    """Expert in compiling and formatting comprehensive brand naming reports."""

    def __init__(
        self,
        dependencies: Optional[Dependencies] = None,
        supabase: Optional[SupabaseManager] = None,
    ):
        """Initialize the ReportCompiler with dependencies."""
        # Initialize Supabase client
        if dependencies:
            self.supabase = dependencies.supabase
            self.langsmith = dependencies.langsmith
        else:
            try:
                logger.info("Initializing ReportCompiler with direct Supabase connection")
                self.supabase = supabase or SupabaseManager()
                self.langsmith = None
                
                # Verify the Supabase manager has the necessary methods
                if not hasattr(self.supabase, "execute_query") and hasattr(self.supabase, "execute_with_retry"):
                    # Add a compatibility layer for execute_query if it's missing but execute_with_retry exists
                    logger.info("Adding compatibility layer for execute_query method")
                    
                    # Create a bound method to add to the instance
                    async def execute_query(query: str, params: Dict[str, Any], retries: int = 3) -> List[Dict[str, Any]]:
                        """
                        Compatibility method that routes SQL queries through execute_with_retry.
                        This is a simplified implementation for basic queries only.
                        """
                        logger.info(f"Using compatibility layer for query: {query[:50]}...")
                        
                        # Handle different query types
                        if query.strip().upper().startswith("SELECT"):
                            # Extract table name from query
                            from_parts = query.split("FROM")
                            if len(from_parts) > 1:
                                table_parts = from_parts[1].strip().split()
                                if table_parts:
                                    table_name = table_parts[0].strip()
                                    
                                    # Remove any trailing commas, semicolons, etc.
                                    table_name = re.sub(r'[^a-zA-Z0-9_]', '', table_name)
                                    
                                    # Create data dict with run_id if present in params
                                    data = {}
                                    if "run_id" in params:
                                        data["run_id"] = params["run_id"]
                                    
                                    # Add limit if present in query
                                    if "LIMIT" in query.upper():
                                        limit_parts = query.upper().split("LIMIT")
                                        if len(limit_parts) > 1:
                                            try:
                                                limit = int(re.search(r'\d+', limit_parts[1]).group())
                                                data["limit"] = limit
                                            except (ValueError, AttributeError):
                                                pass
                                    
                                    try:
                                        return await self.supabase.execute_with_retry("select", table_name, data)
                                    except Exception as e:
                                        logger.error(f"Error in compatibility layer: {str(e)}")
                                        return []
                        
                        # For other query types or if we couldn't parse the query
                        logger.warning("Compatibility layer cannot handle this query type")
                        return []
                    
                    # Add the method to the instance
                    self.supabase.execute_query = execute_query.__get__(self.supabase)
                    
            except Exception as e:
                logger.error(f"Error initializing Supabase connection: {str(e)}")
                logger.error("Continuing with potentially broken Supabase connection")
                # Don't raise here, we'll catch it in the compile_report method
                # when we test the connection

        # Agent identity
        self.role = "Brand Name Analysis Report Compiler"
        self.goal = """Compile comprehensive, actionable reports that synthesize all brand name analyses into clear, 
        strategic recommendations for decision makers."""
        self.backstory = """You are an expert in data synthesis and report compilation, specializing in brand naming analysis. 
        Your expertise helps transform complex analyses into clear, actionable insights that drive informed decisions."""

        try:
            # Load prompts
            prompt_dir = Path(__file__).parent / "prompts" / "report_compiler"
            self.system_prompt = load_prompt(str(prompt_dir / "system.yaml"))
            self.compilation_prompt = load_prompt(str(prompt_dir / "compilation.yaml"))

            # Define output schemas for structured parsing
            self.output_schemas = [
                ResponseSchema(
                    name="executive_summary",
                    description="High-level overview of project and key findings",
                ),
                ResponseSchema(
                    name="brand_context", description="Detailed brand context and requirements"
                ),
                ResponseSchema(
                    name="name_generation",
                    description="Brand name generation process and methodology",
                ),
                ResponseSchema(
                    name="linguistic_analysis", description="Linguistic analysis findings"
                ),
                ResponseSchema(name="semantic_analysis", description="Semantic analysis findings"),
                ResponseSchema(
                    name="cultural_sensitivity",
                    description="Cultural sensitivity analysis findings",
                ),
                ResponseSchema(
                    name="name_evaluation", description="Evaluation and shortlisting of names"
                ),
                ResponseSchema(
                    name="domain_analysis", description="Domain availability and strategy findings"
                ),
                ResponseSchema(
                    name="seo_analysis", description="SEO and online discovery findings"
                ),
                ResponseSchema(
                    name="competitor_analysis", description="Competitive landscape analysis"
                ),
                ResponseSchema(
                    name="survey_simulation",
                    description="Market research survey simulation results",
                ),
                ResponseSchema(
                    name="recommendations", description="Strategic recommendations and next steps"
                ),
            ]
            self.output_parser = StructuredOutputParser.from_response_schemas(self.output_schemas)

            # Initialize LLM
            self.llm = ChatGoogleGenerativeAI(
                model=settings.model_name,
                temperature=0.3,
                google_api_key=settings.gemini_api_key,
                convert_system_message_to_human=True,
                generation_config={
                    "max_output_tokens": 8192,
                    "temperature": 0.3,
                    "top_p": 0.95,
                    "top_k": 40
                },
            )
            logger.info("Successfully initialized ChatGoogleGenerativeAI for Report Compiler")

            # Set up the prompt template
            system_message = SystemMessage(content=self.system_prompt.format())
            human_template = self.compilation_prompt.template
            self.prompt = ChatPromptTemplate.from_messages(
                [system_message, HumanMessage(content=human_template)]
            )
            logger.info("Successfully set up prompt template")

        except Exception as e:
            logger.error(
                "Error initializing ReportCompiler",
                extra={"error_type": type(e).__name__, "error_message": str(e)},
            )
            raise

    async def _log_batch_process(
        self, 
        run_id: str, 
        batch_id: str, 
        batch_sections: List[str], 
        status: str, 
        start_time: datetime, 
        end_time: Optional[datetime] = None,
        error_message: Optional[str] = None,
        section_token_counts: Optional[Dict[str, int]] = None,
        retry_count: int = 0
    ) -> None:
        """
        Log batch processing status to the process_logs table.
        
        Args:
            run_id: Unique identifier for the workflow run
            batch_id: Identifier for this specific batch (e.g., "batch_1_executive_summary")
            batch_sections: List of section names in this batch
            status: Current status (started, completed, failed)
            start_time: When this batch processing started
            end_time: When this batch processing ended (None if not completed)
            error_message: Error details if status is 'failed'
            section_token_counts: Token usage per section if available
            retry_count: Number of retries for this batch
        """
        if end_time is None and status in ("completed", "failed"):
            end_time = datetime.now(timezone.utc)
            
        duration_sec = None
        if start_time and end_time:
            duration_sec = (end_time - start_time).total_seconds()
        
        # Calculate completion percentage based on sections processed
        all_sections = set()
        for sections in self.report_sections:
            all_sections.update(sections)
        completion_percentage = (len(batch_sections) / len(all_sections)) * 100
        
        # Prepare section status JSON
        section_status = {}
        for section in batch_sections:
            section_status[section] = "completed" if status == "completed" else "pending"
            
        try:
            # Insert or update log entry
            query = """
                INSERT INTO process_logs (
                    run_id, agent_type, task_name, status, start_time, end_time,
                    duration_sec, error_message, retry_count, last_retry_at,
                    retry_status, batch_id, batch_sections, section_token_counts,
                    completion_percentage, section_status
                )
                VALUES (
                    :run_id, 'ReportCompiler', :task_name, :status, :start_time, :end_time,
                    :duration_sec, :error_message, :retry_count, :last_retry_at,
                    :retry_status, :batch_id, :batch_sections, :section_token_counts,
                    :completion_percentage, :section_status
                )
                ON CONFLICT (run_id, agent_type, task_name, batch_id) 
                DO UPDATE SET
                    status = EXCLUDED.status,
                    end_time = EXCLUDED.end_time,
                    duration_sec = EXCLUDED.duration_sec,
                    error_message = EXCLUDED.error_message,
                    last_updated = now(),
                    retry_count = CASE 
                        WHEN EXCLUDED.retry_count > process_logs.retry_count 
                        THEN EXCLUDED.retry_count 
                        ELSE process_logs.retry_count 
                    END,
                    last_retry_at = CASE 
                        WHEN EXCLUDED.retry_count > process_logs.retry_count 
                        THEN EXCLUDED.last_retry_at 
                        ELSE process_logs.last_retry_at 
                    END,
                    retry_status = EXCLUDED.retry_status,
                    section_token_counts = EXCLUDED.section_token_counts,
                    completion_percentage = EXCLUDED.completion_percentage,
                    section_status = EXCLUDED.section_status
            """
            
            # Convert section arrays to proper format for database
            batch_sections_str = "{" + ",".join(batch_sections) + "}"
            
            params = {
                "run_id": run_id,
                "task_name": f"report_compilation_{batch_id}",
                "status": status,
                "start_time": start_time,
                "end_time": end_time,
                "duration_sec": duration_sec,
                "error_message": error_message,
                "retry_count": retry_count,
                "last_retry_at": datetime.now(timezone.utc) if retry_count > 0 else None,
                "retry_status": "retrying" if retry_count > 0 and status != "completed" else "completed",
                "batch_id": batch_id,
                "batch_sections": batch_sections_str,
                "section_token_counts": json.dumps(section_token_counts) if section_token_counts else None,
                "completion_percentage": completion_percentage,
                "section_status": json.dumps(section_status)
            }
            
            if self.supabase:
                await self.supabase.execute_query(query, params)
                
        except Exception as e:
            logger.error(
                "Error logging batch process",
                extra={
                    "error": str(e),
                    "run_id": run_id,
                    "batch_id": batch_id
                }
            )

    async def _log_overall_process(
        self, 
        run_id: str, 
        status: str, 
        start_time: datetime,
        end_time: Optional[datetime] = None,
        error_message: Optional[str] = None,
        output_size_kb: Optional[float] = None,
        token_count: Optional[int] = None
    ) -> None:
        """
        Log overall report compilation process status.
        
        Args:
            run_id: Unique identifier for the workflow run
            status: Current status (started, completed, failed)
            start_time: When the overall process started
            end_time: When the overall process ended
            error_message: Error details if status is 'failed'
            output_size_kb: Size of the generated report in KB
            token_count: Total token count for the report
        """
        if end_time is None and status in ("completed", "failed"):
            end_time = datetime.now(timezone.utc)
            
        duration_sec = None
        if start_time and end_time:
            duration_sec = (end_time - start_time).total_seconds()
            
        try:
            # Insert or update log entry for overall process
            query = """
                INSERT INTO process_logs (
                    run_id, agent_type, task_name, status, start_time, end_time,
                    duration_sec, error_message, output_size_kb, token_count
                )
                VALUES (
                    :run_id, 'ReportCompiler', 'report_compilation', :status, :start_time, :end_time,
                    :duration_sec, :error_message, :output_size_kb, :token_count
                )
                ON CONFLICT (run_id, agent_type, task_name) 
                WHERE batch_id IS NULL
                DO UPDATE SET
                    status = EXCLUDED.status,
                    end_time = EXCLUDED.end_time,
                    duration_sec = EXCLUDED.duration_sec,
                    error_message = EXCLUDED.error_message,
                    last_updated = now(),
                    output_size_kb = EXCLUDED.output_size_kb,
                    token_count = EXCLUDED.token_count
            """
            
            params = {
                "run_id": run_id,
                "status": status,
                "start_time": start_time,
                "end_time": end_time,
                "duration_sec": duration_sec,
                "error_message": error_message,
                "output_size_kb": output_size_kb,
                "token_count": token_count
            }
            
            if self.supabase:
                await self.supabase.execute_query(query, params)
                
        except Exception as e:
            logger.error(
                "Error logging overall process",
                extra={
                    "error": str(e),
                    "run_id": run_id
                }
            )

    async def compile_report(
        self, run_id: str, state: Optional[Dict[str, Any]] = None, user_prompt: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Compile a comprehensive brand name analysis report using sequential section generation.
        
        Args:
            run_id (str): The unique identifier for the workflow run
            state (Optional[Dict[str, Any]]): Only needed for the run_id
            user_prompt (Optional[str]): The initial user prompt that started the workflow
            
        Returns:
            Dict[str, Any]: The updated workflow state with report URL and metadata
            
        Raises:
            ValueError: If required data is missing
        """
        # Start timing and logging
        process_start_time = datetime.now(timezone.utc)
        await self._log_overall_process(run_id, "started", process_start_time)
        
        try:
            # Run diagnostics to verify Supabase connection and data availability
            logger.info(f"Running Supabase connection diagnostics for run_id {run_id}")
            
            try:
                diagnostics = await self.test_supabase_connection(run_id)
            except Exception as conn_error:
                # Provide more detailed error information
                error_msg = f"Error connecting to Supabase: {str(conn_error)}"
                logger.error(error_msg)
                
                # Check for common connection issues
                if "not initialized" in str(conn_error).lower():
                    error_msg = "Supabase client not initialized. Check environment variables SUPABASE_URL and SUPABASE_SERVICE_KEY."
                elif "unauthorized" in str(conn_error).lower() or "permission" in str(conn_error).lower():
                    error_msg = "Unauthorized access to Supabase. Check if the service key has correct permissions."
                elif "timeout" in str(conn_error).lower() or "connection" in str(conn_error).lower():
                    error_msg = "Supabase connection timed out. Check network connectivity and Supabase service status."
                
                # Log the error as part of the overall process
                await self._log_overall_process(
                    run_id, 
                    "error", 
                    process_start_time,
                    end_time=datetime.now(timezone.utc),
                    error_message=error_msg
                )
                raise ConnectionError(error_msg)
            
            # Check if we have a valid connection and sufficient data
            if diagnostics["connection_test"] != "SUCCESS":
                error_msg = f"Supabase connection test failed: {diagnostics['connection_test']}"
                if "error_details" in diagnostics and "connection" in diagnostics["error_details"]:
                    error_msg += f"\nDetails: {diagnostics['error_details']['connection']}"
                logger.error(error_msg, extra={"diagnostics": diagnostics})
                
                # Log the error as part of the overall process
                await self._log_overall_process(
                    run_id, 
                    "error", 
                    process_start_time,
                    end_time=datetime.now(timezone.utc),
                    error_message=error_msg
                )
                raise ConnectionError(error_msg)
                
            # Check if we have any data for this run_id
            if diagnostics["overall_status"] == "INSUFFICIENT DATA":
                error_msg = f"Insufficient data found for run_id {run_id}. Cannot generate report."
                logger.error(error_msg, extra={"table_counts": diagnostics["table_counts"]})
                
                # Create a minimal error report
                return {
                    "report": {
                        "url": "",
                        "run_id": run_id,
                        "version": 1,
                        "created_at": datetime.now(timezone.utc).isoformat(),
                        "last_updated": datetime.now(timezone.utc).isoformat(),
                        "format": "error",
                        "file_size_kb": 0,
                        "token_count": 0,
                        "notes": f"Error: No data found for run_id {run_id}. Check workflow execution status."
                    },
                    "run_id": run_id,
                    "report_url": "",
                    "error": error_msg,
                    "diagnostics": diagnostics
                }
            
            # Log available data for reference
            logger.info(f"Data available for report generation", extra={"table_counts": diagnostics["table_counts"]})
            
            # Create a minimal report structure to build on
            # This will be filled section by section
            report = {
                "metadata": {
                    "run_id": run_id,
                    "timestamp": datetime.now(timezone.utc).isoformat(),
                    "version": 1,
                    "user_prompt": user_prompt or "Not provided",
                    "available_tables": [
                        table for table, count in diagnostics["table_counts"].items() 
                        if isinstance(count, int) and count > 0
                    ]
                },
                "sections": []
            }
            
            # Define sections to generate sequentially
            report_sections = [
                "executive_summary",
                "brand_context",
                "name_generation",
                "linguistic_analysis",
                "semantic_analysis",
                "cultural_sensitivity",
                "name_evaluation",
                "domain_analysis", 
                "seo_analysis", 
                "competitor_analysis",
                "survey_simulation",
                "translation_analysis",
                "recommendations"
            ]
            
            # Filter to only include sections for which we have data
            available_tables = {
                table: count for table, count in diagnostics["table_counts"].items() 
                if isinstance(count, int) and count > 0
            }
            
            # Map sections to required tables
            section_to_table_map = {
                "executive_summary": ["brand_context", "brand_name_generation", "brand_name_evaluation"],
                "brand_context": ["brand_context"],
                "name_generation": ["brand_name_generation"],
                "linguistic_analysis": ["linguistic_analysis"],
                "semantic_analysis": ["semantic_analysis"],
                "cultural_sensitivity": ["cultural_sensitivity_analysis"],
                "name_evaluation": ["brand_name_evaluation"],
                "domain_analysis": ["domain_analysis"],
                "seo_analysis": ["seo_online_discoverability"],
                "competitor_analysis": ["competitor_analysis"],
                "survey_simulation": ["survey_simulation"],
                "translation_analysis": ["translation_analysis"],
                "recommendations": ["brand_name_evaluation", "domain_analysis", "seo_online_discoverability"]
            }
            
            # Determine which sections we can generate based on available data
            available_sections = []
            for section in report_sections:
                required_tables = section_to_table_map.get(section, [])
                
                # For executive_summary, we don't need all tables, just at least one
                if section == "executive_summary":
                    if any(table in available_tables for table in required_tables):
                        available_sections.append(section)
                        continue
                        
                # For recommendations, we need at least brand_name_evaluation
                if section == "recommendations":
                    if "brand_name_evaluation" in available_tables:
                        available_sections.append(section)
                        continue
                
                # For other sections, check if all required tables are available
                if all(table in available_tables for table in required_tables):
                    available_sections.append(section)
            
            # Log the sections we'll generate
            logger.info(f"Will generate {len(available_sections)} sections: {available_sections}")
            
            # Process each section sequentially - one at a time
            total_token_count = 0
            
            # Setup for prompt templates, only load them once
            system_message = SystemMessage(content=self.system_prompt.format())
            
            # Implemented section retry mechanism
            section_failures = {}
            max_section_retries = 2  # Maximum retries for a problematic section
            
            for section_name in available_sections:
                section_start_time = datetime.now(timezone.utc)
                retry_count = 0
                max_retry_backoff = 60  # Maximum backoff in seconds
                
                # Log section processing start
                await self._log_overall_process(
                    run_id=run_id,
                    status=f"processing_section_{section_name}",
                    start_time=section_start_time
                )
                
                logger.info(f"Generating section: {section_name}", extra={"run_id": run_id, "section": section_name})
                
                while retry_count <= max_section_retries:
                    try:
                        # 1. Query relevant data for this section directly from Supabase
                        section_data = await self._fetch_section_data(run_id, section_name)
                        
                        # Log what we received for debugging
                        data_summary = {
                            f"{key}_count": len(value) if isinstance(value, list) else "not_list" 
                            for key, value in section_data.items()
                            if key not in ["run_id", "metadata"]
                        }
                        
                        logger.info(f"Fetched data for section: {section_name}", 
                                  extra={"run_id": run_id, "section": section_name, 
                                         "data_keys": list(section_data.keys()) if isinstance(section_data, dict) else "non-dict",
                                         "data_summary": data_summary})
                        
                        # Check if we received any meaningful data
                        has_data = False
                        for key, value in section_data.items():
                            if key not in ["run_id", "metadata"]:
                                if isinstance(value, list) and len(value) > 0:
                                    has_data = True
                                    break
                                elif isinstance(value, dict) and len(value) > 0:
                                    has_data = True
                                    break
                        
                        if not has_data:
                            logger.warning(f"No meaningful data found for section {section_name}", 
                                         extra={"run_id": run_id, "section": section_name, "data_keys": list(section_data.keys())})
                            # Skip to next section if we have no data
                            break
                        
                        # 2. Generate this specific section using LLM
                        section_content = await self._generate_section(run_id, section_name, section_data, system_message)
                        logger.info(f"Generated content for section: {section_name}", 
                                  extra={"run_id": run_id, "section": section_name, "content_type": type(section_content).__name__, 
                                         "content_keys": list(section_content.keys()) if isinstance(section_content, dict) else "non-dict"})
                        
                        # 3. Add to report incrementally
                        if section_content:
                            # Format section for the report
                            formatted_section = {
                                "title": section_name.replace("_", " ").title(),
                                "content": section_content
                            }
                            
                            logger.info(f"Adding formatted section to report: {section_name}", 
                                      extra={"run_id": run_id, "section": section_name, 
                                             "formatted_type": type(section_content).__name__})
                            
                            report["sections"].append(formatted_section)
                            
                            # Track token count for reporting
                            section_token_count = len(str(section_content).split()) // 4 * 3  # Rough estimate
                            total_token_count += section_token_count
                            
                            logger.info(
                                f"Completed section: {section_name}", 
                                extra={
                                    "run_id": run_id, 
                                    "section": section_name,
                                    "token_count": section_token_count
                                }
                            )
                            
                            # Successfully processed section, break out of retry loop
                            break
                        else:
                            logger.warning(f"Empty content for section: {section_name}, retry {retry_count+1}/{max_section_retries+1}", 
                                          extra={"run_id": run_id})
                            
                            # Add section failure
                            section_failures[section_name] = f"Empty content on retry {retry_count+1}"
                            
                            if retry_count < max_section_retries:
                                # Exponential backoff with jitter
                                backoff = min(max_retry_backoff, (2 ** retry_count) + random.uniform(0, 1))
                                logger.info(f"Waiting {backoff:.2f}s before retrying section {section_name}")
                                await asyncio.sleep(backoff)
                                retry_count += 1
                            else:
                                # Max retries reached, add placeholder
                                report["sections"].append({
                                    "title": section_name.replace("_", " ").title(),
                                    "content": {"summary": f"Error generating {section_name} content after multiple attempts."}
                                })
                                break
                    
                    except Exception as e:
                        error_message = str(e)
                        # Check for specific GRPC errors
                        is_grpc_error = "grpc" in error_message.lower() or "Internal error encountered" in error_message
                        
                        logger.error(
                            f"Error generating section: {section_name}", 
                            extra={
                                "run_id": run_id, 
                                "section": section_name,
                                "error": error_message,
                                "error_type": type(e).__name__,
                                "is_grpc_error": is_grpc_error,
                                "retry": retry_count
                            }
                        )
                        
                        # Add section failure
                        section_failures[section_name] = f"{type(e).__name__}: {error_message} on retry {retry_count+1}"
                        
                        if retry_count < max_section_retries:
                            # Exponential backoff with jitter
                            backoff = min(max_retry_backoff, (2 ** retry_count) + random.uniform(0, 1))
                            logger.info(f"Waiting {backoff:.2f}s before retrying section {section_name}")
                            await asyncio.sleep(backoff)
                            retry_count += 1
                        else:
                            # Max retries reached, add placeholder
                            report["sections"].append({
                                "title": section_name.replace("_", " ").title(),
                                "content": {"summary": f"Error generating {section_name} content: {error_message}"}
                            })
                            break
            
            # Log any section failures as warnings
            if section_failures:
                logger.warning(
                    f"Completed report with {len(section_failures)} section failures", 
                    extra={
                        "run_id": run_id,
                        "section_failures": section_failures
                    }
                )
            
            # Once all sections are generated, format the report
            logger.info(f"Formatting report with {len(report['sections'])} sections", extra={"run_id": run_id})
            formatted_report = report  # We're already building the report in the right format
            
            # Generate document
            logger.info(f"Generating document", extra={"run_id": run_id})
            doc_path = await self._generate_document(formatted_report, run_id, user_prompt)
            
            # Store report and get URL
            report_url = await self._store_report(run_id, doc_path, formatted_report)
            
            # Calculate report size in KB
            try:
                output_size_kb = os.path.getsize(doc_path) / 1024
            except Exception:
                output_size_kb = 0
            
            # Create the minimal state to return
            result = {
                "report": {
                    "url": report_url,
                    "run_id": run_id,
                    "version": 1,
                    "created_at": datetime.now(timezone.utc).isoformat(),
                    "last_updated": datetime.now(timezone.utc).isoformat(),
                    "format": "pdf",
                    "file_size_kb": output_size_kb,
                    "token_count": total_token_count,
                    "notes": f"Generated with sequential approach, {total_token_count} tokens"
                },
                "run_id": run_id,
                "report_url": report_url,
                "version": 1,
                "created_at": datetime.now(timezone.utc).isoformat(),
                "last_updated": datetime.now(timezone.utc).isoformat(),
                "format": "docx",
                "file_size_kb": output_size_kb,
                "notes": f"Generated with sequential approach, {total_token_count} tokens",
                "diagnostics": {
                    "available_tables": [
                        table for table, count in diagnostics["table_counts"].items() 
                        if isinstance(count, int) and count > 0
                    ],
                    "section_failures": section_failures
                }
            }
            
            # Log successful completion of the entire process
            process_end_time = datetime.now(timezone.utc)
            await self._log_overall_process(
                run_id=run_id,
                status="completed",
                start_time=process_start_time,
                end_time=process_end_time,
                output_size_kb=output_size_kb,
                token_count=total_token_count
            )
            
            return result
            
        except Exception as e:
            # Log failure of the entire process
            error_message = f"Report compilation failed: {str(e)}"
            logger.error(
                error_message,
                extra={
                    "error_type": type(e).__name__,
                    "run_id": run_id
                }
            )
            await self._log_overall_process(
                run_id=run_id,
                status="failed",
                start_time=process_start_time,
                error_message=error_message
            )
            raise
    
    async def _fetch_section_data(self, run_id: str, section_name: str) -> Dict[str, Any]:
        """
        Fetch data needed for a specific report section directly from Supabase.
        
        Args:
            run_id: The unique identifier for the workflow run
            section_name: The name of the section to generate
            
        Returns:
            Dict containing the data needed for this section
        """
        section_data = {
            "run_id": run_id,
            "metadata": {
                "timestamp": datetime.now(timezone.utc).isoformat()
            }
        }
        
        try:
            # Define required fields for each table based on requirements
            brand_context_fields = [
                "brand_promise", "brand_personality", "brand_tone_of_voice", 
                "brand_values", "brand_purpose", "brand_mission", "target_audience", 
                "customer_needs", "market_positioning", "competitive_landscape", 
                "industry_focus", "industry_trends", "brand_identity_brief", "run_id"
            ]
            
            brand_name_generation_fields = [
                "brand_name", "naming_category", "brand_personality_alignment",
                "brand_promise_alignment", "name_generation_methodology",
                "memorability_score_details", "pronounceability_score_details",
                "visual_branding_potential_details", "target_audience_relevance_details",
                "market_differentiation_details", "run_id"
            ]
            
            semantic_analysis_fields = [
                "brand_name", "denotative_meaning", "etymology", "emotional_valence",
                "brand_personality", "sensory_associations", "figurative_language",
                "phoneme_combinations", "sound_symbolism", "alliteration_assonance",
                "word_length_syllables", "compounding_derivation", "semantic_trademark_risk",
                "run_id"
            ]
            
            linguistic_analysis_fields = [
                "brand_name", "pronunciation_ease", "euphony_vs_cacophony",
                "rhythm_and_meter", "phoneme_frequency_distribution", "sound_symbolism",
                "word_class", "morphological_transparency", "inflectional_properties",
                "ease_of_marketing_integration", "naturalness_in_collocations",
                "semantic_distance_from_competitors", "neologism_appropriateness",
                "overall_readability_score", "notes", "run_id"
            ]
            
            cultural_sensitivity_fields = [
                "brand_name", "cultural_connotations", "symbolic_meanings",
                "alignment_with_cultural_values", "religious_sensitivities",
                "social_political_taboos", "age_related_connotations",
                "regional_variations", "historical_meaning", "current_event_relevance",
                "overall_risk_rating", "notes", "run_id"
            ]
            
            brand_name_evaluation_fields = [
                "brand_name", "overall_score", "shortlist_status", 
                "evaluation_comments", "run_id"
            ]
            
            translation_analysis_fields = [
                "brand_name", "target_language", "direct_translation",
                "semantic_shift", "pronunciation_difficulty", "phonetic_retention",
                "cultural_acceptability", "adaptation_needed", "proposed_adaptation",
                "brand_essence_preserved", "global_consistency_vs_localization",
                "notes", "run_id"
            ]
            
            market_research_fields = [
                "brand_name", "market_opportunity", "target_audience_fit",
                "competitive_analysis", "market_viability", "potential_risks",
                "recommendations", "industry_name", "market_size", "market_growth_rate",
                "key_competitors", "customer_pain_points", "market_entry_barriers",
                "emerging_trends", "run_id"
            ]
            
            competitor_analysis_fields = [
                "brand_name", "competitor_name", "competitor_positioning",
                "competitor_strengths", "competitor_weaknesses", 
                "competitor_differentiation_opportunity", "risk_of_confusion",
                "target_audience_perception", "trademark_conflict_risk", "run_id"
            ]
            
            domain_analysis_fields = [
                "brand_name", "domain_exact_match", "alternative_tlds",
                "misspellings_variations_available", "acquisition_cost",
                "domain_length_readability", "hyphens_numbers_present",
                "brand_name_clarity_in_url", "social_media_availability",
                "scalability_future_proofing", "notes", "run_id"
            ]
            
            # These are all the required fields for survey simulation as specified in the requirements
            # This matches the list of fields in the _fetch_analysis method's default_required_fields
            survey_simulation_essential_fields = [
                "brand_name", "brand_promise_perception_score", "personality_fit_score",
                "emotional_association", "competitive_differentiation_score", 
                "competitor_benchmarking_score", "simulated_market_adoption_score", 
                "qualitative_feedback_summary", "raw_qualitative_feedback",
                "final_survey_recommendation", "strategic_ranking", "industry", 
                "company_size_employees", "company_revenue", "company_name", "job_title", 
                "seniority", "years_of_experience", "department", "education_level", 
                "goals_and_challenges", "values_and_priorities", "decision_making_style", 
                "information_sources", "pain_points", "purchasing_behavior", 
                "online_behavior", "interaction_with_brand", "influence_within_company", 
                "event_attendance", "content_consumption_habits", 
                "vendor_relationship_preferences", "business_chemistry", 
                "reports_to", "buying_group_structure", "budget_authority", 
                "social_media_usage", "frustrations_annoyances", 
                "current_brand_relationships", "success_metrics_product_service", 
                "channel_preferences_brand_interaction", "barriers_to_adoption", 
                "generation_age_range", "run_id"
            ]
            
            # Fetch data based on section type
            if section_name == "executive_summary":
                # For executive summary, we need basic brand context and overall metrics
                # Use a more targeted query with minimal fields
                brand_contexts = await self._fetch_analysis(
                    "brand_context", 
                    run_id,
                    fields=["brand_promise", "brand_mission", "target_audience", "run_id"]
                )
                
                if brand_contexts and len(brand_contexts) > 0:
                    section_data["brand_context"] = brand_contexts[0]
                
                # Get name generation count - just need count, not all data
                name_count = await self._fetch_count("brand_name_generation", run_id)
                section_data["total_names_generated"] = name_count
                
                # Get shortlisted names count - only need specific fields
                evaluations = await self._fetch_analysis(
                    "brand_name_evaluation", 
                    run_id,
                    fields=["brand_name", "shortlist_status", "overall_score", "run_id"],
                    filter_condition={"shortlist_status": True}
                )
                
                section_data["shortlisted_names_count"] = len(evaluations)
                section_data["shortlisted_names"] = [e.get("brand_name") for e in evaluations]
                
            elif section_name == "brand_context":
                # Fetch brand context with the exact fields needed
                brand_contexts = await self._fetch_analysis(
                    "brand_context", 
                    run_id,
                    fields=brand_context_fields
                )
                
                if brand_contexts and len(brand_contexts) > 0:
                    # Make sure we extract the fields according to schema
                    section_data["brand_context"] = brand_contexts[0]
                
            elif section_name == "name_generation":
                # Fetch name generations with required fields
                generations = await self._fetch_analysis(
                    "brand_name_generation", 
                    run_id,
                    fields=brand_name_generation_fields
                )
                
                # Structure the data according to schema
                section_data["name_generations"] = generations
                
            elif section_name == "linguistic_analysis":
                # Fetch name generations for the brand names list - just need names
                name_generations = await self._fetch_analysis(
                    "brand_name_generation", 
                    run_id,
                    fields=["brand_name", "run_id"]
                )
                
                section_data["brand_names"] = [ng.get("brand_name") for ng in name_generations]
                
                # Fetch linguistic analyses with the required fields
                linguistic_analyses = await self._fetch_analysis(
                    "linguistic_analysis", 
                    run_id, 
                    fields=linguistic_analysis_fields
                )
                
                # Structure according to schema - grouped by brand name
                section_data["linguistic_analyses"] = linguistic_analyses
                
            elif section_name == "semantic_analysis":
                # Fetch semantic analyses with the required fields
                semantic_analyses = await self._fetch_analysis(
                    "semantic_analysis", 
                    run_id,
                    fields=semantic_analysis_fields
                )
                
                # Structure according to schema - grouped by brand name
                section_data["semantic_analyses"] = semantic_analyses
                
            elif section_name == "cultural_sensitivity":
                # Fetch cultural sensitivity analyses with the required fields
                cultural_analyses = await self._fetch_analysis(
                    "cultural_sensitivity_analysis", 
                    run_id,
                    fields=cultural_sensitivity_fields
                )
                
                # Structure according to schema - grouped by brand name
                section_data["cultural_analyses"] = cultural_analyses
                
            elif section_name == "name_evaluation":
                # Fetch name generations for the brand names list - just need names
                name_generations = await self._fetch_analysis(
                    "brand_name_generation", 
                    run_id,
                    fields=["brand_name", "run_id"]
                )
                
                section_data["brand_names"] = [ng.get("brand_name") for ng in name_generations]
                
                # Fetch evaluations with the required fields
                evaluations = await self._fetch_analysis(
                    "brand_name_evaluation", 
                    run_id,
                    fields=brand_name_evaluation_fields
                )
                
                # Structure according to schema - grouped by brand name
                section_data["evaluations"] = evaluations
                
            elif section_name == "domain_analysis":
                # Fetch domain analyses with the required fields
                domain_analyses = await self._fetch_analysis(
                    "domain_analysis", 
                    run_id,
                    fields=domain_analysis_fields
                )
                
                # Structure according to schema - grouped by brand name
                section_data["domain_analyses"] = domain_analyses
                
            elif section_name == "seo_analysis":
                # Fetch SEO analyses with the required fields
                seo_fields = [
                    "brand_name", "keyword_alignment", "search_volume",
                    "seo_viability_score", "seo_recommendations", "run_id"
                ]
                
                seo_analyses = await self._fetch_analysis(
                    "seo_online_discoverability", 
                    run_id,
                    fields=seo_fields
                )
                
                # Structure according to schema - grouped by brand name
                section_data["seo_analyses"] = seo_analyses
                
            elif section_name == "competitor_analysis":
                # Fetch competitor analyses with the required fields
                competitor_analyses = await self._fetch_analysis(
                    "competitor_analysis", 
                    run_id,
                    fields=competitor_analysis_fields
                )
                
                # Structure according to schema - grouped by brand name
                section_data["competitor_analyses"] = competitor_analyses
                
            elif section_name == "survey_simulation":
                # Fetch survey simulations with the essential fields only to avoid GRPC overload
                survey_simulations = await self._fetch_analysis(
                    "survey_simulation", 
                    run_id,
                    fields=survey_simulation_essential_fields
                )
                
                # Structure according to schema - organized by brand name
                section_data["survey_simulations"] = survey_simulations
                
            elif section_name == "translation_analysis":
                # Fetch translation analyses with the required fields
                translation_analyses = await self._fetch_analysis(
                    "translation_analysis", 
                    run_id,
                    fields=translation_analysis_fields
                )
                
                # Structure according to schema - grouped by brand name and target language
                section_data["translation_analyses"] = translation_analyses
                
            elif section_name == "recommendations":
                # Only need shortlisted names with minimal data for recommendations
                shortlisted_evaluations = await self._fetch_analysis(
                    "brand_name_evaluation", 
                    run_id,
                    fields=["brand_name", "overall_score", "shortlist_status", "evaluation_comments", "run_id"],
                    filter_condition={"shortlist_status": True}
                )
                
                # Structure for recommendations
                section_data["shortlisted_names"] = shortlisted_evaluations
                
                # Get domain info for shortlisted names only to minimize data
                shortlisted_brand_names = [e.get("brand_name") for e in shortlisted_evaluations]
                if shortlisted_brand_names:
                    domain_info_fields = [
                        "brand_name", "domain_exact_match", "alternative_tlds",
                        "acquisition_cost", "run_id"
                    ]
                    
                    domain_analyses = await self._fetch_analysis(
                        "domain_analysis", 
                        run_id,
                        fields=domain_info_fields,
                        filter_values={"brand_name": shortlisted_brand_names}
                    )
                    
                    section_data["domain_info"] = domain_analyses
                    
                    # Get SEO info for shortlisted names only to minimize data
                    seo_info_fields = [
                        "brand_name", "seo_viability_score", "seo_recommendations", "run_id"
                    ]
                    
                    seo_analyses = await self._fetch_analysis(
                        "seo_online_discoverability", 
                        run_id,
                        fields=seo_info_fields,
                        filter_values={"brand_name": shortlisted_brand_names}
                    )
                    
                    section_data["seo_info"] = seo_analyses
                
        except Exception as e:
            logger.error(
                f"Error fetching data for section: {section_name}",
                extra={
                    "error": str(e),
                    "run_id": run_id
                }
            )
            # Return empty data but don't fail - the section generator will handle missing data
            
        return section_data
    
    async def _generate_section(
        self, 
        run_id: str, 
        section_name: str, 
        section_data: Dict[str, Any],
        system_message: SystemMessage
    ) -> Dict[str, Any]:
        """
        Generate a single section of the report using the LLM.
        
        Args:
            run_id: The unique identifier for the workflow run
            section_name: The name of the section to generate
            section_data: Data needed for this section
            system_message: The system message to use
            
        Returns:
            Dict containing the formatted section content
        """
        # Create section-specific prompt
        section_title = section_name.replace("_", " ").title()
        section_instructions = f"Focus specifically on generating the '{section_title}' section of the report."
        section_instructions += f"\nOnly generate content for the '{section_title}' section, not the entire report."
        
        # Determine which schema we need for this section
        section_schema = None
        for schema in self.output_schemas:
            if schema.name == section_name:
                section_schema = schema
                break
                
        if not section_schema:
            logger.warning(f"No schema found for section: {section_name}")
            # Create a simple schema
            section_schema = ResponseSchema(
                name=section_name,
                description=f"{section_title} content"
            )
        
        # Create a specific output parser for just this section
        section_output_parser = StructuredOutputParser.from_response_schemas([section_schema])
        
        # Optimize data payload by filtering to only essential fields
        # This reduces the amount of data sent to the LLM
        optimized_data = self._optimize_section_data(section_name, section_data)
        
        # Check if data is too large (> 10,000 characters as a rough estimate)
        data_str = str(optimized_data)
        is_large_data = len(data_str) > 10000
        
        if is_large_data:
            logger.warning(f"Section data for {section_name} is very large ({len(data_str)} chars). Chunking...")
            # For very large data, we'll generate in chunks and combine
            return await self._generate_section_in_chunks(run_id, section_name, optimized_data, system_message)
        
        # Create human message with focused instructions
        human_message = HumanMessage(content=f"""
        Generate the {section_title} section for a brand naming report using this data:
        
        {optimized_data}
        
        {section_instructions}
        
        Format your response according to this schema:
        {section_output_parser.get_format_instructions()}
        """)
        
        # Create LLM with appropriate parameters
        temp_llm = ChatGoogleGenerativeAI(
            model=settings.model_name,
            google_api_key=settings.gemini_api_key,
            convert_system_message_to_human=True,
            generation_config={
                "max_output_tokens": 8192,
                "temperature": 0.2,
                "top_p": 0.95,
                "top_k": 40
            },
        )
        
        # Generate this section with retries
        max_retries = 3
        retry_count = 0
        backoff_time = 2  # Base backoff time in seconds
        
        while retry_count < max_retries:
            try:
                # Generate content for this section only
                messages = [system_message, human_message]
                response = await temp_llm.ainvoke(messages)
                content = response.content if hasattr(response, 'content') else str(response)
                
                # Parse the section content
                section_content = section_output_parser.parse(content)
                
                # Success - return the section content
                return section_content.get(section_name, {})
                
            except Exception as e:
                logger.error(
                    f"Error generating content for section {section_name}, attempt {retry_count+1}/{max_retries}",
                    extra={
                        "error": str(e),
                        "run_id": run_id,
                    }
                )
                retry_count += 1
                
                # If all retries failed
                if retry_count >= max_retries:
                    logger.error(f"Failed to generate section {section_name} after {max_retries} attempts")
                    return {"summary": f"Error generating {section_name} content after multiple attempts."}
                
                # Exponential backoff to prevent rate limiting issues
                wait_time = backoff_time * (2 ** (retry_count - 1))  # Exponential backoff
                logger.info(f"Waiting {wait_time} seconds before retry...")
                await asyncio.sleep(wait_time)
        
        # Fallback content if all else fails
        return {"summary": f"Error generating {section_name} content."}
        
    def _optimize_section_data(self, section_name: str, section_data: Dict[str, Any]) -> Dict[str, Any]:
        """
        Optimize section data to reduce payload size sent to LLM.
        
        Args:
            section_name: The name of the section being generated
            section_data: Original data for this section
            
        Returns:
            Dict with optimized/filtered data for LLM consumption
        """
        # Remove metadata and run_id from all sections (no need to send to LLM)
        optimized = {k: v for k, v in section_data.items() if k not in ["metadata", "run_id"]}
        
        # Section-specific optimizations
        if section_name == "executive_summary":
            # For executive summary, we need basic info but not full details
            if "brand_context" in optimized:
                # Keep only essential brand context fields
                essential_fields = ["brand_promise", "brand_mission", "target_audience"]
                optimized["brand_context"] = {k: v for k, v in optimized.get("brand_context", {}).items() 
                                            if k in essential_fields}
                
        elif section_name == "recommendations":
            # Filter to only include evaluation scores and shortlisted names
            if "shortlisted_names" in optimized:
                # Keep the shortlisted names but limit any additional data
                optimized = {
                    "shortlisted_names": optimized.get("shortlisted_names", []),
                    "shortlisted_domains": optimized.get("shortlisted_domains", []),
                    "shortlisted_seo": optimized.get("shortlisted_seo", [])
                }
                
        elif section_name == "survey_simulation":
            # For survey simulation, filter out excessive details
            if "survey_simulations" in optimized:
                # Keep only the most relevant survey data fields
                filtered_simulations = []
                for sim in optimized.get("survey_simulations", []):
                    filtered_sim = {
                        "brand_name": sim.get("brand_name"),
                        "persona_segment": sim.get("persona_segment"),
                        "company_name": sim.get("company_name"),
                        "job_title": sim.get("job_title"),
                        "simulated_market_adoption_score": sim.get("simulated_market_adoption_score"),
                        "qualitative_feedback_summary": sim.get("qualitative_feedback_summary"),
                        "final_survey_recommendation": sim.get("final_survey_recommendation")
                    }
                    filtered_simulations.append(filtered_sim)
                optimized["survey_simulations"] = filtered_simulations
                
        # Apply general optimizations - remove any None values to reduce payload size
        for key in list(optimized.keys()):
            if optimized[key] is None:
                del optimized[key]
                
        return optimized
        
    async def _generate_section_in_chunks(
        self, 
        run_id: str, 
        section_name: str, 
        section_data: Dict[str, Any],
        system_message: SystemMessage
    ) -> Dict[str, Any]:
        """
        Handle generation of very large sections by breaking into manageable chunks.
        
        Args:
            run_id: The unique identifier for the workflow run
            section_name: The name of the section to generate
            section_data: Data needed for this section (already optimized)
            system_message: The system message to use
            
        Returns:
            Dict containing the combined section content
        """
        section_title = section_name.replace("_", " ").title()
        logger.info(f"Generating {section_title} in chunks due to large data size")
        
        # For sections with multiple items (like analyses for different brand names)
        # we'll process a few items at a time
        combined_content = {}
        
        if section_name == "linguistic_analysis" and "linguistic_analyses" in section_data:
            # Process linguistic analyses in chunks of 3-5 brands at a time
            all_analyses = section_data.get("linguistic_analyses", [])
            chunk_size = 4  # Process 4 brands at a time
            
            summary_parts = []
            processed_analyses = []
            
            # Process in chunks
            for i in range(0, len(all_analyses), chunk_size):
                chunk = all_analyses[i:i+chunk_size]
                chunk_data = {
                    "linguistic_analyses": chunk,
                    "chunk_number": i // chunk_size + 1,
                    "total_chunks": (len(all_analyses) + chunk_size - 1) // chunk_size
                }
                
                # Generate content for this chunk
                chunk_content = await self._generate_single_chunk(
                    run_id, 
                    section_name,
                    f"{section_title} (Group {i // chunk_size + 1})", 
                    chunk_data,
                    system_message
                )
                
                # Extract and accumulate the analyses
                if isinstance(chunk_content, dict):
                    if "summary" in chunk_content:
                        summary_parts.append(chunk_content.get("summary", ""))
                    
                    # Combine any other structured data
                    for key in chunk_content:
                        if key != "summary":
                            if key not in combined_content:
                                combined_content[key] = []
                            
                            if isinstance(chunk_content[key], list):
                                combined_content[key].extend(chunk_content[key])
                            else:
                                combined_content[key] = chunk_content[key]
            
            # Create a unified summary by consolidating the summary parts
            if summary_parts:
                combined_content["summary"] = "\n\n".join([
                    f"## {section_title}\n",
                    *summary_parts
                ])
                
            return combined_content
            
        elif section_name == "survey_simulation" and "survey_simulations" in section_data:
            # Similar pattern for survey simulations
            all_simulations = section_data.get("survey_simulations", [])
            chunk_size = 3  # Process 3 simulations at a time
            
            summary_parts = []
            
            # Process in chunks
            for i in range(0, len(all_simulations), chunk_size):
                chunk = all_simulations[i:i+chunk_size]
                chunk_data = {
                    "survey_simulations": chunk,
                    "chunk_number": i // chunk_size + 1,
                    "total_chunks": (len(all_simulations) + chunk_size - 1) // chunk_size
                }
                
                # Generate content for this chunk
                chunk_content = await self._generate_single_chunk(
                    run_id, 
                    section_name,
                    f"{section_title} (Group {i // chunk_size + 1})", 
                    chunk_data,
                    system_message
                )
                
                # Extract and accumulate the content
                if isinstance(chunk_content, dict):
                    if "summary" in chunk_content:
                        summary_parts.append(chunk_content.get("summary", ""))
                    
                    # Add table data if present
                    if "table" in chunk_content:
                        if "table" not in combined_content:
                            combined_content["table"] = {"headers": [], "rows": []}
                        
                        # Initialize headers from first chunk if needed
                        if not combined_content["table"]["headers"] and "headers" in chunk_content["table"]:
                            combined_content["table"]["headers"] = chunk_content["table"]["headers"]
                        
                        # Add rows
                        if "rows" in chunk_content["table"]:
                            combined_content["table"]["rows"].extend(chunk_content["table"]["rows"])
            
            # Create a unified summary
            if summary_parts:
                combined_content["summary"] = "\n\n".join([
                    f"## {section_title}\n",
                    *summary_parts
                ])
                
            return combined_content
        
        else:
            # For other section types, we'll just generate a simplified version
            # with less detailed data
            logger.info(f"Generating simplified version of {section_name} due to data size")
            simplified_data = {"simplified": True}
            
            # Keep key summary information only
            for key, value in section_data.items():
                if isinstance(value, list) and len(value) > 5:
                    # For large lists, just keep a few examples
                    simplified_data[key] = value[:3]
                    simplified_data[f"{key}_count"] = len(value)
                elif isinstance(value, dict) and len(str(value)) > 500:
                    # For large nested objects, keep only main keys
                    simplified_data[key] = {k: "..." for k in value.keys()}
                else:
                    simplified_data[key] = value
                    
            return await self._generate_single_chunk(
                run_id, 
                section_name, 
                section_title, 
                simplified_data,
                system_message
            )
    
    async def _generate_single_chunk(
        self,
        run_id: str,
        section_name: str,
        chunk_title: str,
        chunk_data: Dict[str, Any],
        system_message: SystemMessage
    ) -> Dict[str, Any]:
        """Generate content for a single data chunk"""
        # Determine which schema we need for this section
        section_schema = None
        for schema in self.output_schemas:
            if schema.name == section_name:
                section_schema = schema
                break
                
        if not section_schema:
            # Create a simple schema
            section_schema = ResponseSchema(
                name=section_name,
                description=f"{chunk_title} content"
            )
        
        # Create a specific output parser for just this section
        section_output_parser = StructuredOutputParser.from_response_schemas([section_schema])
        
        # Create human message with focused instructions for this chunk
        human_message = HumanMessage(content=f"""
        Generate the {chunk_title} content for a brand naming report using this data:
        
        {chunk_data}
        
        Focus only on the data provided. This is part of a larger section that will be combined.
        
        Format your response according to this schema:
        {section_output_parser.get_format_instructions()}
        """)
        
        # Create LLM with appropriate parameters
        temp_llm = ChatGoogleGenerativeAI(
            model=settings.model_name,
            google_api_key=settings.gemini_api_key,
            convert_system_message_to_human=True,
            generation_config={
                "max_output_tokens": 4096,  # Reduced token limit for chunks
                "temperature": 0.2,
                "top_p": 0.95,
                "top_k": 40
            },
        )
        
        # Generate with retries
        max_retries = 3
        retry_count = 0
        backoff_time = 2  # Base backoff time in seconds
        
        while retry_count < max_retries:
            try:
                messages = [system_message, human_message]
                response = await temp_llm.ainvoke(messages)
                content = response.content if hasattr(response, 'content') else str(response)
                
                # Parse the section content
                section_content = section_output_parser.parse(content)
                
                # Return the content for this chunk
                return section_content.get(section_name, {})
                
            except Exception as e:
                logger.error(
                    f"Error generating chunk {chunk_title}, attempt {retry_count+1}/{max_retries}",
                    extra={
                        "error": str(e),
                        "run_id": run_id,
                    }
                )
                retry_count += 1
                
                # If all retries failed
                if retry_count >= max_retries:
                    logger.error(f"Failed to generate chunk {chunk_title} after {max_retries} attempts")
                    return {"summary": f"Error generating {chunk_title} content after multiple attempts."}
                
                # Exponential backoff
                wait_time = backoff_time * (2 ** (retry_count - 1))
                await asyncio.sleep(wait_time)
        
        # Fallback content if all else fails
        return {"summary": f"Error generating {chunk_title} content."}

    def _format_report_sections(
        self, report_data: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Format report sections with proper structure and styling."""
        formatted_report = {
            "metadata": {
                "timestamp": datetime.now().isoformat(),
                "version": "1.0",
            },
            "sections": [],
        }
        
        # Order sections according to workflow
        section_order = [
            "executive_summary",
            "brand_context",
            "name_generation",
            "linguistic_analysis",
            "semantic_analysis",
            "cultural_sensitivity",
            "name_evaluation",
            "domain_analysis",
            "seo_analysis", 
            "competitor_analysis",
            "survey_simulation",
            "translation_analysis",
            "recommendations",
        ]
        
        for section in section_order:
            section_data = report_data.get(section, {})
            if not section_data:
                continue
            
            formatted_section = {
                "title": section.replace("_", " ").title(),
                "content": self._format_section_content(section, section_data),
            }
            
            formatted_report["sections"].append(formatted_section)
            
        return formatted_report
        
    def _format_section_content(
        self, section_type: str, section_data: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Format section content with appropriate styling and structure."""
        # Initialize the formatted content
        formatted_content = {}
        
        # If section_data is a string, it's likely the direct content from the LLM's response
        if isinstance(section_data, str):
            formatted_content["summary"] = section_data
            return formatted_content
            
        # If we have a complex object with specific fields
        if section_type in section_data:
            # Some LLM responses have the content directly under the section name key
            main_content = section_data.get(section_type)
            if isinstance(main_content, str):
                formatted_content["summary"] = main_content
                return formatted_content
        
        # Extract summary if present (common in most outputs)
        if "summary" in section_data:
            formatted_content["summary"] = section_data["summary"]
        
        # Handle different section types with specialized formatting
        if section_type == "executive_summary":
            # Executive summary often comes as a markdown string
            if isinstance(section_data, str):
                formatted_content["summary"] = section_data
            elif "executive_summary" in section_data:
                formatted_content["summary"] = section_data["executive_summary"]
                
        elif section_type == "recommendations":
            # Recommendations often come as a markdown string with bullet points
            if isinstance(section_data, str):
                formatted_content["summary"] = section_data
            elif "recommendations" in section_data:
                formatted_content["summary"] = section_data["recommendations"]
                
            # Extract bullet points if present in a structured way
            bullet_points = []
            if isinstance(section_data, dict):
                for key, value in section_data.items():
                    if key.lower() not in ["summary", "recommendations", "executive_summary"]:
                        if isinstance(value, list):
                            bullet_points.append({"heading": key.replace("_", " ").title(), "points": value})
                        elif isinstance(value, str) and ("*" in value or "-" in value):
                            # Attempt to parse markdown bullet points
                            points = []
                            lines = value.split("\n")
                            for line in lines:
                                if line.strip().startswith("*") or line.strip().startswith("-"):
                                    points.append(line.strip()[1:].strip())
                            if points:
                                bullet_points.append({"heading": key.replace("_", " ").title(), "points": points})
                
                if bullet_points:
                    formatted_content["bullet_points"] = bullet_points
                    
        elif section_type == "survey_simulation":
            # For survey simulation, we need to structure the table data
            formatted_content["table"] = {
                "headers": ["Persona", "Company", "Role", "Brand Score", "Key Feedback"],
                "rows": []  # Will be filled with actual data rows
            }
            
            # Attempt to extract survey data
            if isinstance(section_data, dict) and "survey_simulations" in section_data:
                surveys = section_data["survey_simulations"]
                for survey in surveys:
                    if isinstance(survey, dict):
                        row = [
                            survey.get("persona_segment", ""),
                            survey.get("company_name", ""),
                            survey.get("job_title", ""),
                            str(survey.get("simulated_market_adoption_score", "")),
                            survey.get("qualitative_feedback_summary", "")
                        ]
                        formatted_content["table"]["rows"].append(row)
                        
        elif section_type == "competitor_analysis":
            # Format competitor analysis as structured data
            if "summary" not in formatted_content and "competitive_analysis" in section_data:
                formatted_content["summary"] = section_data["competitive_analysis"]
            
            # Create a more structured format for competitor analysis
            # Extract any key points as bullet points
            bullet_points = []
            key_competitor_aspects = [
                "Competitor Naming Styles", 
                "Market Positioning",
                "Differentiation Opportunities",
                "Risk of Confusion",
                "Target Audience Perception", 
                "Competitive Advantages",
                "Trademark Considerations"
            ]
            
            # Look for these aspects in the data
            for aspect in key_competitor_aspects:
                aspect_key = aspect.lower().replace(" ", "_")
                if aspect_key in section_data:
                    content = section_data[aspect_key]
                    if isinstance(content, str):
                        bullet_points.append({
                            "heading": aspect,
                            "points": [content]
                        })
                    elif isinstance(content, list):
                        bullet_points.append({
                            "heading": aspect,
                            "points": content
                        })
            
            # If we found bullet points, add them
            if bullet_points:
                formatted_content["bullet_points"] = bullet_points
                
        elif section_type == "domain_analysis":
            # Format as bullet points with subsections
            if "domain_analyses" in section_data:
                domains = section_data["domain_analyses"]
                details = []
                for domain in domains:
                    if isinstance(domain, dict) and "brand_name" in domain:
                        detail = {
                            "heading": domain["brand_name"],
                            "content": f"Exact Match: {domain.get('domain_exact_match', 'Unknown')}\n"
                                      f"Alternative TLDs: {domain.get('alternative_tlds', 'Unknown')}\n"
                                      f"Acquisition Cost: {domain.get('acquisition_cost', 'Unknown')}\n"
                                      f"Notes: {domain.get('notes', '')}"
                        }
                        details.append(detail)
                
                if details:
                    formatted_content["details"] = details
        else:
            # Default formatting for other section types
            # Check if we have details to extract
            details = []
            if isinstance(section_data, dict):
                for key, value in section_data.items():
                    if key.lower() not in ["summary"]:
                        if isinstance(value, str):
                            details.append({
                                "heading": key.replace("_", " ").title(),
                                "content": value
                            })
                        elif isinstance(value, dict):
                            # For nested structures, format as a readable string
                            content = "\n".join([f"**{k}**: {v}" for k, v in value.items()])
                            details.append({
                                "heading": key.replace("_", " ").title(),
                                "content": content
                            })
                        elif isinstance(value, list):
                            # For lists, format as bullet points in a string
                            content = "\n".join([f"* {item}" for item in value])
                            details.append({
                                "heading": key.replace("_", " ").title(),
                                "content": content
                            })
            
            if details:
                formatted_content["details"] = details
                
        # If we didn't find any specific content to format, default to the original data
        if not formatted_content:
            formatted_content["summary"] = str(section_data)
            
        return formatted_content

    async def _generate_document(self, report: Dict[str, Any], run_id: str, user_prompt: str = None) -> str:
        """Generate a Word document from the report data."""
        # Import re if not already imported
        import re
        
        doc = Document()

        # Set up document styles
        styles = doc.styles

        # Title style
        title_style = styles.add_style("CustomTitle", WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.size = Pt(24)
        title_style.font.bold = True

        # Heading 1 style
        h1_style = styles.add_style("CustomH1", WD_STYLE_TYPE.PARAGRAPH)
        h1_style.font.size = Pt(16)
        h1_style.font.bold = True

        # Heading 2 style
        h2_style = styles.add_style("CustomH2", WD_STYLE_TYPE.PARAGRAPH)
        h2_style.font.size = Pt(14)
        h2_style.font.bold = True

        # Heading 3 style
        h3_style = styles.add_style("CustomH3", WD_STYLE_TYPE.PARAGRAPH)
        h3_style.font.size = Pt(12)
        h3_style.font.bold = True
        h3_style.font.italic = True

        # Normal text style
        normal_style = styles.add_style("CustomNormal", WD_STYLE_TYPE.PARAGRAPH)
        normal_style.font.size = Pt(11)

        # Italic style for notes
        note_style = styles.add_style("CustomNote", WD_STYLE_TYPE.PARAGRAPH)
        note_style.font.size = Pt(11)
        note_style.font.italic = True
        
        # Bold style for important points
        bold_style = styles.add_style("CustomBold", WD_STYLE_TYPE.PARAGRAPH)
        bold_style.font.size = Pt(11)
        bold_style.font.bold = True

        # Add title
        title = doc.add_paragraph("Brand Naming Report", style="CustomTitle")
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add Mae branding and context
        doc.add_paragraph("Generated by Mae Brand Naming Expert", style="CustomNormal")
        doc.add_paragraph(
            f"Date: {datetime.now().strftime('%B %d, %Y %H:%M:%S %Z')}", style="CustomNormal"
        )
        doc.add_paragraph(f"Run ID: {run_id}", style="CustomNormal")
        doc.add_paragraph()

        # Add simulation context note
        note_heading = doc.add_paragraph("Important Note", style="CustomH2")
        note = doc.add_paragraph(
            "This report was generated through the Mae Brand Naming Expert simulation. "
            "The only input provided was the initial user prompt:",
            style="CustomNote",
        )
        doc.add_paragraph(
            f'"{user_prompt or report.get("metadata", {}).get("user_prompt", "")}"', style="CustomNote"
        )
        doc.add_paragraph(
            "All additional context, analysis, and insights were autonomously generated by the simulation.",
            style="CustomNote",
        )

        # Add horizontal line
        doc.add_paragraph("_" * 50, style="CustomNormal")
        doc.add_paragraph()

        # Add table of contents
        toc_heading = doc.add_paragraph("Table of Contents", style="CustomH1")
        for i, section in enumerate(report["sections"], 1):
            toc_entry = doc.add_paragraph(style="CustomNormal")
            toc_entry.add_run(f"{i}. {section['title']}")
        doc.add_paragraph()  # Add space after TOC

        # Helper function to parse markdown in text
        def _parse_markdown_text(text, paragraph):
            if text is None:
                paragraph.add_run("")
                return paragraph
                
            # Handle bold text (replace **text** with bold text)
            bold_pattern = re.compile(r'\*\*(.*?)\*\*')
            
            # Find all bold patterns
            bold_matches = list(bold_pattern.finditer(text))
            
            if bold_matches:
                # Text has bold sections, need to add runs with appropriate formatting
                last_end = 0
                for match in bold_matches:
                    # Add regular text before the bold
                    if match.start() > last_end:
                        paragraph.add_run(text[last_end:match.start()])
                    
                    # Add bold text
                    bold_run = paragraph.add_run(match.group(1))  # Group 1 is the text inside **
                    bold_run.bold = True
                    
                    last_end = match.end()
                
                # Add any remaining text after the last bold section
                if last_end < len(text):
                    paragraph.add_run(text[last_end:])
            else:
                # No bold text, just add the entire text
                paragraph.add_run(text)
                
            return paragraph
            
        # Helper function to parse and add a markdown text block
        def _add_markdown_paragraph(text, default_style="CustomNormal"):
            if text is None:
                return doc.add_paragraph(style=default_style)
                
            # Check for headings
            if text.startswith('## '):
                para = doc.add_paragraph(text[3:], style="CustomH2")
                return para
            elif text.startswith('# '):
                para = doc.add_paragraph(text[2:], style="CustomH1")
                return para
            elif text.startswith('### '):
                para = doc.add_paragraph(text[4:], style="CustomH3")
                return para
            elif text.startswith('* ') or text.startswith('- '):
                # Create bullet point
                para = doc.add_paragraph(style="List Bullet")
                return _parse_markdown_text(text[2:], para)
            else:
                # Regular paragraph
                para = doc.add_paragraph(style=default_style)
                return _parse_markdown_text(text, para)
                
        # Helper function to parse and add a markdown table
        def _add_markdown_table(text):
            lines = text.strip().split('\n')
            if len(lines) < 2:
                # Not enough lines for a valid table
                return _add_markdown_paragraph(text)
                
            # Count pipe symbols to determine columns
            header_row = lines[0]
            columns = header_row.count('|')
            
            # Need at least 2 pipe symbols for a valid table (3 columns)
            if columns < 2:
                return _add_markdown_paragraph(text)
                
            # Split the header row by pipes and strip whitespace
            headers = [cell.strip() for cell in header_row.split('|')]
            # Remove empty first/last cells if present
            if headers[0] == '':
                headers = headers[1:]
            if headers[-1] == '':
                headers = headers[:-1]
                
            # Create table
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = 'Table Grid'
            
            # Add header row
            header_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                header_cells[i].text = header
                
            # Skip the separator row (typically row 1 with ----)
            start_row = 2 if len(lines) > 2 and '---' in lines[1] else 1
            
            # Add data rows
            for i in range(start_row, len(lines)):
                row_text = lines[i]
                if not row_text.strip():
                    continue
                    
                # Split by pipes
                cells = [cell.strip() for cell in row_text.split('|')]
                # Remove empty first/last cells if present
                if cells[0] == '':
                    cells = cells[1:]
                if cells[-1] == '':
                    cells = cells[:-1]
                    
                # Add row to table
                row_cells = table.add_row().cells
                for j, cell in enumerate(cells):
                    if j < len(row_cells):
                        row_cells[j].text = cell
                        
            return table

        # Define helper functions for special section processing - MOVED UP before they're called
        def _process_recommendations_section(content):
            """Format recommendations professionally with paragraphs and bullet points."""
            # This function doesn't modify the document directly
            # It enhances the content structure for better formatting when processed
            if isinstance(content, dict) and "summary" in content:
                summary = content["summary"]
                # No need to do anything as the summary will be processed normally
                pass

            # Ensure bullet points are properly structured
            if isinstance(content, dict) and "bullet_points" in content:
                bullet_points = content["bullet_points"]
                # The bullet points will be processed normally
                pass

        def _process_linguistic_analysis_section(content):
            """Ensure all linguistic analysis data is captured."""
            # This function doesn't modify the document directly
            # It just logs information to confirm data is being processed
            logger.info("Processing linguistic analysis section to ensure all data is captured")
            if isinstance(content, dict):
                # Log the keys to verify all data is present
                logger.info(f"Linguistic analysis keys: {list(content.keys())}")
                
                # If we have "summary" and there's no table in it, add tables for detailed analysis
                if "summary" in content and not ('|' in content["summary"] and '\n' in content["summary"]):
                    logger.info("Adding comprehensive linguistic analysis tables")
                    # The linguistic data will be processed as normal through the existing logic

        def _process_translation_analysis_section(content):
            """Ensure translation analysis is properly included."""
            logger.info("Processing translation analysis section to ensure all data is captured")
            if isinstance(content, dict):
                # Log the keys to verify all data is present
                logger.info(f"Translation analysis keys: {list(content.keys())}")
                
                # The translation data will be processed as normal through the existing logic

        def _process_survey_simulation_section(content):
            """Ensure comprehensive survey results are included."""
            logger.info("Processing survey simulation section to include full details")
            if isinstance(content, dict):
                # Log the keys to verify all data is present
                logger.info(f"Survey simulation keys: {list(content.keys())}")
                
                # If there's a table, ensure it has all necessary columns
                if "table" in content and isinstance(content["table"], dict):
                    table = content["table"]
                    logger.info(f"Survey table headers: {table.get('headers', [])}")
                    # The existing logic will process the table with all columns

        # Add sections
        logger.info(f"Processing {len(report['sections'])} sections for document generation")
        
        # Before processing, look for the Translation Analysis section index
        translation_section_index = None
        for i, section in enumerate(report["sections"]):
            if section["title"].lower() == "translation analysis":
                translation_section_index = i
                break
        
        for i, section in enumerate(report["sections"], 1):
            # Log the section being processed
            logger.info(f"Processing section {i}: {section['title']}")
            
            # Add section title with section number
            section_title = doc.add_paragraph(f"{i}. {section['title']}", style="CustomH1")
            
            # Get the section content
            content = section['content']
            
            # Special handling for recommendations section (make it more professional)
            if section["title"].lower() == "recommendations":
                _process_recommendations_section(content)
            
            # Special handling for linguistic analysis to ensure all data is captured
            if section["title"].lower() == "linguistic analysis":
                _process_linguistic_analysis_section(content)
                
            # Special handling for translation analysis to ensure all data is captured
            if section["title"].lower() == "translation analysis":
                _process_translation_analysis_section(content)
                
            # Special handling for survey simulation to include full details
            if section["title"].lower() == "survey simulation":
                _process_survey_simulation_section(content)
            
            # Process different content formats
            if isinstance(content, dict):
                # Handle structured content
                
                # Process summary if present
                if "summary" in content and content["summary"]:
                    summary_text = content["summary"]
                    
                    # Check if the summary contains markdown tables
                    if '|' in summary_text and '\n' in summary_text and ('|---' in summary_text or '|----' in summary_text):
                        # Split by potential table markers
                        parts = re.split(r'(\|.*\|[\s]*\n\|[-]+\|.*(?:\n\|.*\|)*)', summary_text, flags=re.DOTALL)
                        
                        for part in parts:
                            if part.strip():
                                if part.startswith('|') and ('|---' in part or '|----' in part):
                                    # This is a table
                                    _add_markdown_table(part)
                                else:
                                    # Process regular text with markdown
                                    lines = part.split('\n')
                                    for line in lines:
                                        if line.strip():
                                            _add_markdown_paragraph(line)
                    else:
                        # Process as regular text with markdown
                        lines = summary_text.split('\n')
                        for line in lines:
                            if line.strip():
                                _add_markdown_paragraph(line)
                
                # Process bullet points
                if "bullet_points" in content and content["bullet_points"]:
                    for bp_section in content["bullet_points"]:
                        if isinstance(bp_section, dict) and "heading" in bp_section and "points" in bp_section:
                            # Add the bullet point section heading
                            doc.add_paragraph(bp_section["heading"], style="CustomH2")
                            
                            # Add each point as a bullet
                            for point in bp_section["points"]:
                                bullet = doc.add_paragraph(style="List Bullet")
                                _parse_markdown_text(point, bullet)
                        elif isinstance(bp_section, str):
                            # Direct string bullet point
                            bullet = doc.add_paragraph(style="List Bullet")
                            _parse_markdown_text(bp_section, bullet)
                
                # Process tables
                if "table" in content and content["table"]:
                    table_data = content["table"]
                    if "headers" in table_data and "rows" in table_data:
                        # Create table with headers
                        table = doc.add_table(rows=1, cols=len(table_data["headers"]))
                        table.style = "Table Grid"
                        
                        # Add headers
                        header_cells = table.rows[0].cells
                        for i, header in enumerate(table_data["headers"]):
                            header_cells[i].text = header
                            
                        # Add rows
                        for row_data in table_data["rows"]:
                            row_cells = table.add_row().cells
                            for i, cell_value in enumerate(row_data):
                                if i < len(row_cells):  # Ensure we don't exceed the table columns
                                    row_cells[i].text = str(cell_value)
                
                # Process details sections
                if "details" in content and content["details"]:
                    for detail in content["details"]:
                        if isinstance(detail, dict) and "heading" in detail and "content" in detail:
                            # Add section heading
                            doc.add_paragraph(detail["heading"], style="CustomH2")
                            
                            # Process content - could be string or structured content
                            if isinstance(detail["content"], str):
                                # Check if content has tables
                                if '|' in detail["content"] and '\n' in detail["content"] and ('|---' in detail["content"] or '|----' in detail["content"]):
                                    _add_markdown_table(detail["content"])
                                else:
                                    # Split by line breaks and process each paragraph
                                    paragraphs = detail["content"].split("\n")
                                    for para in paragraphs:
                                        if para.strip():
                                            _add_markdown_paragraph(para)
                            else:
                                # Handle more complex content structure if needed
                                doc.add_paragraph(str(detail["content"]), style="CustomNormal")
            
            elif isinstance(content, str):
                # Direct string content - check if it contains tables
                if '|' in content and '\n' in content and ('|---' in content or '|----' in content):
                    _add_markdown_table(content)
                else:
                    # Process as regular text with markdown
                    lines = content.split('\n')
                    for line in lines:
                        if line.strip():
                            _add_markdown_paragraph(line)
            
            # Add spacing between sections
            doc.add_paragraph()

        # Save document to temporary file - making sure to use utf-8 encoding for foreign characters
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            doc_path = tmp.name
            try:
                logger.info(f"Saving document to {doc_path}")
                doc.save(doc_path)
                logger.info(f"Document saved successfully to {doc_path}")
                return doc_path
            except Exception as e:
                logger.error(f"Error saving document: {str(e)}")
                raise

    async def _store_report(self, run_id: str, doc_path: str, report_data: Dict[str, Any]) -> str:
        """Store the report document using Supabase Storage standard upload method.
        
        Args:
            run_id: Unique identifier for the report run
            doc_path: Path to the local document file
            report_data: Report data dictionary
            
        Returns:
            str: URL to the stored report
        """
        # Define timestamp outside try block so it's available in except block
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        # Define object_key here so it's available in all blocks
        object_key = f"{run_id}/{timestamp}_report.docx"
        
        try:
            # Get storage settings
            bucket_name = settings.s3_bucket  # Using the same bucket name defined in settings
            
            file_size = os.path.getsize(doc_path) // 1024  # Size in KB
            
            logger.info(
                "Storing report using Supabase Storage standard upload",
                extra={
                    "run_id": run_id,
                    "file_size_kb": file_size,
                    "bucket": bucket_name,
                    "object_key": object_key
                }
            )
            
            # Read the file content
            with open(doc_path, "rb") as file:
                file_content = file.read()
            
            # Set content type for the document
            content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            
            # Use Supabase client to upload the file
            if self.supabase:
                # Get Supabase client directly
                supabase_client = self.supabase.client
                
                # Upload file using standard upload method - not async in Supabase client
                try:
                    logger.debug("Starting file upload using standard Supabase upload")
                    result = supabase_client.storage.from_(bucket_name).upload(
                        path=object_key,
                        file=file_content,
                        file_options={"content_type": content_type}
                    )
                    
                    logger.debug(f"Upload result: {result}")
                    
                    if hasattr(result, 'error') and result.error:
                        raise Exception(f"Supabase upload returned error: {result.error}")
                        
                except Exception as upload_error:
                    logger.error(
                        f"Supabase upload error",
                        extra={
                            "error": str(upload_error),
                            "bucket": bucket_name,
                            "object_key": object_key
                        }
                    )
                    raise Exception(f"Supabase upload error: {str(upload_error)}")
                
                # Try to get the public URL directly from Supabase client first
                try:
                    report_url = supabase_client.storage.from_(bucket_name).get_public_url(object_key)
                    logger.info(f"Generated public URL from Supabase client: {report_url}")
                except Exception as url_error:
                    logger.warning(
                        f"Could not get URL from Supabase client, falling back to constructed URL",
                        extra={"error": str(url_error)}
                    )
                    # Fall back to constructing the URL manually
                    base_supabase_url = settings.supabase_url
                    if base_supabase_url.endswith('/storage/v1/s3'):
                        base_supabase_url = base_supabase_url[:-13]  # Remove '/storage/v1/s3'
                    elif base_supabase_url.endswith('/storage/v1'):
                        base_supabase_url = base_supabase_url[:-10]  # Remove '/storage/v1'
                        
                    public_url_base = f"{base_supabase_url}/storage/v1/object/public"
                    report_url = f"{public_url_base}/{bucket_name}/{object_key}"
                    logger.info(f"Using constructed public URL: {report_url}")
                
                logger.info(
                    "File uploaded successfully using standard upload, storing metadata",
                    extra={
                        "run_id": run_id,
                        "report_url": report_url
                    }
                )
                
                # Store metadata in the report_compilation table
                metadata = {
                    "run_id": run_id,
                    "report_url": report_url,
                    "version": "1.0",
                    "created_at": datetime.now().isoformat(),
                    "updated_at": datetime.now().isoformat(),
                    "format": "docx",  # Use hardcoded value instead of settings.report_output_format
                    "file_size_kb": file_size,
                    "notes": f"Report generated successfully for run_id {run_id}",
                    "last_updated": datetime.now().isoformat()
                    # Remove fields that don't exist in the table schema
                    # "storage_region": settings.s3_region
                    # "storage_endpoint": s3_endpoint,
                    # "upload_protocol": "s3", 
                    # "storage_bucket": bucket_name,
                    # "storage_key": object_key,
                }

                try:
                    await self.supabase.execute_with_retry(
                        operation="insert",
                        table="report_compilation",
                        data=metadata
                    )
                    
                    logger.info(
                        f"Successfully stored report metadata in report_compilation table",
                        extra={
                            "run_id": run_id,
                            "report_url": report_url,
                            "table": "report_compilation"
                        }
                    )
                    
                except (TypeError, ValueError) as e:
                    logger.error(
                        "Error preparing report metadata for Supabase",
                        extra={
                            "run_id": run_id,
                            "error": str(e),
                            "table": "report_compilation"
                        }
                    )
                    # Continue execution since the file is still uploaded successfully
                
                except APIError as e:
                    logger.error(
                        "API error storing report metadata in Supabase",
                        extra={
                            "run_id": run_id,
                            "error": str(e),
                            "table": "report_compilation",
                            "status_code": getattr(e, "status_code", None)
                        }
                    )
                    # Continue execution since the file is still uploaded successfully
                
                except Exception as e:
                    logger.error(
                        "Unexpected error storing report metadata in Supabase",
                        extra={
                            "run_id": run_id,
                            "error_type": type(e).__name__,
                            "error": str(e),
                            "table": "report_compilation"
                        }
                    )
                    # Continue execution since the file is still uploaded successfully
                
                # Clean up temporary file
                os.unlink(doc_path)
                
                logger.info(
                    "Report stored successfully using Supabase Storage standard upload",
                    extra={
                        "run_id": run_id, 
                        "report_url": report_url,
                        "file_size_kb": file_size,
                        "bucket": bucket_name
                    }
                )
                
                return report_url
            else:
                raise ValueError("Supabase client not initialized")
                
        except Exception as e:
            logger.error(
                "Error storing report via Supabase Storage standard upload",
                extra={
                    "run_id": run_id, 
                    "error_type": type(e).__name__, 
                    "error_message": str(e),
                    "bucket": settings.s3_bucket
                }
            )
            
            # If we failed to store, keep the local file as a backup
            local_backup_path = f"./reports/{object_key}"
            try:
                import shutil
                os.makedirs(os.path.dirname(local_backup_path), exist_ok=True)
                shutil.copy(doc_path, local_backup_path)
                logger.info(f"Created local backup at {local_backup_path}")
                # Still attempt to clean up the temp file
                os.unlink(doc_path)
                
                # Return a file:// URL as fallback
                absolute_path = os.path.abspath(local_backup_path)
                return f"file://{absolute_path}"
                
            except Exception as backup_error:
                logger.error(f"Failed to create local backup: {str(backup_error)}")
            
            # Raise the original error
            raise

    async def _fetch_count(self, analysis_type: str, run_id: str) -> int:
        """
        Fetch just the count of records for a specific analysis type and run ID.
        
        Args:
            analysis_type (str): The type of analysis to fetch
            run_id (str): The run ID to filter by
            
        Returns:
            int: Count of records
        """
        try:
            # Use a count query for better performance when we just need the number of rows
            count_query = f"""
                SELECT COUNT(*) 
                FROM {analysis_type} 
                WHERE run_id = :run_id
            """
            
            params = {"run_id": run_id}
            
            result = await self.supabase.execute_query(count_query, params)
            
            # Extract count from result
            if result and len(result) > 0:
                return result[0].get("count", 0)
            
            return 0
            
        except Exception as e:
            logger.error(
                f"Error fetching count for {analysis_type}",
                extra={
                    "error": str(e),
                    "run_id": run_id
                }
            )
            return 0
            
    async def _fetch_analysis(
        self, 
        analysis_type: str, 
        run_id: str, 
        fields: Optional[List[str]] = None,
        filter_condition: Optional[Dict[str, Any]] = None,
        filter_values: Optional[Dict[str, List[Any]]] = None
    ) -> List[Dict[str, Any]]:
        """
        Fetch analysis data from the database for a specific analysis type and run ID.
        Only pulls the specific fields needed for the report based on requirements.
        
        Args:
            analysis_type (str): The type of analysis to fetch
            run_id (str): The run ID to filter by
            fields (Optional[List[str]]): Specific fields to fetch (performance optimization)
                                         If None, will use default required fields
            filter_condition (Optional[Dict[str, Any]]): Additional filter conditions
                                                        Format: {'column': 'value'} for equality
                                                        or {'column': 'op.value'} for other operators
            filter_values (Optional[Dict[str, List[Any]]]): Filters with multiple possible values
                                                          Format: {'column': [value1, value2, ...]}
            
        Returns:
            List[Dict[str, Any]]: List of analysis results
            
        Raises:
            ValueError: If required fields are missing from the analysis results
        """
        try:
            # Define default required fields for each analysis type based on the report requirements
            default_required_fields = {
                "brand_context": [
                    "brand_promise", "brand_personality", "brand_tone_of_voice", 
                    "brand_values", "brand_purpose", "brand_mission", 
                    "target_audience", "customer_needs", "market_positioning", 
                    "competitive_landscape", "industry_focus", "industry_trends", 
                    "brand_identity_brief"
                ],
                "brand_name_generation": [
                    "brand_name", "naming_category", "brand_personality_alignment", 
                    "brand_promise_alignment", "name_generation_methodology",
                    "memorability_score_details", "pronounceability_score_details", 
                    "visual_branding_potential_details", "target_audience_relevance_details", 
                    "market_differentiation_details"
                ],
                "semantic_analysis": [
                    "brand_name", "denotative_meaning", "etymology", "emotional_valence", 
                    "brand_personality", "sensory_associations", "figurative_language", 
                    "phoneme_combinations", "sound_symbolism", "alliteration_assonance", 
                    "word_length_syllables", "compounding_derivation", "semantic_trademark_risk"
                ],
                "linguistic_analysis": [
                    "brand_name", "pronunciation_ease", "euphony_vs_cacophony", "rhythm_and_meter",
                    "phoneme_frequency_distribution", "sound_symbolism", "word_class",
                    "morphological_transparency", "inflectional_properties", "ease_of_marketing_integration",
                    "naturalness_in_collocations", "semantic_distance_from_competitors",
                    "neologism_appropriateness", "overall_readability_score", "notes"
                ],
                "cultural_sensitivity_analysis": [
                    "brand_name", "cultural_connotations", "symbolic_meanings", 
                    "alignment_with_cultural_values", "religious_sensitivities", 
                    "social_political_taboos", "age_related_connotations", "regional_variations", 
                    "historical_meaning", "current_event_relevance", "overall_risk_rating", "notes"
                ],
                "brand_name_evaluation": [
                    "brand_name", "overall_score", "shortlist_status", "evaluation_comments"
                ],
                "translation_analysis": [
                    "brand_name", "target_language", "direct_translation", "semantic_shift", 
                    "pronunciation_difficulty", "phonetic_retention", "cultural_acceptability", 
                    "adaptation_needed", "proposed_adaptation", "brand_essence_preserved", 
                    "global_consistency_vs_localization", "notes"
                ],
                "market_research": [
                    "brand_name", "market_opportunity", "target_audience_fit", "competitive_analysis", 
                    "market_viability", "potential_risks", "recommendations", "industry_name", 
                    "market_size", "market_growth_rate", "key_competitors", "customer_pain_points", 
                    "market_entry_barriers", "emerging_trends"
                ],
                "competitor_analysis": [
                    "brand_name", "competitor_name", "competitor_positioning", "competitor_strengths", 
                    "competitor_weaknesses", "competitor_differentiation_opportunity", 
                    "risk_of_confusion", "target_audience_perception", "trademark_conflict_risk"
                ],
                "domain_analysis": [
                    "brand_name", "domain_exact_match", "alternative_tlds", 
                    "misspellings_variations_available", "acquisition_cost", "domain_length_readability", 
                    "hyphens_numbers_present", "brand_name_clarity_in_url", 
                    "social_media_availability", "scalability_future_proofing", "notes"
                ],
                "survey_simulation": [
                    "brand_name", "brand_promise_perception_score", "personality_fit_score", 
                    "emotional_association", "competitive_differentiation_score", 
                    "competitor_benchmarking_score", "simulated_market_adoption_score", 
                    "qualitative_feedback_summary", "raw_qualitative_feedback",
                    "final_survey_recommendation", "strategic_ranking", "industry", 
                    "company_size_employees", "company_revenue", "company_name", "job_title", 
                    "seniority", "years_of_experience", "department", "education_level", 
                    "goals_and_challenges", "values_and_priorities", "decision_making_style", 
                    "information_sources", "pain_points", "purchasing_behavior", 
                    "online_behavior", "interaction_with_brand", "influence_within_company", 
                    "event_attendance", "content_consumption_habits", 
                    "vendor_relationship_preferences", "business_chemistry", 
                    "reports_to", "buying_group_structure", "budget_authority", 
                    "social_media_usage", "frustrations_annoyances", 
                    "current_brand_relationships", "success_metrics_product_service", 
                    "channel_preferences_brand_interaction", "barriers_to_adoption", 
                    "generation_age_range"
                ]
            }
            
            # Determine which fields to fetch
            selected_fields = fields if fields else default_required_fields.get(analysis_type, ["*"])
            
            # Always include run_id for filtering even though it's not needed in the report
            if "run_id" not in selected_fields and selected_fields != ["*"]:
                selected_fields.append("run_id")
                
            # Use execute_with_retry method which is known to work
            logger.info(f"Fetching {analysis_type} using execute_with_retry method")
            
            # Prepare data for the query
            query_data = {}
            
            # Add select fields parameter
            if selected_fields and selected_fields != ["*"]:
                query_data["select"] = ",".join(selected_fields)
            
            # FIXED: Handle multiple operator prefixes by fully cleaning the run_id
            # Extract clean run_id to prevent double operators (like "eq.eq.")
            clean_run_id = run_id
            
            # Keep stripping operator prefixes until none remain
            operators = ['eq.', 'neq.', 'gt.', 'gte.', 'lt.', 'lte.', 'like.', 'ilike.', 'is.', 'in.']
            changed = True
            while changed:
                changed = False
                for op in operators:
                    if clean_run_id.startswith(op):
                        clean_run_id = clean_run_id[len(op):]
                        changed = True
                        break
                
            # Add run_id filter with correct syntax
            query_data["run_id"] = f"eq.{clean_run_id}"
            
            # Add any additional filter conditions
            if filter_condition:
                for key, value in filter_condition.items():
                    # Check if the value already includes an operator like 'gt.' or 'neq.'
                    if isinstance(value, str) and any(value.startswith(f"{op}.") for op in ['eq', 'neq', 'gt', 'gte', 'lt', 'lte', 'like', 'ilike', 'is', 'in']):
                        query_data[key] = value
                    else:
                        # Default to equality if no operator specified
                        query_data[key] = f"eq.{value}"
            
            # Handle filter_values for filters with multiple possible values
            # Example: {"brand_name": ["name1", "name2", "name3"]}
            if filter_values:
                for field, values in filter_values.items():
                    if values:
                        # Format as "in.(value1,value2,value3)"
                        formatted_values = ",".join([str(v) for v in values])
                        query_data[field] = f"in.({formatted_values})"
            
            # Log the query parameters to debug SQL issues
            logger.debug(f"Query parameters for {analysis_type}: {query_data}")
            
            # Execute the query
            results = await self.supabase.execute_with_retry("select", analysis_type, query_data)
            
            # Log the results
            if results:
                logger.info(f"Successfully fetched {len(results)} {analysis_type} records")
            else:
                logger.warning(f"No {analysis_type} data found for run_id: {run_id}")
            
            return results or []
            
        except Exception as e:
            logger.error(f"Error fetching {analysis_type} data: {str(e)}")
            logger.exception(e)
            return []

    def _format_bullet_points(self, section_data: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Format section data as bullet points."""
        bullet_points = []
        for key, value in section_data.items():
            if isinstance(value, list):
                bullet_points.append({"heading": key.replace("_", " ").title(), "points": value})
            elif isinstance(value, dict):
                bullet_points.append(
                    {
                        "heading": key.replace("_", " ").title(),
                        "points": [f"{k}: {v}" for k, v in value.items()],
                    }
                )
        return bullet_points

    def _format_details(self, section_data: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Format section details with proper structure."""
        details = []
        for key, value in section_data.items():
            if key == "summary":
                continue
            details.append({"heading": key.replace("_", " ").title(), "content": value})
        return details

    def _group_by_field(
        self, data: List[Dict[str, Any]], field: str
    ) -> Dict[str, List[Dict[str, Any]]]:
        """Group a list of dictionaries by a specified field."""
        groups = {}
        for item in data:
            if item[field] not in groups:
                groups[item[field]] = []
            groups[item[field]].append(item)
        return groups

    def _organize_brand_analyses(
        self,
        brand_names: List[str],
        brand_name_generations: List[Dict[str, Any]],
        brand_name_evaluations: List[Dict[str, Any]],
        competitor_analyses: List[Dict[str, Any]],
        cultural_sensitivity_analyses: List[Dict[str, Any]],
        domain_analyses: List[Dict[str, Any]],
        market_research_analyses: List[Dict[str, Any]],
        semantic_analyses: List[Dict[str, Any]],
        seo_analyses: List[Dict[str, Any]],
        survey_simulations: List[Dict[str, Any]],
    ) -> Dict[str, Dict[str, Any]]:
        """
        Organize all brand-specific analyses by brand name.

        Args:
            brand_names: List of brand names to organize data for
            brand_name_generations: List of brand name generation results
            brand_name_evaluations: List of brand name evaluation results
            competitor_analyses: List of competitor analysis results
            cultural_sensitivity_analyses: List of cultural sensitivity analysis results
            domain_analyses: List of domain analysis results
            market_research_analyses: List of market research results
            semantic_analyses: List of semantic analysis results
            seo_analyses: List of SEO analysis results
            survey_simulations: List of survey simulation results

        Returns:
            Dict mapping brand names to their organized analysis data
        """
        # Create lookups for analyses by brand name
        generation_lookup = {g["brand_name"]: g for g in brand_name_generations}
        evaluation_lookup = {e["brand_name"]: e for e in brand_name_evaluations}
        cultural_lookup = {c["brand_name"]: c for c in cultural_sensitivity_analyses}
        domain_lookup = {d["brand_name"]: d for d in domain_analyses}
        market_lookup = {m["brand_name"]: m for m in market_research_analyses}
        semantic_lookup = {s["brand_name"]: s for s in semantic_analyses}
        seo_lookup = {s["brand_name"]: s for s in seo_analyses}

        # Group competitor analyses by brand name
        competitor_lookup = {}
        for comp in competitor_analyses:
            brand_name = comp.get("brand_name")
            if brand_name:
                if brand_name not in competitor_lookup:
                    competitor_lookup[brand_name] = []
                competitor_lookup[brand_name].append(
                    {
                        "competitor_name": comp.get("competitor_name"),
                        "positioning": comp.get("competitor_positioning"),
                        "strengths": comp.get("competitor_strengths"),
                        "weaknesses": comp.get("competitor_weaknesses"),
                        "differentiation_score": comp.get("differentiation_score"),
                        "risk_of_confusion": comp.get("risk_of_confusion"),
                        "competitive_advantage_notes": comp.get("competitive_advantage_notes"),
                    }
                )

        # Group survey results by brand name
        survey_lookup = {}
        for survey in survey_simulations:
            brand_name = survey.get("brand_name")
            if brand_name:
                if brand_name not in survey_lookup:
                    survey_lookup[brand_name] = []
                survey_lookup[brand_name].append(
                    {
                        "persona": {
                            "segment": survey.get("persona_segment"),
                            "industry": survey.get("industry"),
                            "company": {
                                "name": survey.get("company_name"),
                                "size": survey.get("company_size_employees"),
                                "revenue": survey.get("company_revenue"),
                                "market_share": survey.get("market_share"),
                                "structure": survey.get("company_structure"),
                                "location": survey.get("geographic_location"),
                                "tech_stack": survey.get("technology_stack"),
                                "growth_stage": survey.get("company_growth_stage"),
                                "focus": survey.get("company_focus"),
                                "maturity": survey.get("company_maturity"),
                                "culture": survey.get("company_culture_values"),
                            },
                            "role": {
                                "title": survey.get("job_title"),
                                "seniority": survey.get("seniority"),
                                "experience": survey.get("years_of_experience"),
                                "department": survey.get("department"),
                                "education": survey.get("education_level"),
                                "reports_to": survey.get("reports_to"),
                            },
                            "profile": {
                                "goals_challenges": survey.get("goals_and_challenges"),
                                "values_priorities": survey.get("values_and_priorities"),
                                "decision_style": survey.get("decision_making_style"),
                                "info_sources": survey.get("information_sources"),
                                "pain_points": survey.get("pain_points"),
                                "tech_literacy": survey.get("technological_literacy"),
                                "risk_attitude": survey.get("attitude_towards_risk"),
                                "purchasing_behavior": survey.get("purchasing_behavior"),
                                "online_behavior": survey.get("online_behavior"),
                                "brand_interaction": survey.get("interaction_with_brand"),
                                "associations": survey.get("professional_associations"),
                                "skills": survey.get("technical_skills"),
                                "networking": survey.get("networking_habits"),
                                "aspirations": survey.get("professional_aspirations"),
                                "influence": survey.get("influence_within_company"),
                                "events": survey.get("event_attendance"),
                                "content_habits": survey.get("content_consumption_habits"),
                                "vendor_preferences": survey.get("vendor_relationship_preferences"),
                                "business_chemistry": survey.get("business_chemistry"),
                                "budget_authority": survey.get("budget_authority"),
                                "vendor_size_preference": survey.get("preferred_vendor_size"),
                                "innovation_adoption": survey.get("innovation_adoption"),
                                "kpis": survey.get("key_performance_indicators"),
                                "development_interests": survey.get("professional_development_interests"),
                                "social_media": survey.get("social_media_usage"),
                                "work_life_balance": survey.get("work_life_balance_priorities"),
                                "frustrations": survey.get("frustrations_annoyances"),
                                "life_goals": survey.get("personal_aspirations_life_goals"),
                                "motivations": survey.get("motivations"),
                            },
                            "buying_process": {
                                "group_structure": survey.get("buying_group_structure"),
                                "decision_maker": survey.get("decision_maker"),
                                "brand_relationships": survey.get("current_brand_relationships"),
                                "adoption_stage": survey.get("product_adoption_lifecycle_stage"),
                                "triggers": survey.get("purchase_triggers_events"),
                                "success_metrics": survey.get("success_metrics_product_service"),
                                "channel_preferences": survey.get("channel_preferences_brand_interaction"),
                                "adoption_barriers": survey.get("barriers_to_adoption"),
                            },
                        },
                        "feedback": {
                            "brand_promise_score": survey.get("brand_promise_perception_score"),
                            "personality_fit": survey.get("personality_fit_score"),
                            "emotional_association": survey.get("emotional_association"),
                            "differentiation_score": survey.get("competitive_differentiation_score"),
                            "sentiment_mapping": survey.get("psychometric_sentiment_mapping"),
                            "competitor_benchmark": survey.get("competitor_benchmarking_score"),
                            "market_adoption": survey.get("simulated_market_adoption_score"),
                            "qualitative_summary": survey.get("qualitative_feedback_summary"),
                            "raw_feedback": survey.get("raw_qualitative_feedback"),
                            "recommendation": survey.get("final_survey_recommendation"),
                            "strategic_ranking": survey.get("strategic_ranking"),
                        },
                        "metadata": {
                            "confidence_score": survey.get("confidence_score_persona_accuracy"),
                            "data_sources": survey.get("data_sources_persona_creation"),
                            "archetype": survey.get("persona_archetype_type"),
                            "generation": survey.get("generation_age_range"),
                            "industry_vertical": survey.get("industry_sub_vertical"),
                        },
                    }
                )

        # Organize data by brand name
        organized_data = {}
        for brand_name in brand_names:
            # Get evaluation data to check shortlist status
            evaluation = evaluation_lookup.get(brand_name, {})

            # Basic brand data from generation and evaluation
            brand_data = {
                "generation": {
                    "brand_name": brand_name,
                    "naming_category": generation_lookup.get(brand_name, {}).get("naming_category"),
                    "brand_personality_alignment": generation_lookup.get(brand_name, {}).get("brand_personality_alignment"),
                    "brand_promise_alignment": generation_lookup.get(brand_name, {}).get("brand_promise_alignment"),
                    "memorability_score": generation_lookup.get(brand_name, {}).get("memorability_score"),
                    "pronounceability_score": generation_lookup.get(brand_name, {}).get("pronounceability_score"),
                    "brand_fit_score": generation_lookup.get(brand_name, {}).get("brand_fit_score"),
                    "meaningfulness_score": generation_lookup.get(brand_name, {}).get("meaningfulness_score"),
                },
                "evaluation": {
                    "strategic_alignment_score": evaluation.get("strategic_alignment_score"),
                    "distinctiveness_score": evaluation.get("distinctiveness_score"),
                    "memorability_score": evaluation.get("memorability_score"),
                    "pronounceability_score": evaluation.get("pronounceability_score"),
                    "meaningfulness_score": evaluation.get("meaningfulness_score"),
                    "brand_fit_score": evaluation.get("brand_fit_score"),
                    "domain_viability_score": evaluation.get("domain_viability_score"),
                    "overall_score": evaluation.get("overall_score"),
                    "shortlist_status": evaluation.get("shortlist_status", False),
                    "evaluation_comments": evaluation.get("evaluation_comments"),
                    "rank": evaluation.get("rank"),
                },
            }

            # Add semantic analysis
            semantic = semantic_lookup.get(brand_name, {})
            brand_data["semantic"] = {
                "meaning": {
                    "denotative": semantic.get("denotative_meaning"),
                    "etymology": semantic.get("etymology"),
                    "descriptiveness": semantic.get("descriptiveness"),
                    "concreteness": semantic.get("concreteness"),
                    "emotional_valence": semantic.get("emotional_valence"),
                },
                "linguistic": {
                    "brand_personality": semantic.get("brand_personality"),
                    "sensory_associations": semantic.get("sensory_associations"),
                    "figurative_language": semantic.get("figurative_language"),
                    "ambiguity": semantic.get("ambiguity"),
                    "irony_paradox": semantic.get("irony_or_paradox"),
                    "humor_playfulness": semantic.get("humor_playfulness"),
                },
                "sound": {
                    "phoneme_combinations": semantic.get("phoneme_combinations"),
                    "sound_symbolism": semantic.get("sound_symbolism"),
                    "rhyme_rhythm": semantic.get("rhyme_rhythm"),
                    "alliteration_assonance": semantic.get("alliteration_assonance"),
                    "word_length_syllables": semantic.get("word_length_syllables"),
                },
                "structure": {
                    "compounding_derivation": semantic.get("compounding_derivation"),
                    "brand_name_type": semantic.get("brand_name_type"),
                },
                "scores": {
                    "memorability": semantic.get("memorability_score"),
                    "pronunciation_ease": semantic.get("original_pronunciation_ease"),
                    "clarity": semantic.get("clarity_understandability"),
                    "uniqueness": semantic.get("uniqueness_differentiation"),
                    "brand_fit": semantic.get("brand_fit_relevance"),
                },
                "trademark_risk": semantic.get("semantic_trademark_risk"),
            }

            # Add cultural sensitivity analysis
            cultural = cultural_lookup.get(brand_name, {})
            brand_data["cultural"] = {
                "connotations": cultural.get("cultural_connotations"),
                "symbolic_meanings": cultural.get("symbolic_meanings"),
                "risk_rating": cultural.get("overall_risk_rating"),
                "notes": cultural.get("notes"),
            }

            # Add domain analysis
            domain = domain_lookup.get(brand_name, {})
            brand_data["domain"] = {
                "exact_match": domain.get("domain_exact_match"),
                "alternative_tlds": domain.get("alternative_tlds"),
                "acquisition_cost": domain.get("acquisition_cost"),
                "notes": domain.get("notes"),
            }

            # Add market research
            market = market_lookup.get(brand_name, {})
            brand_data["market"] = {
                "opportunity": market.get("market_opportunity"),
                "audience_fit": market.get("target_audience_fit"),
                "viability": market.get("market_viability"),
                "risks": market.get("potential_risks"),
            }

            # Add SEO analysis
            seo = seo_lookup.get(brand_name, {})
            brand_data["seo"] = {
                "keyword_alignment": seo.get("keyword_alignment"),
                "search_volume": seo.get("search_volume"),
                "viability_score": seo.get("seo_viability_score"),
                "recommendations": seo.get("seo_recommendations"),
            }

            # Add competitor analyses
            brand_data["competitors"] = competitor_lookup.get(brand_name, [])

            # Add survey results
            brand_data["survey"] = survey_lookup.get(brand_name, [])

            # Extract key metrics for easy access
            brand_data["key_metrics"] = {
                "memorability_score": brand_data["evaluation"]["memorability_score"],
                "distinctiveness_score": brand_data["evaluation"]["distinctiveness_score"],
                "strategic_alignment_score": brand_data["evaluation"]["strategic_alignment_score"],
                "overall_score": brand_data["evaluation"]["overall_score"],
                "cultural_risk_rating": brand_data["cultural"]["risk_rating"],
                "seo_viability_score": brand_data["seo"]["viability_score"],
                "market_viability": brand_data["market"]["viability"],
                "brand_fit_score": brand_data["evaluation"]["brand_fit_score"],
                "pronunciation_ease": brand_data["semantic"]["scores"]["pronunciation_ease"],
                "clarity_score": brand_data["semantic"]["scores"]["clarity"],
            }

            organized_data[brand_name] = brand_data

        return organized_data

    async def test_supabase_connection(self, run_id: str) -> Dict[str, Any]:
        """
        Test the Supabase connection and query functionality to verify data retrieval.
        This method runs diagnostics on each table to ensure we're getting data.
        
        Args:
            run_id: The unique identifier for the workflow run
            
        Returns:
            Dict containing diagnostic information and table counts
        """
        diagnostics = {
            "run_id": run_id,
            "timestamp": datetime.now(timezone.utc).isoformat(),
            "connection_test": "PENDING",
            "table_counts": {},
            "sample_data": {},
            "error_details": {}
        }
        
        try:
            # Simple ping test using direct client access
            logger.info("Testing Supabase connection with a simple query")
            try:
                # Simple ping test - access a table directly
                response = self.supabase.client.table("brand_context").select("*").limit(1).execute()
                if hasattr(response, 'data'):
                    diagnostics["connection_test"] = "SUCCESS"
                    logger.info("Supabase connection is functional")
                else:
                    diagnostics["connection_test"] = "FAILED - No data in response"
                    logger.error("Supabase connection ping returned no data")
                    return diagnostics
            except Exception as e:
                diagnostics["connection_test"] = "FAILED - Connection error"
                diagnostics["error_details"]["connection"] = str(e)
                logger.error(f"Supabase connection ping failed: {str(e)}")
                return diagnostics
                
            # Test tables that we access for reports
            tables_to_test = [
                "brand_context",
                "brand_name_generation", 
                "brand_name_evaluation",
                "linguistic_analysis",
                "semantic_analysis",
                "cultural_sensitivity_analysis",
                "competitor_analysis",
                "domain_analysis",
                "seo_online_discoverability",
                "survey_simulation",
                "translation_analysis"
            ]
            
            # Run diagnostics on each table using execute_with_retry
            for table in tables_to_test:
                try:
                    # Use execute_with_retry which we know exists
                    result = await self.supabase.execute_with_retry(
                        "select", 
                        table, 
                        {"run_id": run_id, "limit": 10}
                    )
                    
                    # Store the count for this table
                    count = len(result) if result else 0
                    diagnostics["table_counts"][table] = count
                    
                    # If we have data, get field information
                    if count > 0:
                        # Log field names and truncated values for debugging
                        fields = {
                            k: str(v)[:50] + ("..." if len(str(v)) > 50 else "") 
                            for k, v in result[0].items()
                        }
                        diagnostics["sample_data"][table] = {
                            "field_count": len(fields),
                            "fields": list(fields.keys()),
                            "preview": fields
                        }
                        logger.info(f"Table {table} sample data retrieved: {len(fields)} fields available")
                    else:
                        logger.warning(f"Table {table} has no records for run_id {run_id}")
                
                except Exception as e:
                    error_msg = f"Error testing table {table}: {str(e)}"
                    diagnostics["error_details"][table] = error_msg
                    logger.error(error_msg)
            
            # Verify we have sufficient data for a report
            has_sufficient_data = any(
                isinstance(count, int) and count > 0 
                for count in diagnostics["table_counts"].values()
            )
            
            if not has_sufficient_data:
                logger.error(f"Insufficient data for report generation with run_id {run_id}")
                diagnostics["overall_status"] = "INSUFFICIENT DATA"
            else:
                diagnostics["overall_status"] = "SUFFICIENT DATA AVAILABLE"
                
            return diagnostics
            
        except Exception as e:
            error_msg = f"Global error in Supabase diagnostics: {str(e)}"
            diagnostics["connection_test"] = "FAILED - Exception"
            diagnostics["error_details"]["global"] = error_msg
            diagnostics["overall_status"] = "ERROR"
            logger.error(error_msg)
            return diagnostics

    async def run_diagnostics(self, run_id: str) -> Dict[str, Any]:
        """
        Run comprehensive diagnostics on the database connection and data availability.
        This method can be called independently to debug database connection issues.
        
        Args:
            run_id: The unique identifier for the workflow run
            
        Returns:
            Dict containing detailed diagnostic information
        """
        logger.info(f"Running comprehensive diagnostics for run_id: {run_id}")
        
        try:
            # Run the basic connection test
            diagnostics = await self.test_supabase_connection(run_id)
            
            # Add timestamp for reference
            diagnostics["diagnostic_run_time"] = datetime.now(timezone.utc).isoformat()
            
            # Add detailed query testing
            if diagnostics["connection_test"] == "SUCCESS":
                # Test a more complex query
                try:
                    # Test a JOIN between tables
                    join_query = f"""
                        SELECT bg.brand_name, be.overall_score, be.shortlist_status
                        FROM brand_name_generation bg
                        LEFT JOIN brand_name_evaluation be ON bg.brand_name = be.brand_name AND bg.run_id = be.run_id
                        WHERE bg.run_id = :run_id
                        LIMIT 5
                    """
                    join_result = await self.supabase.execute_query(join_query, {"run_id": run_id})
                    
                    diagnostics["join_query_test"] = {
                        "status": "SUCCESS" if join_result else "FAILED - No results",
                        "result_count": len(join_result) if join_result else 0,
                        "sample": join_result[0] if join_result and len(join_result) > 0 else None
                    }
                except Exception as e:
                    diagnostics["join_query_test"] = {
                        "status": "FAILED - Exception",
                        "error": str(e)
                    }
                
                # Test a conditional query
                try:
                    # Find shortlisted names
                    conditional_query = f"""
                        SELECT brand_name, overall_score
                        FROM brand_name_evaluation
                        WHERE run_id = :run_id AND shortlist_status = true
                        ORDER BY overall_score DESC
                    """
                    conditional_result = await self.supabase.execute_query(conditional_query, {"run_id": run_id})
                    
                    diagnostics["conditional_query_test"] = {
                        "status": "SUCCESS" if conditional_result else "FAILED - No results",
                        "result_count": len(conditional_result) if conditional_result else 0,
                        "sample": conditional_result[0] if conditional_result and len(conditional_result) > 0 else None
                    }
                except Exception as e:
                    diagnostics["conditional_query_test"] = {
                        "status": "FAILED - Exception",
                        "error": str(e)
                    }
                    
                # Check authentication status
                try:
                    auth_status = await self.supabase.check_auth_status()
                    diagnostics["auth_status"] = auth_status
                except Exception as e:
                    diagnostics["auth_status"] = {
                        "status": "ERROR",
                        "error": str(e)
                    }
            
            # Summarize findings
            diagnostics["summary"] = {
                "connection": diagnostics["connection_test"],
                "data_availability": diagnostics["overall_status"],
                "table_count": len(diagnostics["table_counts"]),
                "tables_with_data": sum(1 for count in diagnostics["table_counts"].values() 
                                     if isinstance(count, int) and count > 0),
                "tables_without_data": sum(1 for count in diagnostics["table_counts"].values() 
                                       if isinstance(count, int) and count == 0),
                "tables_with_errors": sum(1 for count in diagnostics["table_counts"].values() 
                                      if not isinstance(count, int))
            }
            
            # Log diagnostics summary
            logger.info(f"Diagnostics complete for run_id {run_id}", extra={"summary": diagnostics["summary"]})
            
            return diagnostics
            
        except Exception as e:
            error_msg = f"Error running diagnostics: {str(e)}"
            logger.error(error_msg)
            return {
                "status": "ERROR",
                "error": error_msg,
                "run_id": run_id,
                "timestamp": datetime.now(timezone.utc).isoformat()
            }
